import pytest
from docx import Document
from docx.oxml import OxmlElement

from tvba_core_title import (
    identify_numeric_title_level,
    identify_level_from_number,
    normalize_number_string,
    apply_title_style,
    auto_detect_and_format,
    identify_list_item,
    identify_chinese_title,
    _strip_list_prefix,
)
from tvba_settings import TitleLevelSettings, BodySettings, FormatSettings


class TestNormalizeNumberString:
    def test_fullwidth_dot_to_halfwidth(self):
        assert normalize_number_string("1．2．3") == "1.2.3"

    def test_removes_trailing_dot(self):
        assert normalize_number_string("1.2.") == "1.2"

    def test_strips_whitespace(self):
        assert normalize_number_string("  1.1  ") == "1.1"

    def test_multiple_fullwidth_dots(self):
        assert normalize_number_string("１．２．３．４") == "1.2.3.4"


class TestIdentifyLevelFromNumber:
    def test_level_1_no_dot(self):
        assert identify_level_from_number("1") == 1

    def test_level_1_with_dot_zero(self):
        assert identify_level_from_number("1.0") == 1

    def test_level_2_one_dot(self):
        assert identify_level_from_number("1.1") == 2

    def test_level_3_two_dots(self):
        assert identify_level_from_number("1.1.2") == 3

    def test_level_4_three_dots(self):
        assert identify_level_from_number("1.1.2.3") == 4

    def test_level_5_four_dots(self):
        assert identify_level_from_number("1.1.2.3.4") == 5

    def test_level_0_too_many_dots(self):
        assert identify_level_from_number("1.1.2.3.4.5") == 0

    def test_level_0_empty(self):
        assert identify_level_from_number("") == 0

    def test_level_0_list_marker(self):
        assert identify_level_from_number("1）") == 0
        assert identify_level_from_number("(1)") == 0


class TestIdentifyNumericTitleLevel:
    def test_level_1_simple(self):
        assert identify_numeric_title_level("1 引言") == 1

    def test_level_2_simple(self):
        assert identify_numeric_title_level("1.1 背景") == 2

    def test_level_3_simple(self):
        assert identify_numeric_title_level("1.1.1 详细背景") == 3

    def test_no_number_returns_0(self):
        assert identify_numeric_title_level("引言") == 0

    def test_number_without_space_returns_0(self):
        assert identify_numeric_title_level("1引言") == 0

    def test_fullwidth_dots_work(self):
        assert identify_numeric_title_level("1．1 背景") == 2

    def test_tab_separator_works(self):
        assert identify_numeric_title_level("1.1\t背景") == 2

    def test_trailing_dot_normalized(self):
        assert identify_numeric_title_level("1. 引言") == 1

    def test_level_1_requires_space_or_tab(self):
        assert identify_numeric_title_level("1引言") == 0
        assert identify_numeric_title_level("1 引言") == 1

    def test_too_many_dots_returns_0(self):
        assert identify_numeric_title_level("1.1.2.3.4.5 太深") == 0


    def test_plain_integer_quantity_sentence_returns_0(self):
        text = (
            "21 \u4e2a\u5730\u5e02\u7684\u57fa\u7840\u4e0a\uff0c"
            "\u52a0\u5feb\u5efa\u8bbe\u5929\u7136\u6c14\u4e3b\u5e72\u7ba1\u9053"
            "\u201c\u53bf\u53bf\u901a\u5de5\u7a0b\u201d\uff0c"
            "\u52302022\u5e74\uff0c\u57fa\u672c\u5b9e\u73b0\u5929\u7136\u6c14"
            "\u4e3b\u5e72\u7ba1\u9053\u201c\u53bf\u53bf\u901a\u201d\u3002"
        )
        assert identify_numeric_title_level(text) == 0

    def test_plain_integer_heading_still_matches(self):
        assert identify_numeric_title_level("20 \u7ba1\u9053\u5de5\u7a0b\u7ef4\u4fee\u3001\u62a2\u4fee") == 1


class TestApplyTitleStyle:
    def test_applies_outline_level(self):
        doc = Document()
        para = doc.add_paragraph("1 标题")
        settings = TitleLevelSettings(font="黑体", size="三号", bold=True)
        body = BodySettings()
        apply_title_style(para, 1, settings, body)
        pPr = para._element.find(".//w:pPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        outline = pPr.find("w:outlineLvl", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        assert outline is not None
        assert outline.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "0"

    def test_applies_font_and_bold(self):
        doc = Document()
        para = doc.add_paragraph("1 标题")
        settings = TitleLevelSettings(font="黑体", size="三号", bold=True)
        body = BodySettings()
        apply_title_style(para, 1, settings, body)
        run = para.runs[0]
        assert run.font.bold is True

    def test_applies_center_alignment(self):
        doc = Document()
        para = doc.add_paragraph("1 标题")
        settings = TitleLevelSettings(alignment="居中")
        body = BodySettings()
        apply_title_style(para, 1, settings, body)
        assert para.alignment == 1


    def test_applies_title_indent_and_normalizes_brackets(self):
        doc = Document()
        para = doc.add_paragraph("1.1.1.1 \u56db\u7ea7\u6807\u9898(\u8bd5\u8fd0\u884c)")
        settings = TitleLevelSettings(
            alignment="\u5de6\u5bf9\u9f50",
            left_indent_chars=2.0,
            special_indent="\u65e0",
            normalize_brackets=True,
        )
        body = BodySettings(modify_content=False)

        apply_title_style(para, 4, settings, body)

        assert "\uff08\u8bd5\u8fd0\u884c\uff09" in para.text
        pPr = para._element.find(".//w:pPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        ind = pPr.find("w:ind", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        assert ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}leftChars") == "200"
        assert ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left") is None
        assert ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}firstLine") == "0"
        assert ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hanging") is None

    def test_no_special_indent_overrides_inherited_first_line(self):
        doc = Document()
        para = doc.add_paragraph("1.1 前言")
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        pPr = para._element.get_or_add_pPr()
        ind = pPr.find(f"{{{W}}}ind")
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(f"{{{W}}}firstLine", "198")

        settings = TitleLevelSettings(special_indent="无", special_indent_chars=0.0)
        apply_title_style(para, 2, settings, BodySettings())

        ind = para._element.find(".//w:ind", {"w": W})
        assert ind.get(f"{{{W}}}firstLine") == "0"
        assert ind.get(f"{{{W}}}firstLineChars") == "0"
        assert ind.get(f"{{{W}}}hanging") is None
        assert ind.get(f"{{{W}}}hangingChars") is None


class TestAutoDetectAndFormat:
    def test_detects_and_formats_titles(self):
        doc = Document()
        doc.add_paragraph("1 一级标题")
        doc.add_paragraph("1.1 二级标题")
        doc.add_paragraph("正文段落")
        settings = FormatSettings()
        auto_detect_and_format(doc, settings)

        # Check first paragraph has outline level 0 (level 1)
        p1 = doc.paragraphs[0]
        pPr = p1._element.find(".//w:pPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        outline = pPr.find("w:outlineLvl", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        assert outline.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "0"

        # Check second paragraph has outline level 1 (level 2)
        p2 = doc.paragraphs[1]
        pPr2 = p2._element.find(".//w:pPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        outline2 = pPr2.find("w:outlineLvl", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        assert outline2.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "1"

        # Check body paragraph has no outline level
        p3 = doc.paragraphs[2]
        pPr3 = p3._element.find(".//w:pPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        if pPr3 is not None:
            outline3 = pPr3.find("w:outlineLvl", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            assert outline3 is None
        # If pPr3 is None, body paragraph has no formatting at all — also correct

    def test_does_not_overwrite_existing_outline_level(self):
        """Paragraphs that already have an outline level should keep it."""
        from tvba_core_oox import set_outline_level
        doc = Document()
        para = doc.add_paragraph("1.1 研究方法")
        # User manually set this as a level-1 title (outline level 0) in Word
        set_outline_level(para, 0)
        settings = FormatSettings()
        auto_detect_and_format(doc, settings)
        # Outline level should remain 0, not be overwritten to 1 (level 2)
        pPr = para._element.find(".//w:pPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        outline = pPr.find("w:outlineLvl", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        assert outline.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "0"

    def test_list_paragraph_not_treated_as_title(self):
        """Body list items like '1）第一项' should NOT be treated as titles."""
        from lxml import etree
        from tvba_core_numbering import DocxListResolver

        doc = Document()
        para = doc.add_paragraph("1）第一项")
        # Simulate Word multilevel list formatting (w:numPr/w:ilvl=0)
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        pPr = para._element.find(f"{{{W}}}pPr")
        if pPr is None:
            pPr = etree.SubElement(para._element, f"{{{W}}}pPr")
        numPr = etree.SubElement(pPr, f"{{{W}}}numPr")
        ilvl = etree.SubElement(numPr, f"{{{W}}}ilvl")
        ilvl.set(f"{{{W}}}val", "0")

        settings = FormatSettings()
        resolver = DocxListResolver(doc)
        auto_detect_and_format(doc, settings, resolver)

        # Should NOT have outline level set
        outline = pPr.find(f"{{{W}}}outlineLvl")
        assert outline is None, f"List paragraph '1）第一项' should not be treated as title, got outline={outline}"

    def test_plain_large_document_does_not_call_list_resolver(self):
        """Regression guard: ordinary body paragraphs must not trigger COM/list calls.

        The GUI used to appear stuck at "Detecting titles..." on large documents
        because title detection queried the list resolver for every paragraph.
        Plain paragraphs have no numbering hint, so they should be skipped before
        any expensive resolver/COM call.
        """

        class CountingResolver:
            def __init__(self):
                self.level_calls = 0
                self.text_calls = 0

            def get_list_level(self, para):
                self.level_calls += 1
                return None

            def get_list_text(self, para):
                self.text_calls += 1
                return None

        doc = Document()
        for i in range(500):
            doc.add_paragraph(f"Plain body paragraph {i}")

        resolver = CountingResolver()
        auto_detect_and_format(doc, FormatSettings(), resolver)

        assert resolver.level_calls == 0
        assert resolver.text_calls == 0

    def test_only_numbered_hint_paragraph_calls_list_resolver(self):
        """Only paragraphs with w:numPr should reach the expensive resolver path."""
        from lxml import etree

        class CountingResolver:
            def __init__(self):
                self.level_calls = 0
                self.text_calls = 0

            def get_list_level(self, para):
                self.level_calls += 1
                return 2

            def get_list_text(self, para):
                self.text_calls += 1
                return "1.1"

        doc = Document()
        doc.add_paragraph("Plain body paragraph")
        numbered = doc.add_paragraph("Rendered list heading")

        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        pPr = numbered._element.get_or_add_pPr()
        numPr = etree.SubElement(pPr, f"{{{W}}}numPr")
        ilvl = etree.SubElement(numPr, f"{{{W}}}ilvl")
        ilvl.set(f"{{{W}}}val", "1")
        numId = etree.SubElement(numPr, f"{{{W}}}numId")
        numId.set(f"{{{W}}}val", "1")

        resolver = CountingResolver()
        auto_detect_and_format(doc, FormatSettings(), resolver)

        assert resolver.level_calls == 1
        assert resolver.text_calls == 1


class TestIdentifyListItem:
    """Legacy list markers must not be treated as title levels."""

    def test_level4_bracket_number(self):
        assert identify_list_item("(1) 适用范围") == 0

    def test_level4_fullwidth_bracket(self):
        assert identify_list_item("1）适用范围") == 0

    def test_level4_parenthesis(self):
        assert identify_list_item("1) 适用范围") == 0

    def test_level4_fullwidth_double(self):
        assert identify_list_item("（1）适用范围") == 0

    def test_level4_chinese_separator(self):
        assert identify_list_item("1、适用范围") == 0

    def test_level5_letter_dot(self):
        assert identify_list_item("a. 说明") == 0

    def test_level5_letter_paren(self):
        assert identify_list_item("a) 说明") == 0

    def test_level5_letter_fullwidth(self):
        assert identify_list_item("a）说明") == 0

    def test_level5_letter_chinese_sep(self):
        assert identify_list_item("a、说明") == 0

    def test_not_title_plain_text(self):
        assert identify_list_item("普通正文") == 0

    def test_not_title_number_in_middle(self):
        # "1)kg" shouldn't match because it has no space after prefix
        assert identify_list_item("1)kg") == 0

    def test_not_title_year(self):
        assert identify_list_item("2026.05") == 0


class TestStripListPrefix:
    def test_level4_bracket(self):
        assert _strip_list_prefix("(1) 适用范围内容", 4) == ""

    def test_level4_fullwidth(self):
        assert _strip_list_prefix("1）适用范围", 4) == ""

    def test_level5_dot(self):
        assert _strip_list_prefix("a. 说明内容", 5) == ""


class TestAutoDetectLevel45:
    """Verify level 4/5 dotted-number detection in auto_detect_and_format."""

    def test_level4_bracket_does_not_apply_title_style(self):
        doc = Document()
        para = doc.add_paragraph("(1) 适用范围说明")
        settings = FormatSettings()
        auto_detect_and_format(doc, settings)
        from tvba_core_oox import get_effective_outline_level
        ol = get_effective_outline_level(para)
        assert ol is None

    def test_level4_fullwidth_bracket_does_not_apply_title_style(self):
        doc = Document()
        para = doc.add_paragraph("1）适用范围说明")
        settings = FormatSettings()
        auto_detect_and_format(doc, settings)
        from tvba_core_oox import get_effective_outline_level
        ol = get_effective_outline_level(para)
        assert ol is None

    def test_level5_letter_does_not_apply_title_style(self):
        doc = Document()
        para = doc.add_paragraph("a. 分项说明")
        settings = FormatSettings()
        auto_detect_and_format(doc, settings)
        from tvba_core_oox import get_effective_outline_level
        ol = get_effective_outline_level(para)
        assert ol is None

    def test_level4_dotted_number_applies_title_style(self):
        doc = Document()
        para = doc.add_paragraph("1.1.1.1 适用范围说明")
        settings = FormatSettings()
        auto_detect_and_format(doc, settings)
        from tvba_core_oox import get_effective_outline_level
        ol = get_effective_outline_level(para)
        assert ol == 3

    def test_level5_dotted_number_applies_title_style(self):
        doc = Document()
        para = doc.add_paragraph("1.1.1.1.1 适用范围说明")
        settings = FormatSettings()
        auto_detect_and_format(doc, settings)
        from tvba_core_oox import get_effective_outline_level
        ol = get_effective_outline_level(para)
        assert ol == 4

    def test_body_like_1kg_not_detected(self):
        """'1)kg' should not be treated as a title."""
        doc = Document()
        para = doc.add_paragraph("1)kg")
        settings = FormatSettings()
        auto_detect_and_format(doc, settings)
        from tvba_core_oox import get_effective_outline_level
        ol = get_effective_outline_level(para)
        assert ol is None

    def test_short_prefix_not_detected(self):
        """'a. xy' too short, no Chinese, should not be a title."""
        doc = Document()
        para = doc.add_paragraph("a. xy")
        settings = FormatSettings()
        auto_detect_and_format(doc, settings)
        from tvba_core_oox import get_effective_outline_level
        ol = get_effective_outline_level(para)
        assert ol is None


class TestAutoNumberedListFallback:
    """Auto-numbered paragraphs whose text lacks the numbering prefix."""

    def _make_auto_numbered_para(self, doc, text):
        """Create a paragraph with w:numPr (simulating Word auto-numbering)."""
        from lxml import etree
        para = doc.add_paragraph(text)
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        pPr = para._element.find(f"{{{W}}}pPr")
        if pPr is None:
            pPr = etree.SubElement(para._element, f"{{{W}}}pPr")
        numPr = etree.SubElement(pPr, f"{{{W}}}numPr")
        ilvl = etree.SubElement(numPr, f"{{{W}}}ilvl")
        ilvl.set(f"{{{W}}}val", "3")  # list level 4 (0-indexed)
        numId = etree.SubElement(numPr, f"{{{W}}}numId")
        numId.set(f"{{{W}}}val", "1")
        return para

    class _MockDocxResolver:
        """Docx resolver that returns a valid level but has no get_list_text."""
        def get_list_level(self, para):
            return 4
        # No get_list_text attribute → docx fallback path

    def test_fallback_rejects_auto_numbered_level4_with_chinese_without_list_text(self):
        """Docx fallback must not infer title level from body list numbering."""
        doc = Document()
        para = self._make_auto_numbered_para(
            doc,
            "《国家管网集团广东省天然气管网珠中江干线项目雷电灾害风险评估报告》（终稿），广东省气候中心，2022年9月；"
        )
        settings = FormatSettings()
        resolver = self._MockDocxResolver()
        auto_detect_and_format(doc, settings, list_resolver=resolver)
        from tvba_core_oox import get_effective_outline_level
        ol = get_effective_outline_level(para)
        assert ol is None

    def test_com_list_text_rejects_bracket_marker(self):
        class Resolver:
            def get_list_level(self, para):
                return 4

            def get_list_text(self, para):
                return "1）"

        doc = Document()
        para = self._make_auto_numbered_para(doc, "适用范围说明")
        auto_detect_and_format(doc, FormatSettings(), list_resolver=Resolver())
        from tvba_core_oox import get_effective_outline_level
        assert get_effective_outline_level(para) is None

    def test_com_list_text_accepts_dotted_level4_marker(self):
        class Resolver:
            def get_list_level(self, para):
                return 4

            def get_list_text(self, para):
                return "1.1.1.1"

        doc = Document()
        para = self._make_auto_numbered_para(doc, "适用范围说明")
        auto_detect_and_format(doc, FormatSettings(), list_resolver=Resolver())
        from tvba_core_oox import get_effective_outline_level
        assert get_effective_outline_level(para) == 3

    def test_fallback_rejects_auto_numbered_no_chinese(self):
        """Auto-numbered para without Chinese → not a title (false positive guard)."""
        doc = Document()
        para = self._make_auto_numbered_para(doc, "Some English text only here")
        settings = FormatSettings()
        resolver = self._MockDocxResolver()
        auto_detect_and_format(doc, settings, list_resolver=resolver)
        from tvba_core_oox import get_effective_outline_level
        ol = get_effective_outline_level(para)
        assert ol is None

    def test_fallback_rejects_auto_numbered_too_short(self):
        """Auto-numbered para too short → not a title."""
        doc = Document()
        para = self._make_auto_numbered_para(doc, "短")
        settings = FormatSettings()
        resolver = self._MockDocxResolver()
        auto_detect_and_format(doc, settings, list_resolver=resolver)
        from tvba_core_oox import get_effective_outline_level
        ol = get_effective_outline_level(para)
        assert ol is None
