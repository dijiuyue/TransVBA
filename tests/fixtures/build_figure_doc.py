"""Build a test .docx with figure captions."""
from docx import Document
import sys
from pathlib import Path


def build(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("图 1-1 示例图片")
    doc.add_paragraph("正文段落")
    doc.add_paragraph("图 2-1 另一个图片")
    doc.save(path)


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("figure_test.docx")
    build(out)
    print(f"Built: {out}")
