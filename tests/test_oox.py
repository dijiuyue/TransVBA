import tempfile
from pathlib import Path
import pytest
from docx import Document
from lxml import etree

from tvba_core_oox import (
    set_far_east_font,
    set_ascii_font,
    set_run_font_size,
    set_style_font_size,
    set_outline_level,
    apply_indent_chars,
    apply_paragraph_spacing,
    set_table_layout_window,
    set_table_layout_content,
    set_table_borders,
    set_row_height_at_least,
)

NSMAP = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

class TestSetFarEastFont:
    def test_sets_rFonts_eastAsia_attribute(self):
        doc = Document()
        para = doc.add_paragraph("测试")
        run = para.runs[0]
        set_far_east_font(run, "黑体")
        rPr = run._element.find(".//w:rPr", NSMAP)
        assert rPr is not None
        rFonts = rPr.find("w:rFonts", NSMAP)
        assert rFonts is not None
        assert rFonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia") == "黑体"

class TestSetAsciiFont:
    def test_sets_font_name(self):
        doc = Document()
        para = doc.add_paragraph("Hello")
        run = para.runs[0]
        set_ascii_font(run, "Times New Roman")
        assert run.font.name == "Times New Roman"

class TestSetOutlineLevel:
    def test_sets_outline_level_zero(self):
        doc = Document()
        para = doc.add_paragraph("Title")
        set_outline_level(para, 0)
        pPr = para._element.find(".//w:pPr", NSMAP)
        assert pPr is not None
        outline = pPr.find("w:outlineLvl", NSMAP)
        assert outline is not None
        assert outline.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "0"

    def test_sets_outline_level_four(self):
        doc = Document()
        para = doc.add_paragraph("Title")
        set_outline_level(para, 4)
        pPr = para._element.find(".//w:pPr", NSMAP)
        outline = pPr.find("w:outlineLvl", NSMAP)
        assert outline.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "4"

class TestApplyIndentChars:
    def test_applies_left_indent_in_twips(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        apply_indent_chars(
            para.paragraph_format,
            left_chars=2.0,
            right_chars=0.0,
            special_kind="无",
            special_chars=0.0,
        )
        pPr = para._element.find(".//w:pPr", NSMAP)
        ind = pPr.find("w:ind", NSMAP)
        assert ind is not None
        # 2 chars * 12 points/char * 20 twips/point = 480 twips
        assert ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left") == "480"

    def test_applies_first_line_indent(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        apply_indent_chars(
            para.paragraph_format,
            left_chars=0.0,
            right_chars=0.0,
            special_kind="首行缩进",
            special_chars=2.0,
        )
        pPr = para._element.find(".//w:pPr", NSMAP)
        ind = pPr.find("w:ind", NSMAP)
        assert ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}firstLine") == "480"

    def test_applies_hanging_indent(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        apply_indent_chars(
            para.paragraph_format,
            left_chars=0.0,
            right_chars=0.0,
            special_kind="悬挂缩进",
            special_chars=2.0,
        )
        pPr = para._element.find(".//w:pPr", NSMAP)
        ind = pPr.find("w:ind", NSMAP)
        assert ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hanging") == "480"

class TestApplyParagraphSpacing:
    def test_sets_beforeLines_and_afterLines(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        apply_paragraph_spacing(
            para.paragraph_format,
            before_lines=0.5,
            after_lines=0.5,
        )
        pPr = para._element.find(".//w:pPr", NSMAP)
        spacing = pPr.find("w:spacing", NSMAP)
        assert spacing is not None
        assert spacing.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}beforeLines") == "50"
        assert spacing.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}afterLines") == "50"

    def test_sets_line_spacing(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        apply_paragraph_spacing(
            para.paragraph_format,
            line_spacing=1.5,
        )
        pPr = para._element.find(".//w:pPr", NSMAP)
        spacing = pPr.find("w:spacing", NSMAP)
        assert spacing is not None
        assert spacing.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}line") == "360"
        assert spacing.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lineRule") == "auto"

    def test_cleans_old_spacing_attrs(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        # Ensure pPr exists, then inject dirty spacing attribute
        from lxml import etree
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        pPr = para._element.find(f".//{{{W_NS}}}pPr")
        if pPr is None:
            pPr = etree.SubElement(para._element, f"{{{W_NS}}}pPr")
        sp = etree.SubElement(pPr, f"{{{W_NS}}}spacing")
        sp.set(f"{{{W_NS}}}beforeAutospacing", "1")
        apply_paragraph_spacing(
            para.paragraph_format,
            before_lines=0.0,
            after_lines=0.0,
        )
        spacing = pPr.find(f"{{{W_NS}}}spacing")
        assert spacing.get(f"{{{W_NS}}}beforeAutospacing") == "0"

class TestTableLayout:
    def test_set_window_layout(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        set_table_layout_window(table)
        tblPr = table._element.find(".//w:tblPr", NSMAP)
        assert tblPr is not None
        layout = tblPr.find("w:tblLayout", NSMAP)
        assert layout is not None
        assert layout.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "autofit"

    def test_set_content_layout(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        set_table_layout_content(table)
        tblPr = table._element.find(".//w:tblPr", NSMAP)
        layout = tblPr.find("w:tblLayout", NSMAP)
        assert layout is not None
        assert layout.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "fixed"

class TestTableBorders:
    def test_sets_all_borders(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        set_table_borders(table, line_width_pt=1.5)
        tblPr = table._element.find(".//w:tblPr", NSMAP)
        borders = tblPr.find("w:tblBorders", NSMAP)
        assert borders is not None
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = borders.find(f"w:{side}", NSMAP)
            assert border is not None, f"Missing {side} border"
            # 1.5 pt = 30 half-points
            assert border.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz") == "30"
            assert border.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "single"

class TestRowHeight:
    def test_sets_row_height_at_least(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        row = table.rows[0]
        set_row_height_at_least(row, height_cm=0.7)
        trPr = row._tr.find("w:trPr", NSMAP)
        assert trPr is not None
        trHeight = trPr.find("w:trHeight", NSMAP)
        assert trHeight is not None
        # 0.7 cm = 0.7 * 28.3465 pt * 20 twips/pt = ~397 twips
        assert trHeight.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") is not None
        assert trHeight.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hRule") == "atLeast"


class TestSetRunFontSize:
    def test_sets_both_w_sz_and_w_szCs(self):
        doc = Document()
        para = doc.add_paragraph("测试")
        run = para.runs[0]
        set_run_font_size(run, 12.0)
        rPr = run._element.find(".//w:rPr", NSMAP)
        assert rPr is not None
        sz = rPr.find("w:sz", NSMAP)
        szCs = rPr.find("w:szCs", NSMAP)
        assert sz is not None
        assert szCs is not None
        # 12 pt = 24 half-points
        assert sz.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "24"
        assert szCs.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "24"

    def test_creates_rPr_if_missing(self):
        doc = Document()
        para = doc.add_paragraph("测试")
        run = para.runs[0]
        # Ensure no rPr exists initially
        rPr = run._element.find("w:rPr", NSMAP)
        if rPr is not None:
            run._element.remove(rPr)
        set_run_font_size(run, 10.5)
        rPr = run._element.find(".//w:rPr", NSMAP)
        assert rPr is not None
        sz = rPr.find("w:sz", NSMAP)
        szCs = rPr.find("w:szCs", NSMAP)
        assert sz is not None
        assert szCs is not None
        # 10.5 pt = 21 half-points
        assert sz.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "21"
        assert szCs.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "21"


class TestSetStyleFontSize:
    def test_sets_both_w_sz_and_w_szCs_on_style(self):
        doc = Document()
        style = doc.styles["Normal"]
        set_style_font_size(style, 12.0)
        rPr = style.element.find(".//w:rPr", NSMAP)
        assert rPr is not None
        sz = rPr.find("w:sz", NSMAP)
        szCs = rPr.find("w:szCs", NSMAP)
        assert sz is not None
        assert szCs is not None
        assert sz.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "24"
        assert szCs.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "24"

    def test_creates_rPr_if_missing(self):
        doc = Document()
        style = doc.styles["Heading 1"]
        # Remove existing rPr if any
        rPr = style.element.find("w:rPr", NSMAP)
        if rPr is not None:
            style.element.remove(rPr)
        set_style_font_size(style, 16.0)
        rPr = style.element.find(".//w:rPr", NSMAP)
        assert rPr is not None
        sz = rPr.find("w:sz", NSMAP)
        szCs = rPr.find("w:szCs", NSMAP)
        assert sz is not None
        assert szCs is not None
        # 16 pt = 32 half-points
        assert sz.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "32"
        assert szCs.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "32"
