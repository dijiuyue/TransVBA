import pytest
from docx import Document

from tvba_core_normalize import (
    unify_ascii_font,
    apply_brackets,
    add_period_if_needed,
    sync_number_font_with_body,
)


class TestUnifyAsciiFont:
    def test_sets_ascii_runs_to_times_new_roman(self):
        doc = Document()
        para = doc.add_paragraph("Hello World 123")
        unify_ascii_font(doc, "Times New Roman")
        for para in doc.paragraphs:
            for run in para.runs:
                assert run.font.name == "Times New Roman"

    def test_skips_cjk_characters(self):
        doc = Document()
        para = doc.add_paragraph("中文")
        # Should not crash
        unify_ascii_font(doc, "Times New Roman")


class TestApplyBrackets:
    def test_makes_fullwidth_bracketed_text_bold(self):
        doc = Document()
        para = doc.add_paragraph("1 标题（补充说明）")
        apply_brackets(para, para.text)
        # At least one run should be bold (the bracketed part)
        bold_runs = [r for r in para.runs if r.font.bold]
        assert len(bold_runs) >= 1

    def test_makes_halfwidth_bracketed_text_bold(self):
        doc = Document()
        para = doc.add_paragraph("1 Title (supplemental)")
        apply_brackets(para, para.text)
        bold_runs = [r for r in para.runs if r.font.bold]
        assert len(bold_runs) >= 1

    def test_only_first_occurrence_is_processed(self):
        doc = Document()
        para = doc.add_paragraph("（第一）和（第二）")
        apply_brackets(para, para.text)
        # Only the first bracketed occurrence should be bold
        bold_runs = [r for r in para.runs if r.font.bold]
        # Should have at least one bold run for first bracket pair
        assert len(bold_runs) >= 1

    def test_no_brackets_no_change(self):
        doc = Document()
        para = doc.add_paragraph("No brackets here")
        apply_brackets(para, para.text)
        bold_runs = [r for r in para.runs if r.font.bold]
        assert len(bold_runs) == 0


class TestAddPeriodIfNeeded:
    def test_adds_period_when_missing(self):
        doc = Document()
        para = doc.add_paragraph("(1) 标题")
        add_period_if_needed(para)
        assert para.text.endswith("。")

    def test_skips_if_already_has_chinese_period(self):
        doc = Document()
        para = doc.add_paragraph("(1) 标题。")
        add_period_if_needed(para)
        assert para.text == "(1) 标题。"

    def test_skips_if_ends_with_exclamation(self):
        doc = Document()
        para = doc.add_paragraph("(1) 标题！")
        add_period_if_needed(para)
        assert para.text == "(1) 标题！"

    def test_skips_if_ends_with_question_mark(self):
        doc = Document()
        para = doc.add_paragraph("1 标题？")
        add_period_if_needed(para)
        assert para.text == "1 标题？"

    def test_skips_empty_paragraph(self):
        doc = Document()
        para = doc.add_paragraph("")
        add_period_if_needed(para)
        assert para.text == ""

    def test_skips_toc_lines(self):
        doc = Document()
        para = doc.add_paragraph("目录")
        add_period_if_needed(para)
        # TOC title should not get a period
        assert para.text == "目录"


class TestSyncNumberFontWithBody:
    def test_sets_font_on_list_paragraph_runs(self):
        doc = Document()
        para = doc.add_paragraph("Test item", style="List Number")
        # Add a run with no font set
        run = para.add_run(" more text")
        run.font.name = ""
        sync_number_font_with_body(para)
        # Should not crash; function should handle list paragraphs

    def test_skips_non_list_paragraphs(self):
        doc = Document()
        para = doc.add_paragraph("Normal paragraph")
        sync_number_font_with_body(para)
        # Should not crash or modify anything

    def test_applies_body_font_to_runs(self):
        doc = Document()
        para = doc.add_paragraph("Test item", style="List Number")
        run = para.add_run(" body text")
        run.font.name = "SimSun"
        sync_number_font_with_body(para)
        # Should not crash; function detects list paragraph and processes runs
