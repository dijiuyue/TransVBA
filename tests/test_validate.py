"""TDD tests for validation engine (tvba_core_validate.py).

Requirements ref: 多模板自动格式化与检查插件需求说明书 section 6
- Common checks: 非宋体中文, 非TNR数字/英文, 半角括号, 缺句号
- Mode A: 表格非五号字, 表格行高非0.6cm
- Mode B: 封面非二号, 附件正文非小五, 一级标题未对齐网格, 题注空格, 禁忌词
"""
import pytest
from docx import Document
from lxml import etree

from tvba_core_validate import (
    validate_document,
    ValidationIssue,
    _check_chinese_font,
    _check_ascii_font,
    _check_brackets,
    _check_period,
    _check_forbidden_words,
    _check_table_font_size,
    _check_table_row_height,
    _check_grid_alignment,
    _check_appendix_colon,
    _check_figure_table_space,
    _check_cover_title_size,
    _check_appendix_body_size,
    _check_chairman_number,
    _check_paragraph_spacing,
    _check_caption_alignment,
    _check_table_fixed_dimensions,
)
from tvba_settings import FormatSettings, ValidationRules, TableSettings

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _set_run_fonts(run, east_asia=None, ascii=None):
    """Set fonts on a run at the XML level."""
    rPr = run._element.find(f"{{{W}}}rPr")
    if rPr is None:
        rPr = etree.SubElement(run._element, f"{{{W}}}rPr")
    rFonts = rPr.find(f"{{{W}}}rFonts")
    if rFonts is None:
        rFonts = etree.SubElement(rPr, f"{{{W}}}rFonts")
    if east_asia:
        rFonts.set(f"{{{W}}}eastAsia", east_asia)
    if ascii:
        rFonts.set(f"{{{W}}}ascii", ascii)


def _set_run_size(run, half_pts: int):
    """Set font size on a run at the XML level."""
    rPr = run._element.find(f"{{{W}}}rPr")
    if rPr is None:
        rPr = etree.SubElement(run._element, f"{{{W}}}rPr")
    sz = rPr.find(f"{{{W}}}sz")
    if sz is None:
        sz = etree.SubElement(rPr, f"{{{W}}}sz")
    sz.set(f"{{{W}}}val", str(half_pts))


class TestCheckChineseFont:
    """Section 5: 所有中文字符 强制 宋体"""

    def test_flags_non_song_chinese_font(self):
        doc = Document()
        para = doc.add_paragraph("一段中文测试")
        _set_run_fonts(para.runs[0], east_asia="黑体")

        issues: list[ValidationIssue] = []
        _check_chinese_font(para, issues)
        assert len(issues) > 0
        assert "非宋体" in issues[0].description or "黑体" in issues[0].description

    def test_passes_song_chinese_font(self):
        doc = Document()
        para = doc.add_paragraph("一段中文测试")
        _set_run_fonts(para.runs[0], east_asia="宋体")

        issues: list[ValidationIssue] = []
        _check_chinese_font(para, issues)
        assert len(issues) == 0

    def test_skips_non_chinese_text(self):
        doc = Document()
        para = doc.add_paragraph("English only text")
        _set_run_fonts(para.runs[0], east_asia="楷体")  # wrong font

        issues: list[ValidationIssue] = []
        _check_chinese_font(para, issues)
        # No issues because there's no Chinese text to check
        assert len(issues) == 0

    def test_skips_empty_text(self):
        doc = Document()
        para = doc.add_paragraph("")

        issues: list[ValidationIssue] = []
        _check_chinese_font(para, issues)
        assert len(issues) == 0


class TestCheckAsciiFont:
    """Section 5: 所有阿拉伯数字/英文字母 强制 Times New Roman"""

    def test_flags_non_tnr_ascii_font(self):
        doc = Document()
        para = doc.add_paragraph("Test123")
        _set_run_fonts(para.runs[0], ascii="Arial")

        issues: list[ValidationIssue] = []
        _check_ascii_font(para, issues)
        assert len(issues) > 0
        assert "Times New Roman" in issues[0].description or "Arial" in issues[0].description

    def test_passes_tnr_ascii_font(self):
        doc = Document()
        para = doc.add_paragraph("Test123")
        _set_run_fonts(para.runs[0], ascii="Times New Roman")

        issues: list[ValidationIssue] = []
        _check_ascii_font(para, issues)
        assert len(issues) == 0

    def test_skips_non_ascii_text(self):
        doc = Document()
        para = doc.add_paragraph("纯中文无数字")  # No ASCII letters/digits
        _set_run_fonts(para.runs[0], ascii="Arial")

        issues: list[ValidationIssue] = []
        _check_ascii_font(para, issues)
        assert len(issues) == 0


class TestCheckBrackets:
    """Section 5: 列表项半角括号 () 转全角 （）"""

    def test_flags_halfwidth_brackets(self):
        doc = Document()
        para = doc.add_paragraph("(1) 列表项内容")

        issues: list[ValidationIssue] = []
        _check_brackets(para, issues)
        assert len(issues) > 0
        assert "半角括号" in issues[0].description or "全角" in issues[0].description

    def test_passes_fullwidth_brackets(self):
        doc = Document()
        para = doc.add_paragraph("（1）列表项内容")

        issues: list[ValidationIssue] = []
        _check_brackets(para, issues)
        assert len(issues) == 0

    def test_skips_no_brackets(self):
        doc = Document()
        para = doc.add_paragraph("普通正文没有括号")

        issues: list[ValidationIssue] = []
        _check_brackets(para, issues)
        assert len(issues) == 0


class TestCheckPeriod:
    """Section 5: 列表最后一项缺失句号 → 报错提示"""

    def test_flags_missing_period_in_list_item(self):
        doc = Document()
        para = doc.add_paragraph("(1) 列表项缺少句号")

        issues: list[ValidationIssue] = []
        _check_period(para, issues)
        assert len(issues) > 0
        assert "句号" in issues[0].description

    def test_passes_list_item_with_period(self):
        doc = Document()
        para = doc.add_paragraph("(1) 列表项有句号。")

        issues: list[ValidationIssue] = []
        _check_period(para, issues)
        assert len(issues) == 0

    def test_skips_title_paragraphs(self):
        """标题 '1 概述' 不应检查句号"""
        doc = Document()
        para = doc.add_paragraph("1 项目概述")

        issues: list[ValidationIssue] = []
        _check_period(para, issues)
        assert len(issues) == 0, "Title paragraphs should not be checked for period"

    def test_skips_normal_body_text(self):
        """普通正文不应检查句号"""
        doc = Document()
        para = doc.add_paragraph("这是一段正常的正文内容，不需要强制句号检查。")

        issues: list[ValidationIssue] = []
        _check_period(para, issues)
        assert len(issues) == 0

    def test_detects_numbered_list_item(self):
        """'1） 列表项' 格式也应检测"""
        doc = Document()
        para = doc.add_paragraph("1）列表项缺句号")

        issues: list[ValidationIssue] = []
        _check_period(para, issues)
        assert len(issues) > 0

    def test_detects_letter_list_item(self):
        """'a. 列表项' 格式也应检测"""
        doc = Document()
        para = doc.add_paragraph("a. letter list item")

        issues: list[ValidationIssue] = []
        _check_period(para, issues)
        assert len(issues) > 0


class TestCheckForbiddenWords:
    """Section 4.2 Mode B: 禁止'附图/附表' → 改为'附件'"""

    def test_flags_forbidden_word_futu(self):
        doc = Document()
        para = doc.add_paragraph("详见附图1")

        issues: list[ValidationIssue] = []
        _check_forbidden_words(para, ("附图", "附表"), issues)
        assert len(issues) > 0
        assert "附图" in issues[0].description

    def test_flags_forbidden_word_fubiao(self):
        doc = Document()
        para = doc.add_paragraph("见附表2")

        issues: list[ValidationIssue] = []
        _check_forbidden_words(para, ("附图", "附表"), issues)
        assert len(issues) > 0
        assert "附表" in issues[0].description

    def test_passes_normal_text(self):
        doc = Document()
        para = doc.add_paragraph("详见附件一")

        issues: list[ValidationIssue] = []
        _check_forbidden_words(para, ("附图", "附表"), issues)
        assert len(issues) == 0

    def test_empty_text(self):
        doc = Document()
        para = doc.add_paragraph("")

        issues: list[ValidationIssue] = []
        _check_forbidden_words(para, ("附图", "附表"), issues)
        assert len(issues) == 0


class TestCheckTableFontSize:
    """Section 6 Mode A: 检查表格内非五号字"""

    def test_flags_wrong_table_font_size(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        para = cell.paragraphs[0]
        run = para.add_run("表格内容")
        _set_run_size(run, 24)  # 12pt = 小四, not 五号(10.5pt=21 half-pt)

        settings = FormatSettings(table=TableSettings(body_size="五号"))
        issues: list[ValidationIssue] = []
        _check_table_font_size(table, settings, issues)
        assert len(issues) > 0

    def test_passes_correct_table_font_size(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        para = cell.paragraphs[0]
        run = para.add_run("表格内容")
        _set_run_size(run, 21)  # 10.5pt = 五号, 21 half-pt

        settings = FormatSettings(table=TableSettings(body_size="五号"))
        issues: list[ValidationIssue] = []
        _check_table_font_size(table, settings, issues)
        assert len(issues) == 0


class TestCheckTableRowHeight:
    """Check that fixed table row heights are flagged (should be auto-fit)."""

    def test_flags_exact_row_height(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        tr = table.rows[0]._tr
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        trPr = etree.SubElement(tr, f"{{{W}}}trPr")
        trHeight = etree.SubElement(trPr, f"{{{W}}}trHeight")
        trHeight.set(f"{{{W}}}val", "1000")
        trHeight.set(f"{{{W}}}hRule", "exact")

        settings = FormatSettings()
        issues: list[ValidationIssue] = []
        _check_table_row_height(table, settings, issues)
        assert len(issues) > 0

    def test_passes_without_fixed_height(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        # No trHeight element at all — auto-fit, should pass
        settings = FormatSettings()
        issues: list[ValidationIssue] = []
        _check_table_row_height(table, settings, issues)
        assert len(issues) == 0

    def test_passes_at_least_reasonable(self):
        """atLeast with reasonable value should not flag."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        tr = table.rows[0]._tr
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        trPr = etree.SubElement(tr, f"{{{W}}}trPr")
        trHeight = etree.SubElement(trPr, f"{{{W}}}trHeight")
        trHeight.set(f"{{{W}}}val", "340")
        trHeight.set(f"{{{W}}}hRule", "atLeast")

        settings = FormatSettings()
        issues: list[ValidationIssue] = []
        _check_table_row_height(table, settings, issues)
        assert len(issues) == 0


class TestCheckCoverTitleSize:
    """Section 6 Mode B: 封面非二号字 → 警告"""

    def test_flags_wrong_cover_size(self):
        doc = Document()
        para = doc.add_paragraph("测试封面标题")
        para.alignment = 1  # centered
        _set_run_size(para.runs[0], 24)  # 12pt = 小四, should be 二号(22pt=44 half-pt)

        settings = FormatSettings()
        issues: list[ValidationIssue] = []
        _check_cover_title_size(doc.paragraphs, settings, issues)
        assert len(issues) > 0

    def test_passes_correct_cover_size(self):
        doc = Document()
        para = doc.add_paragraph("测试封面标题")
        para.alignment = 1
        _set_run_size(para.runs[0], 44)  # 22pt = 二号

        settings = FormatSettings()
        issues: list[ValidationIssue] = []
        _check_cover_title_size(doc.paragraphs, settings, issues)
        assert len(issues) == 0

    def test_skips_non_centered_text(self):
        doc = Document()
        para = doc.add_paragraph("左对齐文字")
        para.alignment = 0

        settings = FormatSettings()
        issues: list[ValidationIssue] = []
        _check_cover_title_size(doc.paragraphs, settings, issues)
        assert len(issues) == 0


class TestCheckAppendixBodySize:
    """Section 6 Mode B: 附件正文非小五 → 警告"""

    def test_flags_wrong_appendix_body_size(self):
        doc = Document()
        doc.add_paragraph("附件1：测试")
        para = doc.add_paragraph("附录正文")
        _set_run_size(para.runs[0], 24)  # 12pt = 小四

        settings = FormatSettings()
        issues: list[ValidationIssue] = []
        _check_appendix_body_size(doc.paragraphs, settings, issues)
        assert len(issues) > 0

    def test_passes_correct_appendix_body_size(self):
        doc = Document()
        doc.add_paragraph("附件1：测试")
        para = doc.add_paragraph("附录正文")
        _set_run_size(para.runs[0], 18)  # 9pt = 小五

        settings = FormatSettings()
        issues: list[ValidationIssue] = []
        _check_appendix_body_size(doc.paragraphs, settings, issues)
        assert len(issues) == 0

    def test_skips_body_before_appendix(self):
        """Paragraphs before appendix title should not be checked."""
        doc = Document()
        para = doc.add_paragraph("普通正文")
        _set_run_size(para.runs[0], 30)  # wrong size but not appendix

        settings = FormatSettings()
        issues: list[ValidationIssue] = []
        _check_appendix_body_size(doc.paragraphs, settings, issues)
        assert len(issues) == 0


class TestCheckGridAlignment:
    """Section 6 Mode B: 一级标题未勾选'对齐到网格'"""

    def test_flags_missing_snap_to_grid(self):
        doc = Document()
        para = doc.add_paragraph("1 一级标题")
        # Set outline level to 0 (Level 1)
        pPr = para._element.find(f"{{{W}}}pPr")
        if pPr is None:
            pPr = etree.SubElement(para._element, f"{{{W}}}pPr")
        outline = etree.SubElement(pPr, f"{{{W}}}outlineLvl")
        outline.set(f"{{{W}}}val", "0")
        # No snapToGrid or autoSpaceDE set → should be flagged

        issues: list[ValidationIssue] = []
        _check_grid_alignment([para], issues)
        assert len(issues) > 0
        assert "对齐" in issues[0].description

    def test_passes_with_grid_alignment(self):
        doc = Document()
        para = doc.add_paragraph("1 一级标题")
        pPr = para._element.find(f"{{{W}}}pPr")
        if pPr is None:
            pPr = etree.SubElement(para._element, f"{{{W}}}pPr")
        outline = etree.SubElement(pPr, f"{{{W}}}outlineLvl")
        outline.set(f"{{{W}}}val", "0")
        snap = etree.SubElement(pPr, f"{{{W}}}snapToGrid")
        snap.set(f"{{{W}}}val", "true")
        auto = etree.SubElement(pPr, f"{{{W}}}autoSpaceDE")
        auto.set(f"{{{W}}}val", "true")

        issues: list[ValidationIssue] = []
        _check_grid_alignment([para], issues)
        assert len(issues) == 0

    def test_skips_non_level_1(self):
        doc = Document()
        para = doc.add_paragraph("1.1 二级标题")
        pPr = para._element.find(f"{{{W}}}pPr")
        if pPr is None:
            pPr = etree.SubElement(para._element, f"{{{W}}}pPr")
        outline = etree.SubElement(pPr, f"{{{W}}}outlineLvl")
        outline.set(f"{{{W}}}val", "1")  # Level 2

        issues: list[ValidationIssue] = []
        _check_grid_alignment([para], issues)
        assert len(issues) == 0


class TestCheckAppendixColon:
    """Section 6 Mode B: 附件标题未使用冒号"""

    def test_flags_missing_colon(self):
        doc = Document()
        para = doc.add_paragraph("附件1 测试")

        issues: list[ValidationIssue] = []
        _check_appendix_colon([para], issues)
        assert len(issues) > 0
        assert "冒号" in issues[0].description

    def test_passes_with_colon(self):
        doc = Document()
        para = doc.add_paragraph("附件1：测试")

        issues: list[ValidationIssue] = []
        _check_appendix_colon([para], issues)
        assert len(issues) == 0

    def test_skips_non_appendix(self):
        doc = Document()
        para = doc.add_paragraph("附录内容")

        issues: list[ValidationIssue] = []
        _check_appendix_colon([para], issues)
        assert len(issues) == 0


class TestCheckFigureTableSpace:
    """Section 6 Mode B: 图题/Rev后缺少空格"""

    def test_flags_multiple_spaces_in_caption(self):
        doc = Document()
        para = doc.add_paragraph("表 1.1-1  多余的  空格")

        issues: list[ValidationIssue] = []
        _check_figure_table_space([para], issues)
        assert len(issues) > 0

    def test_passes_single_space_in_caption(self):
        doc = Document()
        para = doc.add_paragraph("表 1.1-1 正常空格")

        issues: list[ValidationIssue] = []
        _check_figure_table_space([para], issues)
        assert len(issues) == 0

    def test_skips_non_caption(self):
        doc = Document()
        para = doc.add_paragraph("普通段落")

        issues: list[ValidationIssue] = []
        _check_figure_table_space([para], issues)
        assert len(issues) == 0


class TestCheckChairmanNumber:
    """Section 6: 检查负责人/审定人信息."""

    def test_flags_empty_chairman(self):
        doc = Document()
        para = doc.add_paragraph("负责人：")
        issues: list[ValidationIssue] = []
        _check_chairman_number([para], [], issues)
        assert len(issues) > 0

    def test_passes_valid_chairman(self):
        doc = Document()
        para = doc.add_paragraph("负责人：张三")
        issues: list[ValidationIssue] = []
        _check_chairman_number([para], [], issues)
        assert len(issues) == 0

    def test_passes_shenheren(self):
        doc = Document()
        para = doc.add_paragraph("审定人：李四")
        issues: list[ValidationIssue] = []
        _check_chairman_number([para], [], issues)
        assert len(issues) == 0

    def test_flags_no_chairman_at_all(self):
        doc = Document()
        para = doc.add_paragraph("普通段落内容")
        issues: list[ValidationIssue] = []
        _check_chairman_number([para], [], issues)
        assert len(issues) > 0  # Should warn about missing chairman

    def test_skips_empty_document(self):
        issues: list[ValidationIssue] = []
        _check_chairman_number([], [], issues)
        assert len(issues) == 0


class TestValidateDocument:
    """Integration tests for the validate_document orchestrator."""

    def test_validate_with_no_rules_returns_empty(self, tmp_path):
        """When no validation rules are enabled, should return empty list."""
        doc = Document()
        doc.add_paragraph("任何文字格式")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        settings = FormatSettings(
            validation=ValidationRules()  # all checks disabled
        )
        issues = validate_document(path, settings)
        assert issues == []

    def test_validate_with_chinese_font_check(self, tmp_path):
        """Chinese font check should detect non-宋体."""
        doc = Document()
        para = doc.add_paragraph("一段中文")
        _set_run_fonts(para.runs[0], east_asia="黑体")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        settings = FormatSettings(
            validation=ValidationRules(check_chinese_font=True)
        )
        issues = validate_document(path, settings)
        assert len(issues) > 0

    def test_validate_with_forbidden_words(self, tmp_path):
        """Forbidden words check should detect 附图/附表."""
        doc = Document()
        doc.add_paragraph("详见附图1")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        settings = FormatSettings(
            validation=ValidationRules(
                check_forbidden_words=True,
                forbidden_words=("附图", "附表"),
            )
        )
        issues = validate_document(path, settings)
        assert len(issues) > 0

    def test_validate_passes_clean_document(self, tmp_path):
        """A document with correct formatting should pass validation."""
        doc = Document()
        para = doc.add_paragraph("一段中文")
        _set_run_fonts(para.runs[0], east_asia="宋体", ascii="Times New Roman")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        settings = FormatSettings(
            validation=ValidationRules(
                check_chinese_font=True,
                check_ascii_font=True,
            )
        )
        issues = validate_document(path, settings)
        assert issues == []


class TestCheckParagraphSpacing:
    """Check that paragraphs with non-zero before/after spacing are flagged."""

    def test_flags_nonzero_before(self):
        doc = Document()
        para = doc.add_paragraph("正文")
        _set_spacing(para, before=480)
        issues: list[ValidationIssue] = []
        _check_paragraph_spacing(para, issues)
        assert len(issues) > 0

    def test_flags_nonzero_after(self):
        doc = Document()
        para = doc.add_paragraph("正文")
        _set_spacing(para, after=480)
        issues: list[ValidationIssue] = []
        _check_paragraph_spacing(para, issues)
        assert len(issues) > 0

    def test_flags_nonzero_before_lines(self):
        doc = Document()
        para = doc.add_paragraph("正文")
        _set_spacing(para, before_lines=100)
        issues: list[ValidationIssue] = []
        _check_paragraph_spacing(para, issues)
        assert len(issues) > 0

    def test_passes_zero_spacing(self):
        doc = Document()
        para = doc.add_paragraph("正文")
        _set_spacing(para, before=0, after=0, before_lines=0, after_lines=0)
        issues: list[ValidationIssue] = []
        _check_paragraph_spacing(para, issues)
        assert len(issues) == 0

    def test_skips_empty_paragraph(self):
        doc = Document()
        para = doc.add_paragraph("   ")
        _set_spacing(para, before=480)
        issues: list[ValidationIssue] = []
        _check_paragraph_spacing(para, issues)
        assert len(issues) == 0


class TestCheckCaptionAlignment:
    """Check that table/figure captions are centered."""

    def test_flags_non_centered_table_caption(self):
        doc = Document()
        para = doc.add_paragraph("表 1-1 示例")
        _set_jc(para, "left")
        issues: list[ValidationIssue] = []
        _check_caption_alignment(para, issues)
        assert len(issues) > 0

    def test_flags_non_centered_figure_caption(self):
        doc = Document()
        para = doc.add_paragraph("图 1-1 示例")
        _set_jc(para, "both")
        issues: list[ValidationIssue] = []
        _check_caption_alignment(para, issues)
        assert len(issues) > 0

    def test_passes_centered_caption(self):
        doc = Document()
        para = doc.add_paragraph("表 1-1 示例")
        _set_jc(para, "center")
        issues: list[ValidationIssue] = []
        _check_caption_alignment(para, issues)
        assert len(issues) == 0

    def test_skips_non_caption_text(self):
        doc = Document()
        para = doc.add_paragraph("普通正文")
        issues: list[ValidationIssue] = []
        _check_caption_alignment(para, issues)
        assert len(issues) == 0


class TestCheckTableFixedDimensions:
    """Check that tables don't have fixed dimensions."""

    def test_flags_fixed_table_width(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        # Set fixed table width via w:tblW
        tblPr = table._element.find(f"{{{W}}}tblPr")
        if tblPr is None:
            tblPr = etree.SubElement(table._element, f"{{{W}}}tblPr")
        tblW = etree.SubElement(tblPr, f"{{{W}}}tblW")
        tblW.set(f"{{{W}}}type", "dxa")
        tblW.set(f"{{{W}}}w", "5000")
        issues: list[ValidationIssue] = []
        _check_table_fixed_dimensions(table, issues)
        assert len(issues) > 0

    def test_flags_fixed_layout(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        tblPr = table._element.find(f"{{{W}}}tblPr")
        if tblPr is None:
            tblPr = etree.SubElement(table._element, f"{{{W}}}tblPr")
        tblLayout = etree.SubElement(tblPr, f"{{{W}}}tblLayout")
        tblLayout.set(f"{{{W}}}type", "fixed")
        issues: list[ValidationIssue] = []
        _check_table_fixed_dimensions(table, issues)
        assert len(issues) > 0

    def test_passes_auto_fit_table(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        issues: list[ValidationIssue] = []
        _check_table_fixed_dimensions(table, issues)
        assert len(issues) == 0


def _set_spacing(para, before=0, after=0, before_lines=0, after_lines=0):
    """Set spacing attributes on a paragraph at OOXML level."""
    pPr = para._element.find(f"{{{W}}}pPr")
    if pPr is None:
        pPr = etree.SubElement(para._element, f"{{{W}}}pPr")
    spacing = pPr.find(f"{{{W}}}spacing")
    if spacing is None:
        spacing = etree.SubElement(pPr, f"{{{W}}}spacing")
    spacing.set(f"{{{W}}}before", str(before))
    spacing.set(f"{{{W}}}after", str(after))
    spacing.set(f"{{{W}}}beforeLines", str(before_lines))
    spacing.set(f"{{{W}}}afterLines", str(after_lines))


def _set_jc(para, val: str):
    """Set w:jc alignment on a paragraph at OOXML level."""
    pPr = para._element.find(f"{{{W}}}pPr")
    if pPr is None:
        pPr = etree.SubElement(para._element, f"{{{W}}}pPr")
    jc = pPr.find(f"{{{W}}}jc")
    if jc is None:
        jc = etree.SubElement(pPr, f"{{{W}}}jc")
    jc.set(f"{{{W}}}val", val)
