"""TDD tests for appendix formatting (tvba_core_appendix.py).

Requirements ref: 多模板自动格式化与检查插件需求说明书 section 4.2 (Mode B)
附件标题: 宋体/TNR, 小四, 加粗, 1.5倍, 序号后强制冒号 (:)
附件正文: 宋体/TNR, 小五, 常规, 单倍
"""
import pytest
from docx import Document

from tvba_core_appendix import is_appendix_title, format_appendix
from tvba_core_oox import set_outline_level
from tvba_settings import AppendixSettings

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _get_run_font_size_pt(para) -> float | None:
    sz = para._element.find(f".//{{{W}}}r/{{{W}}}rPr/{{{W}}}sz")
    if sz is None:
        return None
    return int(sz.get(f"{{{W}}}val")) / 2


def _get_run_bold(para) -> bool:
    return para._element.find(f".//{{{W}}}r/{{{W}}}rPr/{{{W}}}b") is not None


def _get_line_spacing_val(para) -> int | None:
    spacing = para._element.find(f".//{{{W}}}pPr/{{{W}}}spacing")
    if spacing is None:
        return None
    val = spacing.get(f"{{{W}}}line")
    return int(val) if val else None


class TestIsAppendixTitle:
    """RED phase: appendix title detection — 匹配 "附件..." 模式."""

    def test_attachment_with_number(self):
        assert is_appendix_title("附件1") is True

    def test_attachment_with_number_and_colon(self):
        assert is_appendix_title("附件1：") is True

    def test_attachment_with_number_and_halfwidth_colon(self):
        assert is_appendix_title("附件1:") is True

    def test_attachment_without_number(self):
        assert is_appendix_title("附件") is True

    def test_attachment_with_spaces_after_number(self):
        assert is_appendix_title("附件 1：") is True

    def test_not_attachment(self):
        assert is_appendix_title("附录一 项目概况") is False

    def test_not_attachment_body(self):
        assert is_appendix_title("这是附件中的一段文字") is False

    def test_empty_string(self):
        assert is_appendix_title("") is False

    def test_attachment_with_longer_number(self):
        assert is_appendix_title("附件12：") is True


class TestFormatAppendixTitle:
    """RED phase: appendix title formatting — 小四/加粗/1.5倍."""

    def test_formats_title_size_as_xiaosi(self):
        """Appendix title should be 小四 (12pt)."""
        doc = Document()
        doc.add_paragraph("附件1：相关文件列表")

        format_appendix(doc)

        para = doc.paragraphs[0]
        size_pt = _get_run_font_size_pt(para)
        assert size_pt == 12.0, f"Expected 小四 (12pt), got {size_pt}pt"

    def test_formats_title_bold(self):
        doc = Document()
        doc.add_paragraph("附件1：相关文件列表")

        format_appendix(doc)

        para = doc.paragraphs[0]
        assert _get_run_bold(para) is True

    def test_formats_title_1_5_line_spacing(self):
        doc = Document()
        doc.add_paragraph("附件1：相关文件列表")

        format_appendix(doc)

        line = _get_line_spacing_val(doc.paragraphs[0])
        assert line is not None
        assert line == 360, f"Expected 360 twips (1.5倍), got {line}"

    def test_formats_title_left_aligned(self):
        """Appendix title should be left-aligned."""
        doc = Document()
        para = doc.add_paragraph("附件1：相关文件列表")
        # Set to center first to verify it gets changed
        para.alignment = 1

        format_appendix(doc)

        assert para.alignment == 0, f"Expected left (0), got {para.alignment}"

    def test_adds_colon_after_number_when_missing(self):
        """When appendix title has no colon after number, one must be added."""
        doc = Document()
        doc.add_paragraph("附件1 相关文件")

        format_appendix(doc)

        para = doc.paragraphs[0]
        assert "：" in para.text, f"Expected fullwidth colon in '{para.text}'"

    def test_preserves_existing_colon(self):
        doc = Document()
        doc.add_paragraph("附件1：已有冒号")

        format_appendix(doc)

        para = doc.paragraphs[0]
        # Should still have exactly one colon
        assert para.text.count("：") >= 1


class TestFormatAppendixBody:
    """RED phase: appendix body — 小五/常规/单倍."""

    def test_formats_body_size_as_xiaowu(self):
        """Body paragraphs after appendix title should be 小五 (9pt)."""
        doc = Document()
        doc.add_paragraph("附件1：测试")
        doc.add_paragraph("这是附录的正文内容")

        format_appendix(doc)

        body_para = doc.paragraphs[1]
        size_pt = _get_run_font_size_pt(body_para)
        assert size_pt == 9.0, f"Expected 小五 (9pt), got {size_pt}pt"

    def test_formats_body_not_bold(self):
        doc = Document()
        doc.add_paragraph("附件1：测试")
        doc.add_paragraph("附录正文不应该加粗")

        format_appendix(doc)

        body_para = doc.paragraphs[1]
        assert _get_run_bold(body_para) is False

    def test_formats_body_single_line_spacing(self):
        doc = Document()
        doc.add_paragraph("附件1：测试")
        doc.add_paragraph("附录正文单倍行距")

        format_appendix(doc)

        line = _get_line_spacing_val(doc.paragraphs[1])
        assert line is not None
        assert line == 240, f"Expected 240 twips (1.0倍), got {line}"

    def test_skips_body_before_appendix(self):
        """Paragraphs before the first appendix title should NOT be formatted as appendix body."""
        doc = Document()
        doc.add_paragraph("普通正文段落")
        doc.add_paragraph("附件1：测试")
        doc.add_paragraph("真正的附录内容")

        format_appendix(doc)

        # Paragraph before appendix title should remain unchanged
        before_para = doc.paragraphs[0]
        size = _get_run_font_size_pt(before_para)
        assert size != 9.0, f"Body before appendix should not get 小五 format"


class TestFormatAppendixBoundary:
    """RED phase: appendix zone boundary — 遇到下一个标题时停止."""

    def test_stops_at_next_heading(self):
        """Appendix body zone ends when a heading paragraph is encountered."""
        doc = Document()
        doc.add_paragraph("附件1：测试")
        doc.add_paragraph("这是附录正文")
        heading = doc.add_paragraph("5 其他章节")
        set_outline_level(heading, 0)  # Level 1 heading
        doc.add_paragraph("这是普通正文")

        format_appendix(doc)

        # The heading should NOT be formatted as appendix body
        size = _get_run_font_size_pt(heading)
        assert size != 9.0, f"Heading should not get appendix body format"

        # The paragraph after heading should also NOT be appendix body
        after = doc.paragraphs[3]
        size2 = _get_run_font_size_pt(after)
        assert size2 != 9.0, f"Text after heading should not get appendix body format"

    def test_multiple_appendix_blocks(self):
        """Multiple appendix titles each start new appendix zones."""
        doc = Document()
        doc.add_paragraph("附件1：第一个附录")
        doc.add_paragraph("第一个附录的内容")
        doc.add_paragraph("附件2：第二个附录")
        doc.add_paragraph("第二个附录的内容")

        format_appendix(doc)

        # Both appendix titles should be formatted as title
        title1 = doc.paragraphs[0]
        assert _get_run_bold(title1) is True
        assert _get_run_font_size_pt(title1) == 12.0

        title2 = doc.paragraphs[2]
        assert _get_run_bold(title2) is True
        assert _get_run_font_size_pt(title2) == 12.0

        # Both appendix bodies should be formatted as body
        body1 = doc.paragraphs[1]
        assert _get_run_font_size_pt(body1) == 9.0

        body2 = doc.paragraphs[3]
        assert _get_run_font_size_pt(body2) == 9.0


class TestAppendixEdgeCases:
    """Edge cases for appendix processing."""

    def test_empty_document(self):
        doc = Document()
        format_appendix(doc)  # should not crash

    def test_no_appendix_in_document(self):
        """Document without any appendix should be unchanged."""
        doc = Document()
        doc.add_paragraph("普通段落")
        doc.add_paragraph("另一个普通段落")

        format_appendix(doc)

        # Neither paragraph should be affected
        for para in doc.paragraphs:
            size = _get_run_font_size_pt(para)
            assert size != 9.0 and size != 12.0

    def test_empty_paragraphs_handled(self):
        """Empty paragraphs within appendix zone should be skipped gracefully."""
        doc = Document()
        doc.add_paragraph("附件1：测试")
        doc.add_paragraph("")
        doc.add_paragraph("附录正文")

        format_appendix(doc)

        # Empty para: no font applied, should not crash
        body_para = doc.paragraphs[2]
        size = _get_run_font_size_pt(body_para)
        assert size == 9.0  # body after empty paragraph still formatted


class TestFormatAppendixWithSettings:
    """Settings-aware appendix formatting — template-driven values."""

    def test_title_uses_custom_font(self):
        doc = Document()
        doc.add_paragraph("附件1：相关文件")

        settings = AppendixSettings(title_font="黑体")
        format_appendix(doc, settings)

        para = doc.paragraphs[0]
        east = para._element.find(f".//{{{W}}}r/{{{W}}}rPr/{{{W}}}rFonts")
        assert east is not None
        assert east.get(f"{{{W}}}eastAsia") == "黑体"

    def test_title_uses_custom_size(self):
        doc = Document()
        doc.add_paragraph("附件1：相关文件")

        settings = AppendixSettings(title_size="三号")  # 16pt
        format_appendix(doc, settings)

        size_pt = _get_run_font_size_pt(doc.paragraphs[0])
        assert size_pt == 16.0

    def test_title_uses_custom_bold(self):
        doc = Document()
        doc.add_paragraph("附件1：相关文件")

        settings = AppendixSettings(title_bold=False)
        format_appendix(doc, settings)

        assert _get_run_bold(doc.paragraphs[0]) is False

    def test_title_uses_custom_line_spacing(self):
        doc = Document()
        doc.add_paragraph("附件1：相关文件")

        settings = AppendixSettings(title_line_spacing=2.0)
        format_appendix(doc, settings)

        line = _get_line_spacing_val(doc.paragraphs[0])
        assert line == 480, f"Expected 480 twips (2.0倍), got {line}"

    def test_body_uses_custom_font(self):
        doc = Document()
        doc.add_paragraph("附件1：相关文件")
        doc.add_paragraph("附录正文内容")

        settings = AppendixSettings(body_font="黑体")
        format_appendix(doc, settings)

        body_para = doc.paragraphs[1]
        east = body_para._element.find(f".//{{{W}}}r/{{{W}}}rPr/{{{W}}}rFonts")
        assert east is not None
        assert east.get(f"{{{W}}}eastAsia") == "黑体"

    def test_body_uses_custom_size(self):
        doc = Document()
        doc.add_paragraph("附件1：相关文件")
        doc.add_paragraph("附录正文内容")

        settings = AppendixSettings(body_size="小四")  # 12pt
        format_appendix(doc, settings)

        size_pt = _get_run_font_size_pt(doc.paragraphs[1])
        assert size_pt == 12.0

    def test_body_uses_custom_bold(self):
        doc = Document()
        doc.add_paragraph("附件1：相关文件")
        doc.add_paragraph("附录正文内容")

        settings = AppendixSettings(body_bold=True)
        format_appendix(doc, settings)

        assert _get_run_bold(doc.paragraphs[1]) is True

    def test_body_uses_custom_line_spacing(self):
        doc = Document()
        doc.add_paragraph("附件1：相关文件")
        doc.add_paragraph("附录正文内容")

        settings = AppendixSettings(body_line_spacing=1.5)
        format_appendix(doc, settings)

        line = _get_line_spacing_val(doc.paragraphs[1])
        assert line == 360, f"Expected 360 twips (1.5倍), got {line}"

    def test_none_settings_uses_defaults(self):
        """Passing None should use AppendixSettings() defaults."""
        doc = Document()
        doc.add_paragraph("附件1：相关文件")
        doc.add_paragraph("附录正文内容")

        format_appendix(doc, None)

        title = doc.paragraphs[0]
        assert _get_run_font_size_pt(title) == 12.0  # 小四
        assert _get_run_bold(title) is True

        body = doc.paragraphs[1]
        assert _get_run_font_size_pt(body) == 9.0  # 小五
        assert _get_run_bold(body) is False
