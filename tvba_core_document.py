"""Document orchestrator.

Corresponds to VBA FormatModule.bas:
  - ApplySettingsToDocument
"""
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from tvba_settings import FormatSettings
from tvba_core_body import apply_normal_style, apply_paragraph
from tvba_core_title import auto_detect_and_format, split_compound_paragraphs, _has_numbering_hint
from tvba_core_toc import is_toc_paragraph, refresh_toc
from tvba_core_table import refresh_all as refresh_tables
from tvba_core_figure import refresh_all as refresh_figures
from tvba_core_convert import ensure_docx
from tvba_core_normalize import unify_ascii_font
from tvba_core_numbering import auto_select
from tvba_core_oox import apply_paragraph_spacing, get_effective_outline_level, sync_numbering_with_titles
from tvba_core_cover import format_cover_title
from tvba_core_appendix import format_appendix
from tvba_logging import log_event, log_exception


MAX_COM_NUMBERING_CANDIDATES = 300


@dataclass
class ApplyWarnings:
    """Warnings collected during document formatting."""
    messages: list[str] = field(default_factory=list)


def apply_settings_to_document(
    docx_path: Path,
    settings: FormatSettings,
    *,
    list_resolver=None,
    output_path: Path | None = None,
    progress_cb=None,
) -> tuple[Path, ApplyWarnings]:
    """Apply all formatting settings to a document.

    Returns (output_path, warnings).
    """
    warnings = ApplyWarnings()
    log_event("apply.start", path=str(docx_path), output_path=str(output_path) if output_path else None)

    if progress_cb:
        progress_cb("Loading document...", 0.0)

    docx_path = ensure_docx(docx_path)
    log_event("apply.ensure_docx.done", path=str(docx_path))

    doc = Document(str(docx_path))
    log_event("apply.document.loaded", paragraphs=len(doc.paragraphs), tables=len(doc.tables))

    # Cache paragraph/table lists to avoid repeated rebuilding.
    # python-docx rebuilds these from XML on every access.
    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)

    # Cache outline level lookups — style chain resolution is expensive
    # and the same paragraph may be checked multiple times.
    outline_cache: dict[int, int | None] = {}

    # Pre-build style-by-id map because python-docx deprecated style_id lookup.
    style_by_id = {style.style_id: style for style in doc.styles}

    # Pre-build TOC style id set: style ids whose names contain "toc".
    # This lets is_toc_paragraph detect custom TOC styles without the expensive
    # paragraph.style lookup.
    toc_style_ids = {style.style_id for style in doc.styles if "toc" in style.name.lower()}

    if progress_cb:
        progress_cb("Applying normal style...", 0.1)
    apply_normal_style(doc, settings.body)

    # Pre-process: split compound paragraphs with multiple concatenated titles
    if progress_cb:
        progress_cb("Splitting compound titles...", 0.15)
    paragraphs = split_compound_paragraphs(doc, _paragraphs=paragraphs)
    log_event("apply.split_compound.done", paragraphs=len(paragraphs))

    numbering_candidate_count = 0
    if settings.auto_detect_numeric_titles and settings.auto_detect_include_list_paragraphs:
        for para in paragraphs:
            text = (para.text or "").strip()
            if text and _has_numbering_hint(para, style_by_id):
                numbering_candidate_count += 1
    prefer_com_for_this_doc = settings.prefer_com_resolver
    if settings.prefer_com_resolver and numbering_candidate_count > MAX_COM_NUMBERING_CANDIDATES:
        prefer_com_for_this_doc = False
        msg = (
            f"检测到 {numbering_candidate_count} 个自动编号候选段落，超过 "
            f"{MAX_COM_NUMBERING_CANDIDATES} 个。为避免 Word COM 长时间卡住，"
            "本次跳过自动编号标题解析，仅使用文本标题识别。"
        )
        warnings.messages.append(msg)
        log_event(
            "apply.com_skipped.too_many_numbering_candidates",
            candidates=numbering_candidate_count,
            limit=MAX_COM_NUMBERING_CANDIDATES,
        )

    # Phase 2 fix: If COM resolver will be used, save preprocessed document
    # to a temp file so the COM paragraph indices align with the in-memory
    # python-docx paragraphs after split. Without this, COM would open the
    # original (pre-split) file and paragraph mapping would be wrong.
    _preprocessed_path = None
    try:
        if settings.auto_detect_include_list_paragraphs and prefer_com_for_this_doc:
            import tempfile
            if progress_cb:
                progress_cb("Saving preprocessed document for Word COM...", 0.18)
            log_event("apply.preprocessed_save.start")
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix="transvba_pre_")
            _preprocessed_path = Path(tmp.name)
            tmp.close()
            doc.save(str(_preprocessed_path))
            log_event(
                "apply.preprocessed_save.done",
                path=str(_preprocessed_path),
                size=_preprocessed_path.stat().st_size if _preprocessed_path.exists() else None,
            )

        # Auto-detect titles
        if progress_cb:
            progress_cb("Detecting titles...", 0.2)
        if settings.auto_detect_numeric_titles:
            if list_resolver is None and settings.auto_detect_include_list_paragraphs:
                resolver_path = str(_preprocessed_path) if _preprocessed_path else str(docx_path)
                if progress_cb:
                    progress_cb(
                        "Opening Word COM list resolver..." if prefer_com_for_this_doc else "Using fast docx list resolver...",
                        0.22,
                    )
                log_event("apply.auto_select.start", prefer_com=prefer_com_for_this_doc, resolver_path=resolver_path)
                list_resolver, resolver_status = auto_select(prefer_com=prefer_com_for_this_doc, docx_path=resolver_path, doc=doc)
                log_event("apply.auto_select.done", mode=resolver_status.mode, warnings=len(resolver_status.warnings))
                if resolver_status.warnings:
                    warnings.messages.extend(resolver_status.warnings)
                # If COM was preferred but fell back, add explicit warning
                if prefer_com_for_this_doc and resolver_status.mode == "docx_fallback":
                    warnings.messages.append(
                        "Word COM 不可用，自动编号标题解析降级为纯 python-docx 模式，"
                        "Word 多级列表标题可能无法被正确识别。"
                    )
            if progress_cb:
                progress_cb("Scanning title candidates...", 0.26)
            log_event("apply.auto_detect.start", paragraphs=len(paragraphs))
            auto_detect_and_format(doc, settings, list_resolver, _paragraphs=paragraphs, _outline_cache=outline_cache, _style_by_id=style_by_id, _toc_style_ids=toc_style_ids)
            log_event("apply.auto_detect.done")

        # Clear outline cache so the body loop sees levels set by auto_detect_and_format.
        # auto_detect_and_format caches outline_level=None before calling apply_title_style
        # which sets the actual outline level; stale cache entries would cause the body
        # loop to call apply_paragraph (body) instead of apply_title_style.
        outline_cache.clear()

        if progress_cb:
            progress_cb("Formatting paragraphs...", 0.4)
        from tvba_core_table import is_table_caption_paragraph
        from tvba_core_figure import is_figure_caption_paragraph

        for para in paragraphs:
            if is_toc_paragraph(para, _toc_style_ids=toc_style_ids):
                continue

            # Skip empty paragraphs — keep their original spacing/layout
            if not para.text or not para.text.strip():
                continue

            # Skip captions — they are formatted later by refresh_tables / refresh_figures
            if is_table_caption_paragraph(para, doc) or is_figure_caption_paragraph(para, doc):
                continue

            # Check effective outline level (direct or from style)
            outline_level = get_effective_outline_level(para, _cache=outline_cache, _style_by_id=style_by_id)

            if outline_level is not None and 0 <= outline_level <= 4:
                # Title paragraph (Word levels 1-5) — apply title formatting
                from tvba_core_title import apply_title_style
                level = outline_level + 1  # Convert 0-4 to 1-5
                apply_title_style(para, level, settings.titles[level - 1], settings.body)
            else:
                # Body text (no outline or level > 4) — apply body formatting
                apply_paragraph(para, settings.body)

        # Sync numbering definitions so auto-generated list numbers match titles
        sync_numbering_with_titles(doc, settings, _paragraphs=paragraphs)

        # Forbidden words replacement (applied during formatting, checked during validation)
        if settings.validation.forbidden_replacements and settings.body.modify_content:
            if progress_cb:
                progress_cb("Replacing forbidden words...", 0.56)
            for para in paragraphs:
                text = para.text
                if not text or not text.strip():
                    continue
                replaced = text
                for forbidden, replacement in settings.validation.forbidden_replacements.items():
                    replaced = replaced.replace(forbidden, replacement)
                if replaced != text:
                    # Replace in the longest run (most likely the content run)
                    runs = list(para.runs)
                    if runs:
                        longest = max(runs, key=lambda r: len(r.text))
                        for forbidden, replacement in settings.validation.forbidden_replacements.items():
                            longest.text = longest.text.replace(forbidden, replacement)

        if progress_cb:
            progress_cb("Formatting TOC...", 0.6)
        refresh_toc(doc, settings.toc, _paragraphs=paragraphs)

        if progress_cb:
            progress_cb("Formatting tables...", 0.7)
        table_warnings = refresh_tables(doc, settings.table, _paragraphs=paragraphs, _tables=tables)
        if table_warnings:
            warnings.messages.extend(table_warnings)

        if progress_cb:
            progress_cb("Formatting figures...", 0.8)
        refresh_figures(doc, settings.figure, _paragraphs=paragraphs)

        if progress_cb:
            progress_cb("Formatting cover page...", 0.82)
        format_cover_title(doc, settings.cover, _paragraphs=paragraphs)

        if progress_cb:
            progress_cb("Formatting appendix...", 0.84)
        format_appendix(doc, settings.appendix, _paragraphs=paragraphs, _outline_cache=outline_cache, _style_by_id=style_by_id)

        if progress_cb:
            progress_cb("Converting presidential orders...", 0.86)
        if settings.format_presidential_order:
            from tvba_core_presidential import format_presidential_order_numbers
            format_presidential_order_numbers(doc, _paragraphs=paragraphs)

        if progress_cb:
            progress_cb("Formatting headers...", 0.88)
        _format_headers(doc, settings.header)

        if progress_cb:
            progress_cb("Normalizing fonts...", 0.9)
        unify_ascii_font(doc, "Times New Roman", _paragraphs=paragraphs, _tables=tables)

        if progress_cb:
            progress_cb("Saving...", 0.95)
    finally:
        # Close COM resolver before saving to release file lock,
        # and clean up temp preprocessed file — must run even on exception
        # to avoid leaking Word processes and temp files.
        if list_resolver is not None and hasattr(list_resolver, "close"):
            try:
                list_resolver.close()
            except Exception:
                pass
        if _preprocessed_path is not None:
            try:
                if _preprocessed_path.exists():
                    _preprocessed_path.unlink()
                    log_event("apply.preprocessed_cleanup.done", path=str(_preprocessed_path))
            except OSError:
                log_exception("apply.preprocessed_cleanup.failed")
                pass

    out = output_path or docx_path
    log_event("apply.save.start", output_path=str(out))
    doc.save(str(out))
    log_event("apply.save.done", output_path=str(out))

    if progress_cb:
        progress_cb("Done", 1.0)

    log_event("apply.done", output_path=str(out), warnings=len(warnings.messages))
    return out, warnings


def _format_headers(doc, settings) -> None:
    """Format header text in all sections using settings values.

    Also normalize 'Rev.' spacing: ensure single space after 'Rev.'
    """
    from tvba_core_oox import format_all_runs_in_paragraph
    from tvba_utils import size_label_to_points

    for section in doc.sections:
        header = section.header
        if header is None:
            continue
        for para in header.paragraphs:
            format_all_runs_in_paragraph(
                para,
                ascii_font="Times New Roman",
                eastasia_font=settings.font,
                size_pt=size_label_to_points(settings.size),
                bold=settings.bold,
            )
            apply_paragraph_spacing(
                para.paragraph_format,
                line_spacing=settings.line_spacing,
            )

            # Normalize "Rev." spacing: "Rev.  " → "Rev. "
            import re
            for run in para.runs:
                if 'Rev.' in run.text:
                    run.text = re.sub(r'Rev\.\s+', 'Rev. ', run.text)
