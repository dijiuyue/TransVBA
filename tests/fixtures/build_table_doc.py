"""Build a test .docx with tables and captions."""
from docx import Document
import sys
from pathlib import Path

def build(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("表 1.1-1 示例表格")
    table = doc.add_table(rows=3, cols=3)
    for row in table.rows:
        for cell in row.cells:
            cell.text = "单元格"
    doc.add_paragraph("正文段落")
    doc.add_paragraph("表 2-1 另一个表格")
    table2 = doc.add_table(rows=2, cols=4)
    doc.save(path)

if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("table_test.docx")
    build(out)
    print(f"Built: {out}")
