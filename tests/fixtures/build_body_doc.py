"""Build a test .docx with various body text paragraphs."""
from docx import Document
from docx.shared import Pt
import sys
from pathlib import Path

def build(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("这是一段普通正文，用于测试正文格式刷新功能。")
    doc.add_paragraph("第二段正文，包含一些数字 123 和英文 ABC。")
    p = doc.add_paragraph()
    p.add_run("第三段有粗体").bold = True
    p.add_run("和普通文本混合。")
    doc.save(path)

if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("body_test.docx")
    build(out)
    print(f"Built: {out}")
