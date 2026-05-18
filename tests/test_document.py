import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.shared import Pt

from tvba_core_document import apply_settings_to_document
from tvba_settings import FormatSettings, BodySettings, FigureSettings, TableSettings


def _get_run_eastasia_font(run):
    """Extract w:rFonts/@w:eastAsia from a run's XML."""
    rPr = run._element.find(".//w:rPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
    if rPr is None:
        return None
    rFonts = rPr.find("w:rFonts", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
    if rFonts is None:
        return None
    return rFonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")


class TestApplySettingsToDocument:
    def test_processes_body_text(self):
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            doc.add_paragraph("正文段落")
            doc.save(path)

            settings = FormatSettings()
            out, warnings = apply_settings_to_document(path, settings)
            assert out.exists()
            assert isinstance(warnings.messages, list)

            # Verify output
            result = Document(out)
            assert len(result.paragraphs) == 1

    def test_processes_titles(self):
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            doc.add_paragraph("1 一级标题")
            doc.add_paragraph("1.1 二级标题")
            doc.add_paragraph("正文")
            doc.save(path)

            settings = FormatSettings()
            out, _ = apply_settings_to_document(path, settings)
            result = Document(out)

            # First paragraph should have outline level
            pPr = result.paragraphs[0]._element.find(
                ".//w:pPr",
                {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            )
            outline = pPr.find("w:outlineLvl", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            assert outline is not None

    def test_calls_progress_callback(self):
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            doc.add_paragraph("正文")
            doc.save(path)

            progress_calls = []
            def cb(msg, pct):
                progress_calls.append((msg, pct))

            settings = FormatSettings()
            apply_settings_to_document(path, settings, progress_cb=cb)
            assert len(progress_calls) > 0

    def test_large_numbered_document_skips_com_resolver(self):
        """Large docs with many numbered paragraphs should not enter slow Word COM."""
        from lxml import etree
        from tvba_core_numbering import DocxListResolver, ResolverStatus

        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "large.docx"
            doc = Document()
            W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            for i in range(305):
                para = doc.add_paragraph(f"Numbered paragraph {i}")
                pPr = para._element.get_or_add_pPr()
                numPr = etree.SubElement(pPr, f"{{{W}}}numPr")
                ilvl = etree.SubElement(numPr, f"{{{W}}}ilvl")
                ilvl.set(f"{{{W}}}val", "0")
                numId = etree.SubElement(numPr, f"{{{W}}}numId")
                numId.set(f"{{{W}}}val", "1")
            doc.save(path)

            calls = []

            def fake_auto_select(*, prefer_com, docx_path, doc):
                calls.append((prefer_com, docx_path))
                return DocxListResolver(doc), ResolverStatus(
                    mode="docx_fallback",
                    reliable_rendered_text=False,
                    warnings=[],
                )

            with patch("tvba_core_document.auto_select", side_effect=fake_auto_select):
                _, warnings = apply_settings_to_document(path, FormatSettings())

            assert calls
            assert calls[0][0] is False
            assert any("跳过自动编号标题解析" in msg for msg in warnings.messages)

    def test_custom_output_path(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "test.docx"
            out = Path(td) / "output.docx"
            doc = Document()
            doc.add_paragraph("正文")
            doc.save(src)

            settings = FormatSettings()
            result, _ = apply_settings_to_document(src, settings, output_path=out)
            assert result == out
            assert out.exists()

    def test_doc_file_converted_before_processing(self):
        with tempfile.TemporaryDirectory() as td:
            doc_path = Path(td) / "test.doc"
            doc_path.write_text("fake doc content")
            expected_docx = Path(td) / "test.docx"

            # Pre-create the expected .docx so Document() can open it after mocked conversion
            Document().save(expected_docx)

            mock_word = MagicMock()
            mock_doc = MagicMock()
            mock_word.Documents.Open.return_value = mock_doc

            settings = FormatSettings()
            with patch("win32com.client.DispatchEx", return_value=mock_word):
                out, _ = apply_settings_to_document(doc_path, settings)

            assert out.suffix == ".docx"
            assert out.exists()

    def test_body_font_actually_changes(self):
        """Verify apply_settings_to_document actually changes paragraph font.

        This test creates a document with paragraphs using 'Arial' font,
        then applies settings with body.font='黑体', and verifies the
        output paragraphs have '黑体' as their eastAsia font.
        """
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            para = doc.add_paragraph("这是正文段落")
            # Set initial font to Arial on the run
            for run in para.runs:
                run.font.name = "Arial"
                # Also set eastAsia font explicitly
                from lxml import etree
                W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                rPr = run._element.find(f"{{{W}}}rPr")
                if rPr is None:
                    rPr = etree.SubElement(run._element, f"{{{W}}}rPr")
                rFonts = rPr.find(f"{{{W}}}rFonts")
                if rFonts is None:
                    rFonts = etree.SubElement(rPr, f"{{{W}}}rFonts")
                rFonts.set(f"{{{W}}}eastAsia", "Arial")
            doc.save(path)

            # Verify initial state: font is Arial
            initial_doc = Document(path)
            initial_font = _get_run_eastasia_font(initial_doc.paragraphs[0].runs[0])
            assert initial_font == "Arial", f"Expected initial font Arial, got {initial_font}"

            # Apply settings with different font
            settings = FormatSettings(body=BodySettings(font="黑体", size="小四", spacing=1.5))
            out, _ = apply_settings_to_document(path, settings)

            # Verify output
            result = Document(out)
            assert len(result.paragraphs) == 1
            run = result.paragraphs[0].runs[0]
            new_font = _get_run_eastasia_font(run)
            assert new_font == "黑体", f"Expected font '黑体' after apply, got {new_font!r}"

    def test_body_size_actually_changes(self):
        """Verify apply_settings_to_document actually changes paragraph font size."""
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            para = doc.add_paragraph("这是正文段落")
            # Set initial size to 16pt (三号)
            for run in para.runs:
                run.font.size = Pt(16)
            doc.save(path)

            # Apply settings with size 五号 (10.5pt)
            settings = FormatSettings(body=BodySettings(font="宋体", size="五号", spacing=1.5))
            out, _ = apply_settings_to_document(path, settings)

            result = Document(out)
            run = result.paragraphs[0].runs[0]
            # Size should be 10.5pt (五号)
            assert run.font.size is not None
            assert abs(run.font.size.pt - 10.5) < 0.1, f"Expected size ~10.5pt, got {run.font.size.pt}pt"

    def test_auto_detect_disabled_does_not_format_text_titles(self):
        """When auto_detect_numeric_titles=False, paragraphs with numeric text
        but no outline level should NOT be treated as titles."""
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            doc.add_paragraph("1 看起来像标题")
            doc.save(path)

            settings = FormatSettings(auto_detect_numeric_titles=False)
            out, _ = apply_settings_to_document(path, settings)

            result = Document(out)
            para = result.paragraphs[0]
            # Should have NO outline level (not treated as title)
            pPr = para._element.find(
                ".//w:pPr",
                {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
            )
            if pPr is not None:
                outline = pPr.find(
                    "w:outlineLvl",
                    {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
                )
                assert outline is None, "Paragraph without outline level should not get one when auto_detect=False"

    def test_respects_existing_outline_levels(self):
        """Paragraphs with outline levels 0-4 should be formatted as titles
        regardless of auto_detect setting."""
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            para = doc.add_paragraph("我的标题")
            # Manually set outline level 2 (level 3 in UI)
            from tvba_core_oox import set_outline_level
            set_outline_level(para, 2)
            doc.save(path)

            settings = FormatSettings(auto_detect_numeric_titles=False)
            out, _ = apply_settings_to_document(path, settings)

            result = Document(out)
            para = result.paragraphs[0]
            # Should keep outline level 2
            pPr = para._element.find(
                ".//w:pPr",
                {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
            )
            outline = pPr.find(
                "w:outlineLvl",
                {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
            )
            assert outline is not None
            assert outline.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "2"

    def test_heading_style_recognized_as_title(self):
        """Paragraphs with Word heading styles (Heading 1, Heading 2, etc.)
        should be recognized and formatted as titles."""
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            para = doc.add_paragraph("一级标题")
            para.style = doc.styles["Heading 1"]
            doc.save(path)

            settings = FormatSettings(auto_detect_numeric_titles=False)
            out, _ = apply_settings_to_document(path, settings)

            result = Document(out)
            para = result.paragraphs[0]
            # Should have outline level 0 (from Heading 1 style)
            pPr = para._element.find(
                ".//w:pPr",
                {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
            )
            outline = pPr.find(
                "w:outlineLvl",
                {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
            )
            assert outline is not None, "Heading 1 style should result in outline level"
            assert outline.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "0"

    def test_heading_2_style_formatted_as_level_2(self):
        """Paragraphs with Heading 2 style should be formatted as level 2 title."""
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            para = doc.add_paragraph("二级标题")
            para.style = doc.styles["Heading 2"]
            doc.save(path)

            settings = FormatSettings(auto_detect_numeric_titles=False)
            out, _ = apply_settings_to_document(path, settings)

            result = Document(out)
            para = result.paragraphs[0]
            # Should have outline level 1 (from Heading 2 style, 0-indexed)
            pPr = para._element.find(
                ".//w:pPr",
                {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
            )
            outline = pPr.find(
                "w:outlineLvl",
                {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
            )
            assert outline is not None, "Heading 2 style should result in outline level"
            assert outline.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "1"

    def test_figure_caption_has_no_indent(self):
        """Figure captions should not inherit body first-line indent."""
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            doc.add_paragraph("图2.2.11-1  管道经过南朗遗物点位置示意图")
            doc.save(path)

            settings = FormatSettings(
                body=BodySettings(
                    font="宋体", size="小四", spacing=1.5,
                    special_indent="首行缩进", special_indent_chars=2.0,
                ),
                figure=FigureSettings(
                    title_font="黑体", title_size="小四", title_bold=True, title_spacing=1.5,
                ),
            )
            out, _ = apply_settings_to_document(path, settings)

            result = Document(out)
            para = result.paragraphs[0]
            pPr = para._element.find(
                ".//w:pPr",
                {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
            )
            ind = pPr.find("w:ind", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if ind is not None:
                first_line = ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}firstLine")
                hanging = ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hanging")
                assert first_line is None, f"Figure caption should not have first-line indent, got {first_line}"
                assert hanging is None, f"Figure caption should not have hanging indent, got {hanging}"

    def test_table_caption_has_no_indent(self):
        """Table captions should not inherit body first-line indent."""
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            doc.add_paragraph("表2.2.9-1  沿线经过环境敏感点统计表")
            doc.add_table(rows=2, cols=2)  # Table needed for caption to be found
            doc.save(path)

            settings = FormatSettings(
                body=BodySettings(
                    font="宋体", size="小四", spacing=1.5,
                    special_indent="首行缩进", special_indent_chars=2.0,
                ),
                table=TableSettings(
                    title_font="黑体", title_size="小四", title_bold=True, title_spacing=1.5,
                ),
            )
            out, _ = apply_settings_to_document(path, settings)

            result = Document(out)
            para = result.paragraphs[0]
            pPr = para._element.find(
                ".//w:pPr",
                {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
            )
            ind = pPr.find("w:ind", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if ind is not None:
                first_line = ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}firstLine")
                hanging = ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hanging")
                assert first_line is None, f"Table caption should not have first-line indent, got {first_line}"
                assert hanging is None, f"Table caption should not have hanging indent, got {hanging}"

