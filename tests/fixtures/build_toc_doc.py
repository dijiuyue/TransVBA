"""Build a test .docx with TOC paragraphs."""
from docx import Document
import sys
from pathlib import Path


def build(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章  绪论\t1")
    doc.add_paragraph("  1.1  研究背景\t2")
    doc.add_paragraph("  1.2  研究意义\t3")
    doc.add_paragraph("    1.2.1  理论意义\t3")
    doc.add_paragraph("第二章  相关工作\t5")
    doc.save(path)


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("toc_test.docx")
    build(out)
    print(f"Built: {out}")
