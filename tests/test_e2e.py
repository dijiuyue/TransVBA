import tempfile
from pathlib import Path
import pytest
from docx import Document

from tvba_core_document import apply_settings_to_document
from tvba_settings import FormatSettings


class TestEndToEnd:
    def test_full_document_processing(self):
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "full.docx"
            doc = Document()
            doc.add_paragraph("目录")
            doc.add_paragraph("第一章\t1")
            doc.add_paragraph("1 引言")
            doc.add_paragraph("1.1 背景")
            doc.add_paragraph("正文段落")
            doc.add_paragraph("表 1-1 表格")
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "单元格"
            doc.add_paragraph("图 1-1 图片")
            doc.save(path)

            settings = FormatSettings()
            out, _ = apply_settings_to_document(path, settings)

            result = Document(out)
            assert len(result.paragraphs) >= 7

            # Title should have outline level
            title_para = None
            for p in result.paragraphs:
                if p.text.startswith("1 引言"):
                    title_para = p
                    break
            assert title_para is not None
            pPr = title_para._element.find(
                ".//w:pPr",
                {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            )
            outline = pPr.find("w:outlineLvl", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            assert outline is not None

    def test_title_levels_correct(self):
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "levels.docx"
            doc = Document()
            doc.add_paragraph("1 一级")
            doc.add_paragraph("1.1 二级")
            doc.add_paragraph("1.1.1 三级")
            doc.add_paragraph("1.1.1.1 四级")
            doc.add_paragraph("1.1.1.1.1 五级")
            doc.save(path)

            settings = FormatSettings()
            out, _ = apply_settings_to_document(path, settings)
            result = Document(out)

            expected_levels = ["0", "1", "2", "3", "4"]
            for i, expected in enumerate(expected_levels):
                pPr = result.paragraphs[i]._element.find(
                    ".//w:pPr",
                    {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                )
                outline = pPr.find("w:outlineLvl", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                assert outline is not None
                actual = outline.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                assert actual == expected, f"Paragraph {i}: expected outline level {expected}, got {actual}"
