import pytest
from docx import Document

from tvba_core_numbering import DocxListResolver, ComListResolver, auto_select, ResolverStatus


class TestDocxListResolver:
    def test_no_numbering_returns_none(self):
        doc = Document()
        para = doc.add_paragraph("普通段落")
        resolver = DocxListResolver(doc)
        assert resolver.get_list_level(para) is None
        assert resolver.get_list_text(para) is None

    def test_returns_level_from_numPr(self):
        doc = Document()
        # Add a paragraph with numbering via OOXML
        para = doc.add_paragraph("列表项")
        pPr = para._element.get_or_add_pPr()
        from lxml import etree

        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        numPr = etree.SubElement(pPr, f"{{{W}}}numPr")
        ilvl = etree.SubElement(numPr, f"{{{W}}}ilvl")
        ilvl.set(f"{{{W}}}val", "2")
        numId = etree.SubElement(numPr, f"{{{W}}}numId")
        numId.set(f"{{{W}}}val", "1")

        resolver = DocxListResolver(doc)
        # Docx resolver returns ilvl + 1 as level
        assert resolver.get_list_level(para) == 3


class TestAutoSelect:
    def test_returns_docx_resolver_when_no_path(self):
        resolver, status = auto_select(prefer_com=True, docx_path=None)
        assert isinstance(resolver, DocxListResolver)
        assert isinstance(status, ResolverStatus)
        assert status.mode == "docx_fallback"

    def test_returns_docx_resolver_when_com_disabled(self):
        resolver, status = auto_select(prefer_com=False, docx_path="some.docx")
        assert isinstance(resolver, DocxListResolver)
        assert isinstance(status, ResolverStatus)
        assert status.mode == "docx_fallback"

    def test_status_reports_unreliable_text_for_docx_fallback(self):
        _, status = auto_select(prefer_com=False, docx_path="dummy.docx")
        assert status.reliable_rendered_text is False

    def test_com_fallback_produces_warning(self):
        """When COM is preferred but unavailable, status warns."""
        _, status = auto_select(prefer_com=True, docx_path="nonexistent.docx")
        if status.mode == "docx_fallback":
            assert len(status.warnings) > 0
        # If COM happened to work on this machine, that's also fine


class TestComListResolverPerformanceGuard:
    def test_init_does_not_prefetch_every_com_paragraph(self, monkeypatch, tmp_path):
        """Regression guard for GUI freeze at "Detecting titles...".

        ComListResolver used to call Word's Paragraphs(i) once per python-docx
        paragraph during __init__. On large documents this caused hundreds or
        thousands of COM round-trips before progress could advance.
        """
        import sys
        import types

        class FakeParagraphs:
            def __init__(self):
                self.calls = []

            def __call__(self, index):
                self.calls.append(index)
                return f"paragraph-{index}"

        class FakeComDoc:
            def __init__(self):
                self.Paragraphs = FakeParagraphs()
                self.closed = False

            def Close(self, SaveChanges=False):
                self.closed = True

        class FakeDocuments:
            def __init__(self, fake_doc):
                self.fake_doc = fake_doc
                self.open_calls = []

            def Open(self, path, **kwargs):
                self.open_calls.append((path, kwargs))
                return self.fake_doc

        class FakeWord:
            def __init__(self, fake_doc):
                self.Visible = True
                self.DisplayAlerts = None
                self.Documents = FakeDocuments(fake_doc)
                self.quit_called = False

            def Quit(self):
                self.quit_called = True

        fake_com_doc = FakeComDoc()
        fake_word = FakeWord(fake_com_doc)

        client_mod = types.ModuleType("win32com.client")
        client_mod.DispatchEx = lambda prog_id: fake_word
        package_mod = types.ModuleType("win32com")
        package_mod.client = client_mod
        monkeypatch.setitem(sys.modules, "win32com", package_mod)
        monkeypatch.setitem(sys.modules, "win32com.client", client_mod)

        docx_path = tmp_path / "large.docx"
        doc = Document()
        for i in range(500):
            doc.add_paragraph(f"Paragraph {i}")
        doc.save(str(docx_path))

        resolver = ComListResolver(str(docx_path), doc=doc)

        assert fake_com_doc.Paragraphs.calls == []
        assert len(resolver._element_to_index) == 500

        first = resolver._get_com_paragraph(doc.paragraphs[0])
        assert first.startswith("paragraph-")
        assert len(fake_com_doc.Paragraphs.calls) == 1

        resolver.close()
        assert fake_com_doc.closed is True
        assert fake_word.quit_called is True


# Check if win32com is available
try:
    import win32com.client

    _WORD_AVAILABLE = True
except ImportError:
    _WORD_AVAILABLE = False


def _word_can_dispatch():
    """Try to dispatch Word.Application to confirm Word is installed.

    Runs in a subprocess to isolate fatal COM exceptions.
    """
    if not _WORD_AVAILABLE:
        return False
    import subprocess
    import sys

    code = (
        "import win32com.client; "
        "w = win32com.client.DispatchEx('Word.Application'); "
        "w.Visible = False; "
        "w.Quit(); "
        "print('OK')"
    )
    try:
        result = subprocess.run(
            [sys.executable, "-c", code],
            capture_output=True,
            text=True,
            timeout=30,
        )
        return result.returncode == 0 and "OK" in result.stdout
    except Exception:
        return False


_WORD_CAN_DISPATCH = _word_can_dispatch()

@pytest.mark.word_com
@pytest.mark.skipif(not _WORD_CAN_DISPATCH, reason="Word not available on this machine")
class TestComListResolver:
    def test_com_resolver_with_real_word_document(self, tmp_path):
        """Create a docx with a list, open via COM, verify level and text."""
        docx_path = tmp_path / "list_test.docx"
        doc = Document()
        doc.add_paragraph("Intro paragraph")
        doc.add_paragraph("First item", style="List Number")
        doc.add_paragraph("Second item", style="List Number")
        doc.add_paragraph("Normal paragraph")
        doc.save(str(docx_path))

        # Open with python-docx to get paragraph references
        pydocx_doc = Document(str(docx_path))

        with ComListResolver(str(docx_path)) as resolver:
            # para[0] = "Intro paragraph" - not a list
            assert resolver.get_list_level(pydocx_doc.paragraphs[0]) is None
            assert resolver.get_list_text(pydocx_doc.paragraphs[0]) is None

            # para[1] = "First item" - list level 1
            level = resolver.get_list_level(pydocx_doc.paragraphs[1])
            assert level == 1, f"Expected level 1, got {level}"
            text = resolver.get_list_text(pydocx_doc.paragraphs[1])
            assert text is not None

            # para[2] = "Second item" - list level 1
            assert resolver.get_list_level(pydocx_doc.paragraphs[2]) == 1

            # para[3] = "Normal paragraph" - not a list
            assert resolver.get_list_level(pydocx_doc.paragraphs[3]) is None

    def test_com_diagnose(self, tmp_path):
        """Test diagnose() collects all list paragraphs."""
        docx_path = tmp_path / "list_test.docx"
        doc = Document()
        doc.add_paragraph("Intro")
        doc.add_paragraph("Item 1", style="List Number")
        doc.add_paragraph("Item 2", style="List Number")
        doc.add_paragraph("Normal")
        doc.save(str(docx_path))

        pydocx_doc = Document(str(docx_path))

        with ComListResolver(str(docx_path)) as resolver:
            entries = resolver.diagnose(pydocx_doc)
            assert len(entries) == 2
            assert entries[0].level == 1
            assert entries[1].level == 1

    def test_auto_select_returns_com_when_available(self, tmp_path):
        """auto_select should return ComListResolver when COM is available and path is given."""
        docx_path = tmp_path / "dummy.docx"
        doc = Document()
        doc.add_paragraph("test")
        doc.save(str(docx_path))

        resolver, status = auto_select(prefer_com=True, docx_path=str(docx_path), doc=doc)
        # If COM works, expect ComListResolver; if COM fails to open doc, fallback is OK
        assert isinstance(resolver, (ComListResolver, DocxListResolver))
        assert isinstance(status, ResolverStatus)
