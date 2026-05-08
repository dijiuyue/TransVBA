import pytest
from docx import Document

from tvba_core_toc import (
    is_toc_entry_line,
    is_toc_title_line,
    is_toc_paragraph,
    identify_toc_level,
    apply_toc_title_style,
    apply_toc_entry_style,
    refresh_toc,
)
from tvba_settings import TocLegacyFixedDefaults


class TestIsTocEntryLine:
    def test_tab_and_page_number(self):
        assert is_toc_entry_line("第一章\t1") is True

    def test_no_tab(self):
        assert is_toc_entry_line("第一章 1") is False

    def test_tab_but_no_number(self):
        assert is_toc_entry_line("第一章\t") is False

    def test_multiple_tabs(self):
        assert is_toc_entry_line("第一章\t\t1") is True

    def test_page_number_with_suffix(self):
        assert is_toc_entry_line("第一章\t1\r") is True


class TestIsTocTitleLine:
    def test_exact_directory(self):
        assert is_toc_title_line("目录") is True

    def test_with_spaces(self):
        assert is_toc_title_line("  目录  ") is True

    def test_other_text(self):
        assert is_toc_title_line("第一章") is False


class TestIsTocParagraph:
    def test_toc_entry_text(self):
        doc = Document()
        para = doc.add_paragraph("第一章\t1")
        assert is_toc_paragraph(para) is True

    def test_toc_title_text(self):
        doc = Document()
        para = doc.add_paragraph("目录")
        assert is_toc_paragraph(para) is True

    def test_toc1_style(self):
        doc = Document()
        para = doc.add_paragraph("Some text")
        para.style = doc.styles.add_style("TOC1", 1)
        assert is_toc_paragraph(para) is True

    def test_toc2_style(self):
        doc = Document()
        para = doc.add_paragraph("Some text")
        para.style = doc.styles.add_style("TOC2", 1)
        assert is_toc_paragraph(para) is True

    def test_toc3_style(self):
        doc = Document()
        para = doc.add_paragraph("Some text")
        para.style = doc.styles.add_style("TOC3", 1)
        assert is_toc_paragraph(para) is True

    def test_non_toc_style(self):
        doc = Document()
        para = doc.add_paragraph("Some text")
        # Default style is "Normal" which does not contain "TOC"
        assert is_toc_paragraph(para) is False

    def test_toc_style_case_insensitive(self):
        doc = Document()
        para = doc.add_paragraph("Some text")
        para.style = doc.styles.add_style("toc Heading", 1)
        assert is_toc_paragraph(para) is True


class TestIdentifyTocLevel:
    def test_level_1_no_indent(self):
        assert identify_toc_level("第一章\t1") == 1

    def test_level_2_two_spaces(self):
        assert identify_toc_level("  1.1\t2") == 2

    def test_level_3_four_spaces(self):
        assert identify_toc_level("    1.1.1\t3") == 3

    def test_level_0_unknown_indent(self):
        assert identify_toc_level("     1.1\t2") == 0

    def test_level_1_from_number(self):
        assert identify_toc_level("1\tIntroduction\t1") == 1

    def test_level_2_from_number(self):
        assert identify_toc_level("1.1\tBackground\t2") == 2

    def test_level_3_from_number(self):
        assert identify_toc_level("1.1.2\tDetails\t3") == 3

    def test_level_from_number_with_spaces(self):
        assert identify_toc_level("  2.1\tSection\t5") == 2

    def test_level_from_fullwidth_number(self):
        assert identify_toc_level("１．１\tSection\t5") == 2


class TestApplyTocTitleStyle:
    def test_applies_bold_and_font(self):
        doc = Document()
        para = doc.add_paragraph("目录")
        defaults = TocLegacyFixedDefaults()
        apply_toc_title_style(para, defaults)
        run = para.runs[0]
        assert run.font.bold is True


class TestRefreshToc:
    def test_formats_toc_entries(self):
        doc = Document()
        doc.add_paragraph("目录")
        doc.add_paragraph("第一章\t1")
        doc.add_paragraph("  1.1\t2")
        defaults = TocLegacyFixedDefaults()
        refresh_toc(doc, defaults)
        # Title should be bold
        title_para = doc.paragraphs[0]
        assert title_para.runs[0].font.bold is True
