"""Tests for presidential order number conversion."""
from docx import Document

from tvba_core_presidential import (
    number_to_chinese,
    format_presidential_order_numbers,
    check_presidential_order_numbers,
    _PRESIDENTIAL_RE,
)


class TestNumberToChinese:
    def test_single_digit(self):
        assert number_to_chinese(1) == "一"
        assert number_to_chinese(5) == "五"
        assert number_to_chinese(9) == "九"

    def test_ten(self):
        assert number_to_chinese(10) == "十"

    def test_ten_plus(self):
        assert number_to_chinese(11) == "十一"
        assert number_to_chinese(19) == "十九"

    def test_twenty(self):
        assert number_to_chinese(20) == "二十"

    def test_twenty_plus(self):
        assert number_to_chinese(25) == "二十五"

    def test_eighty_one(self):
        assert number_to_chinese(81) == "八十一"

    def test_round_hundred(self):
        assert number_to_chinese(100) == "一百"

    def test_hundred_plus(self):
        assert number_to_chinese(101) == "一百零一"
        assert number_to_chinese(110) == "一百一十"
        assert number_to_chinese(111) == "一百一十一"

    def test_round_thousand(self):
        assert number_to_chinese(1000) == "一千"

    def test_thousand_plus(self):
        assert number_to_chinese(1001) == "一千零一"
        assert number_to_chinese(2020) == "二千零二十"


class TestRegexPattern:
    def test_matches_basic(self):
        m = _PRESIDENTIAL_RE.search("第81号")
        assert m is not None
        assert m.group(1) == "81"

    def test_matches_with_prefix(self):
        m = _PRESIDENTIAL_RE.search("主席令第81号")
        assert m is not None
        assert m.group(1) == "81"

    def test_matches_full_title(self):
        m = _PRESIDENTIAL_RE.search("中华人民共和国主席令第81号")
        assert m is not None
        assert m.group(1) == "81"

    def test_not_match_already_chinese(self):
        assert _PRESIDENTIAL_RE.search("第八十一号") is None

    def test_not_match_ordinary_number(self):
        assert _PRESIDENTIAL_RE.search("第3章") is None


class TestFormatPresidentialOrder:
    def test_converts_basic(self):
        doc = Document()
        doc.add_paragraph("第81号")
        n = format_presidential_order_numbers(doc)
        assert n > 0
        assert "第八十一号" in doc.paragraphs[0].text

    def test_converts_with_prefix(self):
        doc = Document()
        doc.add_paragraph("主席令第10号")
        format_presidential_order_numbers(doc)
        assert "主席令第十号" in doc.paragraphs[0].text

    def test_no_double_convert(self):
        doc = Document()
        doc.add_paragraph("第八十一号")
        n = format_presidential_order_numbers(doc)
        assert n == 0

    def test_leaves_ordinary_numbers(self):
        doc = Document()
        doc.add_paragraph("普通文本中的第3章数字")
        n = format_presidential_order_numbers(doc)
        assert n == 0


class TestCheckPresidentialOrder:
    def test_flags_arabic_order(self):
        from tvba_core_validate import ValidationIssue
        doc = Document()
        doc.add_paragraph("第81号")
        issues: list[ValidationIssue] = []
        check_presidential_order_numbers(doc.paragraphs, issues)
        assert len(issues) > 0

    def test_passes_chinese_order(self):
        from tvba_core_validate import ValidationIssue
        doc = Document()
        doc.add_paragraph("第八十一号")
        issues: list[ValidationIssue] = []
        check_presidential_order_numbers(doc.paragraphs, issues)
        assert len(issues) == 0
