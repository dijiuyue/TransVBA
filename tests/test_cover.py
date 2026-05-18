"""TDD tests for cover title formatting (tvba_core_cover.py).

Requirements ref: 多模板自动格式化与检查插件需求说明书 section 4.2
Mode B: 封面标题 = 宋体/TNR, 二号, 加粗, 1.5倍, 居中, 仅首页大标题
"""
import pytest
from docx import Document
from lxml import etree

from tvba_core_cover import format_cover_title
from tvba_core_oox import set_outline_level
from tvba_settings import CoverSettings

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _get_run_font_size_pt(para) -> float | None:
    """Extract font size in points from the first run's w:sz."""
    sz = para._element.find(f".//{{{W}}}r/{{{W}}}rPr/{{{W}}}sz")
    if sz is None:
        return None
    half_pts = int(sz.get(f"{{{W}}}val"))
    return half_pts / 2


def _get_run_bold(para) -> bool:
    b = para._element.find(f".//{{{W}}}r/{{{W}}}rPr/{{{W}}}b")
    return b is not None and b.get(f"{{{W}}}val") != "0"


def _get_run_eastasia_font(para) -> str | None:
    rFonts = para._element.find(f".//{{{W}}}r/{{{W}}}rPr/{{{W}}}rFonts")
    if rFonts is None:
        return None
    return rFonts.get(f"{{{W}}}eastAsia")


def _get_line_spacing(para) -> str | None:
    spacing = para._element.find(f".//{{{W}}}pPr/{{{W}}}spacing")
    if spacing is None:
        return None
    return spacing.get(f"{{{W}}}line")


class TestFormatCoverTitle:
    """RED phase: cover title formatting per requirements section 4.2 (Mode B)."""

    def test_formats_centered_short_text_as_cover_title(self):
        """First centered short paragraph before TOC → 二号/加粗/居中/宋体."""
        doc = Document()
        # Simulate a cover page: centered title
        para = doc.add_paragraph("某可行性研究报告")
        para.alignment = 1  # center

        format_cover_title(doc)

        # should be formatted as 二号 (22pt)
        size_pt = _get_run_font_size_pt(para)
        assert size_pt == 22.0, f"Expected 二号 (22pt), got {size_pt}pt"

        # should be bold
        assert _get_run_bold(para) is True

        # should stay centered
        assert para.alignment == 1

        # should use 宋体 for eastAsian
        east = _get_run_eastasia_font(para)
        assert east == "宋体", f"Expected 宋体, got {east}"

    def test_applies_1_5_line_spacing(self):
        """Cover title line spacing should be 1.5倍 (360 twips)."""
        doc = Document()
        para = doc.add_paragraph("测试报告")
        para.alignment = 1

        format_cover_title(doc)

        line_val = _get_line_spacing(para)
        assert line_val is not None
        # 1.5×240 = 360 twips
        assert int(line_val) == 360, f"Expected 360 twips (1.5倍), got {line_val}"

    def test_skips_paragraphs_with_outline_level(self):
        """Paragraphs that already have outline level (headings) should NOT be treated as cover."""
        doc = Document()
        para = doc.add_paragraph("1 项目概述")  # looks like a heading
        para.alignment = 1
        set_outline_level(para, 0)  # Level 1 heading

        format_cover_title(doc)

        # Should NOT be formatted as cover title (二号)
        size = _get_run_font_size_pt(para)
        assert size != 22.0, f"Heading should not be formatted as cover title"

    def test_skips_non_centered_text(self):
        """Only centered text qualifies as cover title."""
        doc = Document()
        para = doc.add_paragraph("左对齐文本不是封面标题")
        para.alignment = 0  # left

        format_cover_title(doc)

        size = _get_run_font_size_pt(para)
        assert size != 22.0, f"Left-aligned text should not be cover title"

    def test_stops_at_toc_boundary(self):
        """Cover title detection stops when hitting TOC entries."""
        doc = Document()
        # Paragraph AFTER a TOC entry should not be treated as cover
        toc = doc.add_paragraph("目录")
        para = doc.add_paragraph("某报告标题")
        para.alignment = 1

        format_cover_title(doc)

        # para is AFTER "目录" heading, but the function searches
        # sequentially and checks only centered paragraphs before TOC.
        # Since "目录" kills the search at position 0, para at position 1
        # should NOT be formatted.
        # Actually: the function breaks when it sees is_toc_title_line.
        # "目录" matches is_toc_title_line, so break happens at index 0.
        size = _get_run_font_size_pt(para)
        assert size != 22.0, f"Text after TOC should not be cover title"

    def test_skips_text_too_short(self):
        """Text shorter than 2 chars should not be cover title."""
        doc = Document()
        para = doc.add_paragraph("X")
        para.alignment = 1

        format_cover_title(doc)

        size = _get_run_font_size_pt(para)
        assert size != 22.0

    def test_skips_text_too_long(self):
        """Text longer than 80 chars should not be cover title."""
        doc = Document()
        para = doc.add_paragraph("A" * 81)
        para.alignment = 1

        format_cover_title(doc)

        size = _get_run_font_size_pt(para)
        assert size != 22.0

    def test_skips_digit_only_text(self):
        """Pure digit/dot strings should not be cover title."""
        doc = Document()
        para = doc.add_paragraph("1.2.3")
        para.alignment = 1

        format_cover_title(doc)

        size = _get_run_font_size_pt(para)
        assert size != 22.0

    def test_only_formats_first_matching_paragraph(self):
        """Only the first centered, short paragraph should be formatted."""
        doc = Document()
        title1 = doc.add_paragraph("报告标题一")
        title1.alignment = 1
        title2 = doc.add_paragraph("副标题")
        title2.alignment = 1

        format_cover_title(doc)

        # First title: formatted
        assert _get_run_font_size_pt(title1) == 22.0
        # Second title: NOT formatted (break after first match)
        size2 = _get_run_font_size_pt(title2)
        assert size2 != 22.0, f"Second centered paragraph should not be cover title"


class TestCoverTitleDetectionEdgeCases:
    """Edge cases for cover title detection per requirements."""

    def test_empty_document_does_not_crash(self):
        """format_cover_title should handle empty documents gracefully."""
        doc = Document()
        # No paragraphs at all — should just return without error
        format_cover_title(doc)

    def test_no_centered_paragraphs(self):
        """If no paragraph is centered, no formatting should occur."""
        doc = Document()
        doc.add_paragraph("左对齐文字一")
        doc.add_paragraph("左对齐文字二")

        format_cover_title(doc)

        # Neither should be 22pt
        for para in doc.paragraphs:
            size = _get_run_font_size_pt(para)
            assert size != 22.0

    def test_centered_but_long_text_not_cover(self):
        """Long centered text (>80 chars) is description, not title."""
        doc = Document()
        # Generate text over 80 characters — body paragraph, not a title
        long_text = "本项目位于某某省某某市某某县，项目起点位于某某镇某某村，终点位于某某乡某某村，路线全长约120.5公里，设计时速80公里，路基宽度24.5米，双向四车道高速公路标准建设，项目总投资约50亿元。"
        assert len(long_text) > 80, f"Test text must be >80 chars, got {len(long_text)}"
        para = doc.add_paragraph(long_text)
        para.alignment = 1

        format_cover_title(doc)

        size = _get_run_font_size_pt(para)
        assert size != 22.0, "Text >80 chars should not be cover title"


class TestFormatCoverTitleWithSettings:
    """Settings-aware cover formatting — template-driven values."""

    def test_uses_custom_font_size(self):
        """Cover title should use font size from settings, not hardcoded 二号."""
        doc = Document()
        para = doc.add_paragraph("测试报告标题")
        para.alignment = 1

        settings = CoverSettings(size="三号")  # 16pt
        format_cover_title(doc, settings)

        size_pt = _get_run_font_size_pt(para)
        assert size_pt == 16.0, f"Expected 三号 (16pt) from settings, got {size_pt}pt"

    def test_uses_custom_font(self):
        doc = Document()
        para = doc.add_paragraph("测试报告标题")
        para.alignment = 1

        settings = CoverSettings(font="黑体")
        format_cover_title(doc, settings)

        east = _get_run_eastasia_font(para)
        assert east == "黑体", f"Expected 黑体 from settings, got {east}"

    def test_uses_custom_bold(self):
        doc = Document()
        para = doc.add_paragraph("测试报告标题")
        para.alignment = 1

        settings = CoverSettings(bold=False)
        format_cover_title(doc, settings)

        assert _get_run_bold(para) is False

    def test_uses_custom_line_spacing(self):
        doc = Document()
        para = doc.add_paragraph("测试报告标题")
        para.alignment = 1

        settings = CoverSettings(line_spacing=2.0)
        format_cover_title(doc, settings)

        line_val = _get_line_spacing(para)
        assert int(line_val) == 480, f"Expected 480 twips (2.0倍), got {line_val}"

    def test_uses_custom_max_length(self):
        """Custom max_length should control character limit for detection."""
        doc = Document()
        para = doc.add_paragraph("A" * 50)
        para.alignment = 1

        settings = CoverSettings(max_length=40)
        format_cover_title(doc, settings)

        size = _get_run_font_size_pt(para)
        assert size != 22.0, f"50-char text should be skipped with max_length=40"

    def test_respects_custom_search_paragraphs(self):
        """Custom search_paragraphs should limit how many paragraphs are scanned."""
        doc = Document()
        # Fill with non-matching paragraphs first
        for _ in range(5):
            doc.add_paragraph("普通段落")
        title = doc.add_paragraph("报告标题")
        title.alignment = 1

        # search_paragraphs=3 means only first 3 paragraphs scanned
        settings = CoverSettings(search_paragraphs=3)
        format_cover_title(doc, settings)

        size = _get_run_font_size_pt(title)
        assert size != 22.0, f"Title at position 6 should be skipped with search_paragraphs=3"

    def test_none_settings_uses_defaults(self):
        """Passing None should use CoverSettings() defaults (二号/宋体/加粗)."""
        doc = Document()
        para = doc.add_paragraph("报告标题")
        para.alignment = 1

        format_cover_title(doc, None)

        assert _get_run_font_size_pt(para) == 22.0
        assert _get_run_bold(para) is True
