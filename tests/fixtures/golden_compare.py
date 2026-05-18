"""Golden file comparison helpers.

Compare key OOXML properties between actual output and golden reference,
without doing binary full-diff. Focuses on the formatting attributes that
the VBA port is designed to produce.
"""
from pathlib import Path
from dataclasses import dataclass, field


@dataclass
class Diff:
    """A single difference found during comparison."""
    path: str  # e.g. "paragraphs[3].runs[0].font.eastAsia"
    expected: str
    actual: str


@dataclass
class CompareResult:
    """Result of comparing an actual document against a golden reference."""
    matched: bool
    diffs: list[Diff] = field(default_factory=list)

    def report(self) -> str:
        if self.matched:
            return "OK: actual matches golden."
        lines = [f"FAIL: {len(self.diffs)} difference(s):"]
        for d in self.diffs:
            lines.append(f"  - {d.path}: expected={d.expected}, actual={d.actual}")
        return "\n".join(lines)


def compare_docx_properties(actual_path: Path, golden_path: Path) -> CompareResult:
    """Compare formatting properties of two docx files.

    Checks:
    - Paragraph outline levels
    - Run fonts (ascii, hAnsi, eastAsia)
    - Run sizes (sz, szCs)
    - Run bold
    - Paragraph alignment (jc)
    - Paragraph indentation (left, right, firstLine, hanging)
    - Paragraph spacing (beforeLines, afterLines, line)
    - Section count
    - Table count
    - Table row height
    - Table borders

    Does NOT compare:
    - Text content (VBA vs Python may differ in normalization order)
    - Binary OPC structure
    - Word rendering-dependent properties
    """
    from docx import Document

    actual_doc = Document(str(actual_path))
    golden_doc = Document(str(golden_path))
    diffs: list[Diff] = []

    _compare_sections(actual_doc, golden_doc, diffs)
    _compare_paragraphs(actual_doc, golden_doc, diffs)
    _compare_tables(actual_doc, golden_doc, diffs)

    return CompareResult(matched=len(diffs) == 0, diffs=diffs)


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _ns(tag: str) -> str:
    return f"{{{W}}}{tag}"


def _compare_sections(actual, golden, diffs):
    actual_sec = len(actual.sections)
    golden_sec = len(golden.sections)
    if actual_sec != golden_sec:
        diffs.append(Diff("sections.count", str(golden_sec), str(actual_sec)))


def _compare_paragraphs(actual, golden, diffs):
    a_paras = list(actual.paragraphs)
    g_paras = list(golden.paragraphs)

    n = min(len(a_paras), len(g_paras))
    for i in range(n):
        ap = a_paras[i]
        gp = g_paras[i]
        base = f"paragraphs[{i}]"

        # Outline level
        a_level = _get_outline_level(ap)
        g_level = _get_outline_level(gp)
        if a_level != g_level:
            diffs.append(Diff(f"{base}.outline_level", str(g_level), str(a_level)))

        # Paragraph indentation
        _compare_para_indent(ap, gp, base, diffs)

        # Paragraph spacing
        _compare_para_spacing(ap, gp, base, diffs)

        # Alignment (jc)
        _compare_para_alignment(ap, gp, base, diffs)

        # Compare runs
        a_runs = list(ap.runs)
        g_runs = list(gp.runs)
        nr = min(len(a_runs), len(g_runs))
        for j in range(nr):
            ar = a_runs[j]
            gr = g_runs[j]
            rbase = f"{base}.runs[{j}]"

            # Run fonts (ascii, hAnsi, eastAsia via XML)
            _compare_run_fonts(ar, gr, rbase, diffs)

            # Run sizes (sz, szCs)
            _compare_run_sizes(ar, gr, rbase, diffs)

            # Bold
            a_bold = ar.font.bold
            g_bold = gr.font.bold
            if a_bold != g_bold:
                diffs.append(Diff(f"{rbase}.font.bold", str(g_bold), str(a_bold)))


def _compare_para_indent(ap, gp, base, diffs):
    """Compare w:pPr/w:ind attributes."""
    a_ind = _get_indent(ap)
    g_ind = _get_indent(gp)
    for attr in ("left", "right", "firstLine", "hanging"):
        a_val = a_ind.get(attr)
        g_val = g_ind.get(attr)
        if a_val != g_val:
            diffs.append(Diff(f"{base}.indent.{attr}", str(g_val), str(a_val)))


def _compare_para_spacing(ap, gp, base, diffs):
    """Compare w:pPr/w:spacing attributes."""
    a_spc = _get_spacing(ap)
    g_spc = _get_spacing(gp)
    for attr in ("beforeLines", "afterLines", "line"):
        a_val = a_spc.get(attr)
        g_val = g_spc.get(attr)
        if a_val != g_val:
            diffs.append(Diff(f"{base}.spacing.{attr}", str(g_val), str(a_val)))


def _compare_para_alignment(ap, gp, base, diffs):
    """Compare w:pPr/w:jc."""
    a_jc = _get_alignment(ap)
    g_jc = _get_alignment(gp)
    if a_jc != g_jc:
        diffs.append(Diff(f"{base}.alignment", str(g_jc), str(a_jc)))


def _compare_run_fonts(ar, gr, rbase, diffs):
    """Compare w:rPr/w:rFonts attributes."""
    a_fonts = _get_run_fonts(ar)
    g_fonts = _get_run_fonts(gr)
    for attr in ("ascii", "hAnsi", "eastAsia"):
        a_val = a_fonts.get(attr)
        g_val = g_fonts.get(attr)
        if a_val != g_val:
            diffs.append(Diff(f"{rbase}.font.{attr}", str(g_val), str(a_val)))


def _compare_run_sizes(ar, gr, rbase, diffs):
    """Compare w:rPr/w:sz and w:rPr/w:szCs."""
    a_sz = _get_run_sizes(ar)
    g_sz = _get_run_sizes(gr)
    for attr in ("sz", "szCs"):
        a_val = a_sz.get(attr)
        g_val = g_sz.get(attr)
        if a_val != g_val:
            diffs.append(Diff(f"{rbase}.{attr}", str(g_val), str(a_val)))


def _compare_tables(actual, golden, diffs):
    a_tables = list(actual.tables)
    g_tables = list(golden.tables)
    if len(a_tables) != len(g_tables):
        diffs.append(Diff("tables.count", str(len(g_tables)), str(len(a_tables))))
        return

    for i, (at, gt) in enumerate(zip(a_tables, g_tables)):
        tbase = f"tables[{i}]"

        # Row height
        a_rows = list(at.rows)
        g_rows = list(gt.rows)
        nr = min(len(a_rows), len(g_rows))
        for j in range(nr):
            a_height = _get_row_height(a_rows[j])
            g_height = _get_row_height(g_rows[j])
            if a_height != g_height:
                diffs.append(Diff(f"{tbase}.rows[{j}].height",
                                  str(g_height), str(a_height)))

        # Borders
        a_borders = _get_table_border_width(at)
        g_borders = _get_table_border_width(gt)
        if a_borders != g_borders:
            diffs.append(Diff(f"{tbase}.border.width",
                              str(g_borders), str(a_borders)))


def _get_outline_level(para):
    pPr = para._element.find(_ns("pPr"))
    if pPr is not None:
        outline = pPr.find(_ns("outlineLvl"))
        if outline is not None:
            return int(outline.get(_ns("val")))
    return None


def _get_alignment(para):
    pPr = para._element.find(_ns("pPr"))
    if pPr is not None:
        jc = pPr.find(_ns("jc"))
        if jc is not None:
            return jc.get(_ns("val"))
    if para.alignment is not None:
        return str(para.alignment)
    return None


def _get_indent(para) -> dict:
    """Get w:ind attributes as dict."""
    result = {}
    pPr = para._element.find(_ns("pPr"))
    if pPr is not None:
        ind = pPr.find(_ns("ind"))
        if ind is not None:
            for attr in ("left", "right", "firstLine", "hanging"):
                val = ind.get(_ns(attr))
                if val is not None:
                    result[attr] = val
    return result


def _get_spacing(para) -> dict:
    """Get w:spacing attributes as dict."""
    result = {}
    pPr = para._element.find(_ns("pPr"))
    if pPr is not None:
        spacing = pPr.find(_ns("spacing"))
        if spacing is not None:
            for attr in ("beforeLines", "afterLines", "line"):
                val = spacing.get(_ns(attr))
                if val is not None:
                    result[attr] = val
    return result


def _get_run_fonts(run) -> dict:
    """Get w:rFonts attributes from run XML."""
    result = {}
    rPr = run._element.find(_ns("rPr"))
    if rPr is not None:
        rFonts = rPr.find(_ns("rFonts"))
        if rFonts is not None:
            for attr in ("ascii", "hAnsi", "eastAsia"):
                val = rFonts.get(_ns(attr))
                if val is not None:
                    result[attr] = val
    return result


def _get_run_sizes(run) -> dict:
    """Get w:sz and w:szCs values from run XML."""
    result = {}
    rPr = run._element.find(_ns("rPr"))
    if rPr is not None:
        sz = rPr.find(_ns("sz"))
        if sz is not None:
            val = sz.get(_ns("val"))
            if val is not None:
                result["sz"] = val
        szCs = rPr.find(_ns("szCs"))
        if szCs is not None:
            val = szCs.get(_ns("val"))
            if val is not None:
                result["szCs"] = val
    return result


def _get_row_height(row) -> dict:
    """Get w:trHeight val and hRule."""
    result = {}
    tr = row._element
    trPr = tr.find(_ns("trPr"))
    if trPr is not None:
        trHeight = trPr.find(_ns("trHeight"))
        if trHeight is not None:
            val = trHeight.get(_ns("val"))
            rule = trHeight.get(_ns("hRule"))
            if val is not None:
                result["val"] = val
            if rule is not None:
                result["hRule"] = rule
    return result


def _get_table_border_width(table) -> str | None:
    """Get the first border sz value (uniform borders assumed)."""
    tblPr = table._element.find(_ns("tblPr"))
    if tblPr is not None:
        borders = tblPr.find(_ns("tblBorders"))
        if borders is not None:
            top = borders.find(_ns("top"))
            if top is not None:
                return top.get(_ns("sz"))
    return None
