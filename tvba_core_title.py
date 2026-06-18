"""Title detection and formatting.

Corresponds to VBA FormatModule.bas:
  - AutoDetectAndFormatNumericTitles (line 783-826)
  - IdentifyContentTitleLevel (line 840-893)
  - IdentifyContentTitleLevelFromNumber (line 901-949)
  - ApplyContentTitleStyle (line 951+)
  - NormalizeNumberString (line 829-837)
"""
import copy
from docx.oxml import OxmlElement
import re

from lxml import etree

from tvba_core_oox import (
    set_far_east_font,
    set_ascii_font,
    set_run_font_size,
    set_outline_level,
    get_effective_outline_level,
    apply_indent_chars,
    apply_style_indent_chars,
    apply_paragraph_spacing,
    set_snap_to_grid,
    set_auto_space_de,
    format_all_runs_in_paragraph,
    clear_paragraph_formatting,
)
from tvba_core_normalize import (
    apply_brackets,
    sync_number_font_with_body,
)
from tvba_utils import size_label_to_points, clean_para_text

# Matches numeric title prefix: digits (and fullwidth digits) with dots/fullwidth dots,
# optionally ending with a trailing dot, followed by optional space/tab and title text.
_TITLE_RE = re.compile(r"^([\d０-９]+(?:[\.．]\d+){0,3}\.?)([ \t]*)(.+)$")

# Legacy list-item markers such as 1), 1）, （1）, and a. are body list
# markers, not content title levels.
_LIST_ITEM_RE = re.compile(r"(?!)")
_LETTER_ITEM_RE = re.compile(r"(?!)")

_QUANTITY_TITLE_STARTS = (
    "\u4e2a", "\u5ea7", "\u9879", "\u6761", "\u6b21", "\u5e74", "\u6708", "\u65e5",
    "\u516c\u91cc", "\u7c73", "\u5428", "\u4eba", "\u6237", "\u5904", "\u5bb6",
    "km", "m", "%",
)
_SENTENCE_PUNCTUATION_RE = re.compile(r"[\uff0c\uff1b\uff1a\uff01\uff1f\u3002,;:!?]")


def normalize_number_string(s: str) -> str:
    """Normalize number string: fullwidth digits/dots to halfwidth, strip trailing dot, trim."""
    s = s.strip()
    # Fullwidth digits ０-９ (U+FF10-U+FF19) to halfwidth 0-9
    for i in range(10):
        s = s.replace(chr(0xFF10 + i), str(i))
    s = s.replace("．", ".")
    s = s.replace("。", ".")
    if s.endswith("."):
        s = s[:-1]
    return s


def identify_level_from_number(num_str: str) -> int:
    """Map a normalized number string to title level 1-4 (0 = not a title)."""
    if not num_str:
        return 0
    if not re.fullmatch(r"\d+(?:\.\d+){0,3}", num_str):
        return 0
    dot_count = num_str.count(".")
    if dot_count == 0:
        return 1
    if dot_count == 1 and num_str.endswith(".0"):
        return 1
    level = dot_count + 1
    if 1 <= level <= 4:
        return level
    return 0


def identify_numeric_title_level(text: str) -> int:
    """Identify title level from paragraph text. Returns 1-4 or 0.

    Detects Arabic dotted-number titles: 1, 1.1, 1.1.1, etc.
    """
    text = clean_para_text(text)
    if not text:
        return 0
    m = _TITLE_RE.match(text)
    if not m:
        return 0
    num_part = normalize_number_string(m.group(1))
    sep = m.group(2)
    title_text = m.group(3).strip()
    level = identify_level_from_number(num_part)
    if level == 0:
        return 0
    if _looks_like_list_marker_tail(title_text):
        return 0
    if "." in num_part and not sep:
        return 0
    if "." not in num_part and _looks_like_quantity_or_body_sentence(title_text):
        return 0
    return level


def _looks_like_list_marker_tail(title_text: str) -> bool:
    """Reject list-number tails after an Arabic digit, e.g. 1), 1）, 1、."""
    return bool(title_text) and title_text[0] in ")]}）】、."


def _looks_like_quantity_or_body_sentence(title_text: str) -> bool:
    """Reject body text like '21 个地市...' from plain-integer headings."""
    if not title_text:
        return True
    compact = title_text.lstrip()
    lower = compact.lower()
    if any(lower.startswith(prefix) for prefix in _QUANTITY_TITLE_STARTS):
        return True
    return len(compact) >= 40 and bool(_SENTENCE_PUNCTUATION_RE.search(compact))


def identify_chinese_title(text: str) -> bool:
    """Chinese-number headings are intentionally not auto-detected."""
    return False


def identify_list_item(text: str) -> int:
    """Deprecated list-marker title detection.

    Bracket/letter list items are intentionally not treated as titles.  Keep
    this function as a compatibility shim for callers/tests.
    """
    return 0


def _strip_list_prefix(text: str, level: int) -> str:
    """Remove the list-item prefix (e.g. '(1) ', 'a. ') and return remaining text.

    Returns empty string if prefix not found.
    """
    text = clean_para_text(text)
    if level == 4:
        m = _LIST_ITEM_RE.match(text)
        if m:
            return text[m.end():].strip()
    elif level == 5:
        m = _LETTER_ITEM_RE.match(text)
        if m:
            return text[m.end():].strip()
    return ""


def _find_split_positions(text: str) -> list[int]:
    """Find non-zero indices where new title patterns start within paragraph text.

    Used to detect compound paragraphs that contain multiple concatenated
    Arabic-numbered titles (e.g. "1 项目背景1.1 格式问题梳理").
    """
    positions = set()

    # 1. Dotted numbers (1.1, 1.1.1, etc.) at non-start, not preceded by digit
    for m in re.finditer(r'\d+(?:\.\d+){1,3}(?!\.\d)', text):
        pos = m.start()
        if pos == 0:
            continue
        if text[pos - 1].isdigit() or text[pos - 1] == ".":
            continue
        if _looks_like_reference_or_quantity_number(text, pos, m.end()):
            continue
        end = m.end()
        if end < len(text) and text[end] in ' \t.':
            positions.add(pos)
        elif end < len(text) and '一' <= text[end] <= '鿿':
            positions.add(pos)

    # 2. Plain numbers after clear title boundaries
    #    (but NOT after "." which would indicate a dotted number like 1.1).
    #    Level-1 headings may be written as "1 概述" or "1概述".
    for m in re.finditer(
        r'(?<=[。！？；;：:\s])(?:\d+(?=\s+[一-鿿])|[1-9](?=[一-鿿]))', text
    ):
        if _looks_like_plain_number_quantity(text, m.start(), m.end()):
            continue
        positions.add(m.start())

    # Sort and merge nearby positions (within 2 chars)
    sorted_pos = sorted(positions)
    merged = []
    for p in sorted_pos:
        if not merged or p - merged[-1] > 2:
            merged.append(p)
    return merged


def _looks_like_reference_or_quantity_number(text: str, start: int, end: int) -> bool:
    """Reject dotted numbers that are references or quantities, not titles."""
    left = text[max(0, start - 12):start]
    right = text[end:end + 12]
    nearby = text[max(0, start - 12):min(len(text), end + 12)]

    if re.search(r"(?:GB|GB/T|SY/T|DL/T|NB/T|ISO|API)\s*$", left, re.IGNORECASE):
        return True
    if re.search(r"(?:GB|GB/T|SY/T|DL/T|NB/T|ISO|API)\s*", nearby, re.IGNORECASE):
        return True
    if re.search(r"第\s*$", left):
        return True
    if re.match(r"\s*(?:的)?(?:节|章节|条|款|部分|规定|要求|相关要求|方法)", right):
        return True
    if re.match(r"\s*(?:公里|km|m|mm|cm|μm|um|%|MPa|kPa|h|小时|次|处|根|级)", right, re.IGNORECASE):
        return True
    return False


def _looks_like_plain_number_quantity(text: str, start: int, end: int) -> bool:
    """Reject plain numeric quantities like '2 级' and '3 处'."""
    right = text[end:end + 8]
    return bool(re.match(r"\s*(?:级|次|处|根|个|条|项|年|月|日|小时|h|m|mm|cm|km|%)", right, re.IGNORECASE))


def split_compound_paragraphs(doc, *, _paragraphs=None) -> list:
    """Split paragraphs that contain multiple concatenated titles.

    When a single paragraph contains multiple title patterns (e.g.
    "1 项目背景1.1 格式问题梳理"), the system can
    only detect the first title. This function pre-processes the document
    by splitting such compound paragraphs into separate paragraphs at
    title boundaries.

    Returns the updated paragraph list (python-docx paragraph objects).
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    paragraphs = _paragraphs if _paragraphs is not None else list(doc.paragraphs)
    if not paragraphs:
        return paragraphs

    from tvba_core_body import is_manual_list_item

    body = doc.element.body
    # Build a map from paragraph text to its XML element.
    # We iterate by text content rather than element to avoid stale references.
    text_to_elems: dict[str, list] = {}
    for para in paragraphs:
        if is_manual_list_item(para) or _has_numbering_hint(para):
            continue
        text = (para.text or "").strip()
        if text:
            text_to_elems.setdefault(text, []).append(para._element)

    splits_found = []
    for text, elems in text_to_elems.items():
        positions = _find_split_positions(text)
        if len(positions) < 1:
            continue
        for para_elem in elems:
            splits_found.append((para_elem, text, positions))

    if not splits_found:
        return paragraphs

    for para_elem, text, positions in splits_found:
        parent = para_elem.getparent()
        if parent is None:
            continue

        # Build split parts from positions
        parts = []
        prev = 0
        for pos in positions:
            if pos > prev:
                part_text = text[prev:pos].strip()
                if part_text:
                    parts.append(part_text)
            prev = pos
        if prev < len(text):
            part_text = text[prev:].strip()
            if part_text:
                parts.append(part_text)

        if len(parts) < 2:
            continue

        parent_children = list(parent)
        try:
            idx = parent_children.index(para_elem)
        except ValueError:
            continue

        # Copy paragraph properties from original
        orig_pPr = para_elem.find(f"{{{W}}}pPr")

        for i, part_text in enumerate(parts):
            new_para = OxmlElement("w:p")

            # Copy paragraph properties for every new paragraph
            if orig_pPr is not None:
                new_para.append(copy.deepcopy(orig_pPr))

            # Create run with basic font settings
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(f"{{{W}}}ascii", "Times New Roman")
            rFonts.set(f"{{{W}}}eastAsia", "宋体")
            rPr.append(rFonts)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = part_text
            r.append(t)
            new_para.append(r)

            parent.insert(idx + 1 + i, new_para)

        parent.remove(para_elem)

    # Return fresh paragraph list since we modified the document
    return list(doc.paragraphs)


def apply_title_style(paragraph, level: int, level_settings, body_settings) -> None:
    """Apply title formatting to a paragraph."""
    # Clear existing run-level formatting before applying new styling
    clear_paragraph_formatting(paragraph)

    # Set outline level (0-indexed: level 1 -> 0)
    set_outline_level(paragraph, level - 1)

    # Font on each run (including nested runs inside fields/hyperlinks)
    format_all_runs_in_paragraph(
        paragraph,
        ascii_font="Times New Roman",
        eastasia_font=level_settings.font,
        size_pt=size_label_to_points(level_settings.size),
        bold=level_settings.bold,
    )

    # Alignment
    _ALIGNMENT_MAP = {"左对齐": 0, "居中": 1, "右对齐": 2, "两端对齐": 3}
    paragraph.alignment = _ALIGNMENT_MAP.get(level_settings.alignment, 0)

    # Spacing (before/after in lines, line spacing) via unified helper
    apply_paragraph_spacing(
        paragraph.paragraph_format,
        before_lines=level_settings.before_lines,
        after_lines=level_settings.after_lines,
        line_spacing=level_settings.line_spacing,
    )
    apply_indent_chars(
        paragraph.paragraph_format,
        left_chars=level_settings.left_indent_chars,
        right_chars=level_settings.right_indent_chars,
        special_kind=level_settings.special_indent,
        special_chars=level_settings.special_indent_chars,
    )
    apply_style_indent_chars(
        paragraph.style,
        left_chars=level_settings.left_indent_chars,
        right_chars=level_settings.right_indent_chars,
        special_kind=level_settings.special_indent,
        special_chars=level_settings.special_indent_chars,
    )

    # Grid alignment for level 1 titles (大鹏模板要求)
    if level == 1:
        set_snap_to_grid(paragraph, True)
        set_auto_space_de(paragraph, True)

    # Normalize brackets and sync number font (no period for titles).
    # Content modifications run only when body settings explicitly allow.
    if body_settings.modify_content or level_settings.normalize_brackets:
        apply_brackets(paragraph, paragraph.text)
    sync_number_font_with_body(paragraph)


def _has_numbering_hint(para, _style_by_id=None) -> bool:
    """Return True when OOXML suggests this paragraph may be a numbered list.

    Word COM list checks are expensive. This cheap filter avoids calling COM for
    every ordinary body paragraph in large documents.
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    pPr = para._element.find(f"{{{W}}}pPr")
    if pPr is not None:
        if pPr.find(f"{{{W}}}numPr") is not None:
            return True
        pStyle = pPr.find(f"{{{W}}}pStyle")
        style_id = pStyle.get(f"{{{W}}}val") if pStyle is not None else None
    else:
        style_id = None

    if not style_id or _style_by_id is None:
        return False

    style = _style_by_id.get(style_id)
    visited = set()
    while style is not None:
        sid = id(style)
        if sid in visited:
            break
        visited.add(sid)

        style_pPr = style.element.find(f"{{{W}}}pPr")
        if style_pPr is not None and style_pPr.find(f"{{{W}}}numPr") is not None:
            return True

        based_on = style.element.find(f"{{{W}}}basedOn")
        if based_on is None:
            break
        based_id = based_on.get(f"{{{W}}}val")
        if not based_id:
            break
        style = _style_by_id.get(based_id)

    return False


def auto_detect_and_format(doc, settings, list_resolver=None, *, _paragraphs=None, _outline_cache=None, _style_by_id=None, _toc_style_ids=None) -> None:
    """Auto-detect numeric titles and apply title formatting."""
    from tvba_core_body import is_manual_list_item
    from tvba_core_toc import is_toc_paragraph

    paragraphs = _paragraphs if _paragraphs is not None else doc.paragraphs
    for para in paragraphs:
        if is_toc_paragraph(para, _toc_style_ids=_toc_style_ids):
            continue

        if is_manual_list_item(para):
            continue

        # Skip paragraphs that already have an outline level set by the user
        # (either directly on the paragraph or inherited from its style)
        existing_outline = get_effective_outline_level(para, _cache=_outline_cache, _style_by_id=_style_by_id)
        if existing_outline is not None:
            continue

        text = clean_para_text(para.text)
        if not text:
            continue

        # Priority 1: Arabic numeric title text detection
        level = identify_numeric_title_level(text)

        # Priority 1b: Reserved compatibility hook.  Chinese-number titles
        # (一、二、三...) are intentionally not auto-detected.
        if level == 0 and identify_chinese_title(text):
            level = 1

        # Priority 2: Multi-level list resolver (COM or docx) as fallback.
        # Only trust list resolver when it can prove the paragraph is a
        # heading (via list text for COM) to avoid treating body list
        # items (e.g. "1）第一项") as titles.
        if level == 0 and list_resolver is not None and _has_numbering_hint(para, _style_by_id):
            list_level = list_resolver.get_list_level(para)
            if list_level is not None and 1 <= list_level <= 4:
                # COM resolver: verify via rendered list text
                if hasattr(list_resolver, 'get_list_text'):
                    list_text = list_resolver.get_list_text(para)
                    if list_text:
                        normalized = normalize_number_string(list_text)
                        derived_level = identify_level_from_number(normalized)
                        if derived_level > 0:
                            level = derived_level
                # Docx resolver: cannot reliably distinguish heading lists
                # from body lists, so skip to avoid false positives.
                # Text already matched above if it looks like a title.

                # Without rendered list text, do not infer title levels from
                # body list numbering.  This avoids turning "1）" / "（1）"
                # style body items into level-4 titles.

        if 1 <= level <= 5:
            apply_title_style(
                para,
                level,
                settings.titles[level - 1],
                settings.body,
            )
