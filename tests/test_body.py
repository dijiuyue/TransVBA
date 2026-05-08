import tempfile
from pathlib import Path
import pytest
from docx import Document

from tvba_core_body import apply_normal_style, apply_paragraph
from tvba_settings import BodySettings
from tvba_utils import size_label_to_points

NSMAP = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

class TestApplyNormalStyle:
    def test_sets_normal_style_font(self):
        doc = Document()
        body = BodySettings(font="黑体", size="小四", spacing=1.5)
        apply_normal_style(doc, body)
        normal = doc.styles["Normal"]
        assert normal.font.name == "Times New Roman"

    def test_sets_normal_style_size(self):
        doc = Document()
        body = BodySettings(size="小四")
        apply_normal_style(doc, body)
        normal = doc.styles["Normal"]
        assert normal.font.size.pt == pytest.approx(12.0, rel=1e-4)

class TestApplyParagraph:
    def test_applies_font_and_size(self):
        doc = Document()
        para = doc.add_paragraph("正文段落")
        body = BodySettings(font="宋体", size="小四", spacing=1.5)
        apply_paragraph(para, body)
        run = para.runs[0]
        assert run.font.name == "Times New Roman"

    def test_applies_alignment(self):
        doc = Document()
        para = doc.add_paragraph("正文段落")
        body = BodySettings(alignment="居中")
        apply_paragraph(para, body)
        # python-docx alignment: 0=left, 1=center, 2=right, 3=justify
        assert para.alignment == 1

    def test_applies_justified_alignment(self):
        doc = Document()
        para = doc.add_paragraph("正文段落")
        body = BodySettings(alignment="两端对齐")
        apply_paragraph(para, body)
        assert para.alignment == 3

    def test_applies_first_line_indent(self):
        doc = Document()
        para = doc.add_paragraph("正文段落")
        body = BodySettings(special_indent="首行缩进", special_indent_cm=0.74)
        apply_paragraph(para, body)
        pPr = para._element.find(".//w:pPr", NSMAP)
        ind = pPr.find("w:ind", NSMAP)
        assert ind is not None
        assert ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}firstLine") is not None

    def test_applies_line_spacing(self):
        doc = Document()
        para = doc.add_paragraph("正文段落")
        body = BodySettings(spacing=2.0)
        apply_paragraph(para, body)
        pPr = para._element.find(".//w:pPr", NSMAP)
        spacing_el = pPr.find("w:spacing", NSMAP)
        assert spacing_el is not None
        assert spacing_el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}line") == "400"
