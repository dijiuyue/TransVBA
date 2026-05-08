"""Build a comprehensive test .docx with all element types."""
from docx import Document
import sys
from pathlib import Path


def build(path: Path) -> None:
    doc = Document()

    # TOC title
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章  绪论\t1")
    doc.add_paragraph("  1.1  背景\t2")

    # Titles
    doc.add_paragraph("1 引言")
    doc.add_paragraph("1.1 研究背景")
    doc.add_paragraph("1.1.1 详细背景")
    doc.add_paragraph("1.1.1.1 更详细")
    doc.add_paragraph("1.1.1.1.1 最详细")

    # Body
    doc.add_paragraph("这是一段正文，包含数字123和英文ABC。")
    doc.add_paragraph("第二段正文。")

    # Table
    doc.add_paragraph("表 1.1-1 示例表格")
    table = doc.add_table(rows=2, cols=2)
    for row in table.rows:
        for cell in row.cells:
            cell.text = "数据"

    # Figure
    doc.add_paragraph("图 1-1 示例图片")

    doc.save(path)


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("full_test.docx")
    build(out)
    print(f"Built: {out}")
