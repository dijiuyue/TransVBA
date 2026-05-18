import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from lxml import etree

from tvba_core_figure import (
    is_figure_caption_line,
    is_figure_caption_paragraph,
    apply_figure_caption,
    refresh_all,
)
from tvba_settings import FigureSettings

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_NS = {"w": W}
W_VAL = f"{{{W}}}val"


def _ppr(para):
    pPr = para._element.find("w:pPr", W_NS)
    if pPr is None:
        pPr = etree.SubElement(para._element, f"{{{W}}}pPr")
    return pPr


def _set_dirty_caption_paragraph_format(para):
    pPr = _ppr(para)
    jc = pPr.find("w:jc", W_NS)
    if jc is None:
        jc = etree.SubElement(pPr, f"{{{W}}}jc")
    jc.set(W_VAL, "both")
    outline = pPr.find("w:outlineLvl", W_NS)
    if outline is None:
        outline = etree.SubElement(pPr, f"{{{W}}}outlineLvl")
    outline.set(W_VAL, "6")
    ind = pPr.find("w:ind", W_NS)
    if ind is None:
        ind = etree.SubElement(pPr, f"{{{W}}}ind")
    ind.set(f"{{{W}}}firstLine", "482")


def _add_numbering_level(doc, *, num_id="901", abstract_id="901", ilvl="6", lvl_text="图%1.%2.%3-%7  "):
    numbering = doc.part.numbering_part._element
    abstract = etree.SubElement(numbering, f"{{{W}}}abstractNum")
    abstract.set(f"{{{W}}}abstractNumId", abstract_id)
    lvl = etree.SubElement(abstract, f"{{{W}}}lvl")
    lvl.set(f"{{{W}}}ilvl", ilvl)
    start = etree.SubElement(lvl, f"{{{W}}}start")
    start.set(W_VAL, "1")
    num_fmt = etree.SubElement(lvl, f"{{{W}}}numFmt")
    num_fmt.set(W_VAL, "decimal")
    lvl_text_elem = etree.SubElement(lvl, f"{{{W}}}lvlText")
    lvl_text_elem.set(W_VAL, lvl_text)

    num = etree.SubElement(numbering, f"{{{W}}}num")
    num.set(f"{{{W}}}numId", num_id)
    abstract_ref = etree.SubElement(num, f"{{{W}}}abstractNumId")
    abstract_ref.set(W_VAL, abstract_id)
    return num_id


def _set_paragraph_num_pr(para, *, num_id="901", ilvl="6"):
    pPr = _ppr(para)
    num_pr = pPr.find("w:numPr", W_NS)
    if num_pr is None:
        num_pr = etree.SubElement(pPr, f"{{{W}}}numPr")
    ilvl_elem = num_pr.find("w:ilvl", W_NS)
    if ilvl_elem is None:
        ilvl_elem = etree.SubElement(num_pr, f"{{{W}}}ilvl")
    ilvl_elem.set(W_VAL, ilvl)
    num_id_elem = num_pr.find("w:numId", W_NS)
    if num_id_elem is None:
        num_id_elem = etree.SubElement(num_pr, f"{{{W}}}numId")
    num_id_elem.set(W_VAL, num_id)


class TestIsFigureCaptionLine:
    def test_starts_with_tu(self):
        assert is_figure_caption_line("图 1-1 示例图片") is True

    def test_tu_no_space(self):
        assert is_figure_caption_line("图1.1-1 示例图片") is True

    def test_tu_simple_number(self):
        assert is_figure_caption_line("图 1-1 示例图片") is True

    def test_starts_with_figure(self):
        assert is_figure_caption_line("Figure 1-1 Example") is True

    def test_figure_with_dot_number(self):
        assert is_figure_caption_line("Figure 1.1-1 Example") is True

    def test_no_prefix(self):
        assert is_figure_caption_line("示例图片") is False

    def test_case_insensitive(self):
        assert is_figure_caption_line("FIGURE 1-1 Example") is True

    def test_simple_number_matches(self):
        assert is_figure_caption_line("图 1 示例图片") is True

    def test_no_caption_text_fails(self):
        assert is_figure_caption_line("图 1-1") is False

    def test_only_prefix_fails(self):
        assert is_figure_caption_line("图 ") is False

    def test_fig_prefix_fails(self):
        assert is_figure_caption_line("Fig 1 Example") is False

    def test_figure_without_number_pattern_fails(self):
        assert is_figure_caption_line("Figure Example") is False

    def test_fullwidth_space_separator(self):
        """Full-width space (U+3000) between number and text should match."""
        assert is_figure_caption_line("图2.2.11-1　管道经过南朗遗物点位置示意图") is True

    def test_nonbreaking_space_separator(self):
        """Non-breaking space (U+00A0) between number and text should match."""
        assert is_figure_caption_line("图2.2.11-1\xa0管道经过南朗遗物点位置示意图") is True

    def test_multi_dot_number_with_double_space(self):
        """Real captions like 图1.8.3-1  should be recognized."""
        assert is_figure_caption_line("图1.8.3-1  工艺系统示意图") is True

    @pytest.mark.parametrize(
        "caption",
        [
            "图1-1  测试图片",
            "图 1-1  测试图片",
            "图1.8-1  测试图片",
            "图 1.8-1  测试图片",
            "图1.8.3-1  工艺系统示意图",
            "图 1.8.3-1  工艺系统示意图",
            "图1.10.1-1  双水首站与已建工程界面划分示意图",
            "图 1.10.1-1  双水首站与已建工程界面划分示意图",
            "图１．１０．１－１  双水首站与已建工程界面划分示意图",
            "图1.10.1—1  双水首站与已建工程界面划分示意图",
        ],
    )
    def test_figure_caption_number_category_matches(self, caption):
        """图 + optional space + dot-number + hyphen suffix captions are one category."""
        assert is_figure_caption_line(caption) is True


class TestIsFigureCaptionParagraph:
    def test_direct_numbering_with_figure_label_is_caption(self):
        doc = Document()
        _add_numbering_level(doc)
        para = doc.add_paragraph("作业带布置示意图（二）")
        _set_paragraph_num_pr(para)

        assert is_figure_caption_paragraph(para, doc) is True

    def test_direct_numbering_with_table_label_is_not_figure_caption(self):
        doc = Document()
        _add_numbering_level(doc, lvl_text="表%1.%2.%3-%6  ")
        para = doc.add_paragraph("沿线经过环境敏感点统计表")
        _set_paragraph_num_pr(para)

        assert is_figure_caption_paragraph(para, doc) is False

    def test_style_name_with_figure_number_is_caption(self):
        doc = Document()
        style = doc.styles.add_style("图1.1.1-1", WD_STYLE_TYPE.PARAGRAPH)
        para = doc.add_paragraph("拟建珠中江干线示意图")
        para.style = style

        assert is_figure_caption_paragraph(para, doc) is True


class TestApplyFigureCaption:
    def test_applies_font_and_bold(self):
        doc = Document()
        para = doc.add_paragraph("图 1-1 示例")
        settings = FigureSettings(title_font="黑体", title_size="小四", title_bold=True)
        apply_figure_caption(para, settings)
        run = para.runs[0]
        assert run.font.bold is True

    def test_default_caption_is_not_bold_and_clears_cjk_bold(self):
        doc = Document()
        para = doc.add_paragraph("图 1-1 示例")
        rPr = para.runs[0]._element.get_or_add_rPr()
        etree.SubElement(rPr, f"{{{W}}}b")
        etree.SubElement(rPr, f"{{{W}}}bCs")

        apply_figure_caption(para, FigureSettings())

        rPr = para.runs[0]._element.rPr
        assert rPr.find("w:b", W_NS).get(W_VAL) == "0"
        assert rPr.find("w:bCs", W_NS).get(W_VAL) == "0"

    def test_auto_numbered_caption_clears_numbering_bold(self):
        doc = Document()
        _add_numbering_level(doc)
        numbering = doc.part.numbering_part._element
        lvl = numbering.find(".//w:abstractNum[@w:abstractNumId='901']/w:lvl[@w:ilvl='6']", W_NS)
        rPr = etree.SubElement(lvl, f"{{{W}}}rPr")
        etree.SubElement(rPr, f"{{{W}}}b")
        etree.SubElement(rPr, f"{{{W}}}bCs")
        para = doc.add_paragraph("作业带布置示意图（二）")
        _set_paragraph_num_pr(para)

        apply_figure_caption(para, FigureSettings(), doc)

        assert rPr.find("w:b", W_NS).get(W_VAL) == "0"
        assert rPr.find("w:bCs", W_NS).get(W_VAL) == "0"

    def test_multi_dot_caption_is_centered(self):
        doc = Document()
        para = doc.add_paragraph("图1.8.3-1  工艺系统示意图")
        _set_dirty_caption_paragraph_format(para)
        settings = FigureSettings(title_alignment="居中")

        apply_figure_caption(para, settings)

        pPr = para._element.find("w:pPr", W_NS)
        jc = pPr.find("w:jc", W_NS)
        assert jc.get(W_VAL) == "center"

    @pytest.mark.parametrize(
        "caption",
        [
            "图1-1  测试图片",
            "图 1-1  测试图片",
            "图1.8-1  测试图片",
            "图 1.8-1  测试图片",
            "图1.8.3-1  工艺系统示意图",
            "图 1.8.3-1  工艺系统示意图",
            "图1.10.1-1  双水首站与已建工程界面划分示意图",
            "图 1.10.1-1  双水首站与已建工程界面划分示意图",
            "图１．１０．１－１  双水首站与已建工程界面划分示意图",
            "图1.10.1—1  双水首站与已建工程界面划分示意图",
        ],
    )
    def test_figure_caption_category_clears_dirty_alignment_outline_and_indent(self, caption):
        doc = Document()
        para = doc.add_paragraph(caption)
        _set_dirty_caption_paragraph_format(para)
        settings = FigureSettings()

        apply_figure_caption(para, settings)

        pPr = para._element.find("w:pPr", W_NS)
        jc = pPr.find("w:jc", W_NS)
        outline = pPr.find("w:outlineLvl", W_NS)
        ind = pPr.find("w:ind", W_NS)
        assert jc.get(W_VAL) == "center"
        assert outline is None
        if ind is not None:
            assert ind.get(f"{{{W}}}firstLine") is None
            assert ind.get(f"{{{W}}}hanging") is None


class TestRefreshAll:
    def test_finds_and_formats_captions(self):
        doc = Document()
        doc.add_paragraph("图 1-1 测试图片")
        doc.add_paragraph("正文段落")
        doc.add_paragraph("图 2-1 另一个图片")
        settings = FigureSettings()
        refresh_all(doc, settings)
        # Caption should not be bold by default
        assert doc.paragraphs[0].runs[0].font.bold is not True
        # Body paragraph should not be formatted
        assert doc.paragraphs[1].runs[0].font.bold is not True

    def test_formats_multi_dot_caption(self):
        doc = Document()
        para = doc.add_paragraph("图1.8.3-1  工艺系统示意图")
        para.alignment = 3
        settings = FigureSettings(title_alignment="居中")

        refresh_all(doc, settings)

        pPr = para._element.find(".//w:pPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        jc = pPr.find("w:jc", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        assert jc.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "center"

    def test_formats_auto_numbered_caption_without_literal_figure_prefix(self):
        doc = Document()
        _add_numbering_level(doc)
        para = doc.add_paragraph("阴极保护站位置示意图")
        _set_paragraph_num_pr(para)
        settings = FigureSettings(title_alignment="居中")

        refresh_all(doc, settings)

        pPr = para._element.find(".//w:pPr", W_NS)
        jc = pPr.find("w:jc", W_NS)
        assert jc.get(W_VAL) == "center"
