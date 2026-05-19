import tempfile
from pathlib import Path
import pytest
from docx import Document

from tvba_controller import TvbaController, ValidationResult, ApplyResult
from tvba_persistence import SettingsRepository
from tvba_settings import FormatSettings, BodySettings, TitleLevelSettings
from tvba_templates import TemplateManager


class FakeDocumentApplier:
    def __init__(self):
        self.calls = []

    def __call__(self, docx_path, settings, **kwargs):
        self.calls.append((docx_path, settings, kwargs))
        return docx_path


class TestTvbaController:
    def test_settings_property_returns_format_settings(self):
        repo = SettingsRepository()
        applier = FakeDocumentApplier()
        ctrl = TvbaController(repo, applier)
        assert isinstance(ctrl.settings, FormatSettings)

    def test_open_file_sets_opened_file(self):
        repo = SettingsRepository()
        applier = FakeDocumentApplier()
        ctrl = TvbaController(repo, applier)
        ctrl.open_file(Path("C:\\test.docx"))
        assert ctrl.opened_file == Path("C:\\test.docx")

    def test_update_setting_changes_body_font(self):
        repo = SettingsRepository()
        applier = FakeDocumentApplier()
        ctrl = TvbaController(repo, applier)
        result = ctrl.update_setting("body.font", "黑体")
        assert result.valid is True
        assert ctrl.settings.body.font == "黑体"

    def test_update_setting_invalid_path(self):
        repo = SettingsRepository()
        applier = FakeDocumentApplier()
        ctrl = TvbaController(repo, applier)
        result = ctrl.update_setting("invalid.path", "value")
        assert result.valid is False

    def test_apply_calls_applier(self):
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            doc.add_paragraph("test")
            doc.save(path)

            repo = SettingsRepository()
            applier = FakeDocumentApplier()
            ctrl = TvbaController(repo, applier)
            ctrl.open_file(path)
            result = ctrl.apply(save_settings=False)
            assert result.success is True
            assert len(applier.calls) == 1

    def test_apply_includes_elapsed_time(self):
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "test.docx"
            doc = Document()
            doc.add_paragraph("test")
            doc.save(path)

            repo = SettingsRepository()
            applier = FakeDocumentApplier()
            ctrl = TvbaController(repo, applier)
            ctrl.open_file(path)
            result = ctrl.apply(save_settings=False)
            assert result.success is True
            assert hasattr(result, "elapsed_ms")
            assert result.elapsed_ms >= 0

    def test_reset_to_template_defaults(self):
        repo = SettingsRepository()
        applier = FakeDocumentApplier()
        ctrl = TvbaController(repo, applier)
        ctrl.update_setting("body.font", "黑体")
        ctrl.reset_to_template_defaults()
        assert ctrl.settings.body.font == "宋体"

    def test_load_saved_settings_ignores_other_template(self):
        with tempfile.TemporaryDirectory() as td:
            repo = SettingsRepository(Path(td) / "settings.json")
            repo.save(FormatSettings(template_name="general_spec"))
            applier = FakeDocumentApplier()
            dapeng = TemplateManager.load_template("dapeng_internal")
            ctrl = TvbaController(repo, applier, dapeng)

            ctrl.load_saved_settings()

            assert ctrl.settings.template_name == "dapeng_internal"
            assert [title.bold for title in ctrl.settings.titles] == [True, False, False, False, False]

    def test_load_saved_settings_migrates_legacy_dapeng_title_bold(self):
        with tempfile.TemporaryDirectory() as td:
            repo = SettingsRepository(Path(td) / "settings.json")
            legacy_titles = tuple(
                TitleLevelSettings(bold=bold)
                for bold in (True, True, True, False, False)
            )
            repo.save(FormatSettings(template_name="dapeng_internal", titles=legacy_titles))
            applier = FakeDocumentApplier()
            dapeng = TemplateManager.load_template("dapeng_internal")
            ctrl = TvbaController(repo, applier, dapeng)

            ctrl.load_saved_settings()

            assert [title.bold for title in ctrl.settings.titles] == [True, False, False, False, False]

    def test_load_saved_settings_migrates_legacy_dapeng_title_indent(self):
        with tempfile.TemporaryDirectory() as td:
            repo = SettingsRepository(Path(td) / "settings.json")
            legacy_titles = tuple(
                TitleLevelSettings(special_indent="首行缩进", special_indent_chars=2.0)
                for _ in range(5)
            )
            repo.save(FormatSettings(template_name="dapeng_internal", titles=legacy_titles))
            applier = FakeDocumentApplier()
            dapeng = TemplateManager.load_template("dapeng_internal")
            ctrl = TvbaController(repo, applier, dapeng)

            ctrl.load_saved_settings()

            assert all(title.special_indent == "无" for title in ctrl.settings.titles)
            assert all(title.special_indent_chars == 0.0 for title in ctrl.settings.titles)

    def test_switch_template_updates_reset_defaults(self):
        with tempfile.TemporaryDirectory() as td:
            repo = SettingsRepository(Path(td) / "settings.json")
            applier = FakeDocumentApplier()
            ctrl = TvbaController(repo, applier)

            ctrl.switch_template("general_spec")
            ctrl.update_setting("titles.3.left_indent_chars", 0.0)
            ctrl.reset_to_template_defaults()

            assert ctrl.settings.titles[3].left_indent_chars == 2.0

    def test_switch_template_loads_only_that_template_saved_settings(self):
        with tempfile.TemporaryDirectory() as td:
            repo = SettingsRepository(Path(td) / "settings.json")
            repo.save(FormatSettings(template_name="general_spec", body=BodySettings(font="黑体")))
            repo.save(FormatSettings(template_name="dapeng_internal", body=BodySettings(font="楷体")))
            applier = FakeDocumentApplier()
            ctrl = TvbaController(repo, applier)

            ctrl.switch_template("general_spec")
            assert ctrl.settings.body.font == "黑体"

            ctrl.switch_template("dapeng_internal")
            assert ctrl.settings.body.font == "楷体"

    def test_reset_clears_only_current_template_saved_settings(self):
        with tempfile.TemporaryDirectory() as td:
            repo = SettingsRepository(Path(td) / "settings.json")
            repo.save(FormatSettings(template_name="general_spec", body=BodySettings(font="黑体")))
            repo.save(FormatSettings(template_name="dapeng_internal", body=BodySettings(font="楷体")))
            applier = FakeDocumentApplier()
            ctrl = TvbaController(repo, applier)

            ctrl.switch_template("dapeng_internal")
            ctrl.reset_to_template_defaults()
            ctrl.clear_saved_settings()

            assert repo.load_for_template("dapeng_internal") == FormatSettings()
            assert repo.load_for_template("general_spec").body.font == "黑体"

    def test_update_setting_table(self):
        repo = SettingsRepository()
        applier = FakeDocumentApplier()
        ctrl = TvbaController(repo, applier)
        result = ctrl.update_setting("table.title_font", "楷体")
        assert result.valid is True
        assert ctrl.settings.table.title_font == "楷体"

    def test_update_setting_figure(self):
        repo = SettingsRepository()
        applier = FakeDocumentApplier()
        ctrl = TvbaController(repo, applier)
        result = ctrl.update_setting("figure.title_bold", False)
        assert result.valid is True
        assert ctrl.settings.figure.title_bold is False

    def test_update_setting_auto_detect_numeric_titles(self):
        repo = SettingsRepository()
        applier = FakeDocumentApplier()
        ctrl = TvbaController(repo, applier)
        result = ctrl.update_setting("auto_detect_numeric_titles", False)
        assert result.valid is True
        assert ctrl.settings.auto_detect_numeric_titles is False

    def test_update_setting_auto_detect_include_list_paragraphs(self):
        repo = SettingsRepository()
        applier = FakeDocumentApplier()
        ctrl = TvbaController(repo, applier)
        result = ctrl.update_setting("auto_detect_include_list_paragraphs", False)
        assert result.valid is True
        assert ctrl.settings.auto_detect_include_list_paragraphs is False

    def test_update_setting_remember_settings(self):
        repo = SettingsRepository()
        applier = FakeDocumentApplier()
        ctrl = TvbaController(repo, applier)
        result = ctrl.update_setting("remember_settings", False)
        assert result.valid is True
        assert ctrl.settings.remember_settings is False
