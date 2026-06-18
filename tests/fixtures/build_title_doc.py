"""Build a test .docx with numeric titles at levels 1-4."""
from docx import Document
import sys
from pathlib import Path


def build(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("1 一级标题")
    doc.add_paragraph("1.1 二级标题")
    doc.add_paragraph("1.1.1 三级标题")
    doc.add_paragraph("1.1.1.1 四级标题")
    doc.add_paragraph("这是一段正文，在标题之后。")
    doc.add_paragraph("1.0 特殊一级标题")
    doc.save(path)


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("title_test.docx")
    build(out)
    print(f"Built: {out}")
