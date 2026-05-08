"""Table + table caption formatting.

Corresponds to VBA FormatModule.bas:
  - RefreshTableFormat
  - SetTableTitle
  - FindTableCaptionRange
  - FindCaptionInShapes
  - IsTableCaptionLine
"""
from tvba_core_oox import (
    set_far_east_font,
    set_ascii_font,
    set_table_layout_window,
    set_table_layout_content,
    set_table_borders,
    set_row_height_at_least,
    apply_indent_chars,
    set_before_after_lines,
)
from tvba_utils import clean_para_text, size_label_to_points, cm_to_points
from docx.shared import Pt


def is_table_caption_line(text: str) -> bool:
    """Check if text is a table caption."""
    text = clean_para_text(text).lower()
    return text.startswith("表 ") or text.startswith("table ")


def find_table_caption(table, doc, max_up_paragraphs: int = 10):
    """Find the caption paragraph preceding a table."""
    # Find table index in document
    table_index = None
    for i, t in enumerate(doc.tables):
        if t._element is table._element:
            table_index = i
            break

    if table_index is None:
        return None

    # Count paragraphs before this table
    paragraphs_before = 0
    for element in doc.element.body:
        if element is table._element:
            break
        if element.tag.endswith("}p"):
            paragraphs_before += 1

    # Search backwards up to max_up_paragraphs
    for i in range(1, max_up_paragraphs + 1):
        idx = paragraphs_before - i
        if idx < 0:
            break
        para = doc.paragraphs[idx]
        if is_table_caption_line(para.text):
            return para

    # TODO: Shape/TextFrame search (COM fallback if needed)
    return None


def apply_table_caption(para, settings) -> None:
    """Apply formatting to a table caption paragraph."""
    for run in para.runs:
        set_ascii_font(run, "Times New Roman")
        set_far_east_font(run, settings.title_font)
        run.font.size = Pt(size_label_to_points(settings.title_size))
        run.font.bold = settings.title_bold

    para.alignment = 1  # Center alignment for captions

    set_before_after_lines(
        para.paragraph_format,
        before_lines=0.0,
        after_lines=0.0,
    )

    # Line spacing
    pPr = para._element.find(".//w:pPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
    if pPr is not None:
        from lxml import etree
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        spacing = pPr.find(f"{{{W}}}spacing")
        if spacing is None:
            spacing = etree.SubElement(pPr, f"{{{W}}}spacing")
        spacing.set(f"{{{W}}}line", str(int(settings.title_spacing * 240)))
        spacing.set(f"{{{W}}}lineRule", "auto")


def apply_table_body(table, settings) -> None:
    """Apply formatting to table body."""
    # Auto fit
    if settings.auto_fit_window:
        set_table_layout_window(table)
    else:
        set_table_layout_content(table)

    # Borders
    set_table_borders(table, line_width_pt=settings.line_width_pt)

    # Row height
    for row in table.rows:
        set_row_height_at_least(row, settings.row_height_cm)

    # Cell font
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_ascii_font(run, "Times New Roman")
                    set_far_east_font(run, settings.body_font)
                    run.font.size = Pt(size_label_to_points(settings.body_size))
                # Line spacing
                pPr = para._element.find(".//w:pPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                if pPr is not None:
                    from lxml import etree
                    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    spacing = pPr.find(f"{{{W}}}spacing")
                    if spacing is None:
                        spacing = etree.SubElement(pPr, f"{{{W}}}spacing")
                    spacing.set(f"{{{W}}}line", str(int(settings.spacing * 240)))
                    spacing.set(f"{{{W}}}lineRule", "auto")


def refresh_all(doc, settings) -> None:
    """Refresh all tables and their captions."""
    for table in doc.tables:
        caption = find_table_caption(table, doc)
        if caption is not None:
            apply_table_caption(caption, settings)
        apply_table_body(table, settings)
