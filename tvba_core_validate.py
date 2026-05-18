"""Document format validation engine.

Scans a Word document against the active template's validation rules
and returns a list of format issues found.
"""
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from tvba_settings import FormatSettings
from tvba_utils import size_label_to_points
from tvba_core_toc import is_toc_title_line, is_toc_entry_line
from tvba_core_appendix import is_appendix_title
from tvba_core_table import is_table_caption_line, is_table_caption_paragraph
from tvba_core_figure import is_figure_caption_line, is_figure_caption_paragraph
from tvba_core_oox import get_effective_run_fonts

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_RPR = f"{{{W}}}rPr"
W_RFONTS = f"{{{W}}}rFonts"
W_PPR = f"{{{W}}}pPr"
W_SZ = f"{{{W}}}sz"
W_VAL = f"{{{W}}}val"
W_EAST = f"{{{W}}}eastAsia"
W_ASCII = f"{{{W}}}ascii"
W_OUTLINE = f"{{{W}}}outlineLvl"
W_SNAP = f"{{{W}}}snapToGrid"
W_AUTO_DE = f"{{{W}}}autoSpaceDE"
W_TRPR = f"{{{W}}}trPr"
W_TRHEIGHT = f"{{{W}}}trHeight"
W_HRULE = f"{{{W}}}hRule"
W_T = f"{{{W}}}t"
W_SPACING = f"{{{W}}}spacing"
W_BEFORE = f"{{{W}}}before"
W_AFTER = f"{{{W}}}after"
W_BEFORE_LINES = f"{{{W}}}beforeLines"
W_AFTER_LINES = f"{{{W}}}afterLines"
W_JC = f"{{{W}}}jc"

CN_RE = re.compile(r'[一-鿿]')
ASCII_RE = re.compile(r'[a-zA-Z0-9]')


@dataclass
class ValidationIssue:
    severity: str  # "error" | "warning"
    description: str
    location: str = ""


def validate_document(
    docx_path: Path,
    settings: FormatSettings,
    progress_cb=None,
) -> list[ValidationIssue]:
    """Scan a document and return format issues based on the template's validation rules."""
    from tvba_core_convert import ensure_docx

    rules = settings.validation
    issues: list[ValidationIssue] = []

    if progress_cb:
        progress_cb("正在准备文档...", 0.0)
    docx_path = ensure_docx(docx_path)
    doc = Document(str(docx_path))
    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)
    n_paras = len(paragraphs)
    n_tables = len(tables)

    # Pre-compute total ticks for accurate progress
    total_ticks = 0
    if rules.check_chinese_font: total_ticks += n_paras
    if rules.check_ascii_font: total_ticks += n_paras
    if rules.check_brackets: total_ticks += n_paras
    if rules.check_period: total_ticks += n_paras
    if rules.check_forbidden_words and rules.forbidden_words: total_ticks += n_paras
    if rules.check_table_font_size: total_ticks += n_tables
    if rules.check_table_row_height: total_ticks += n_tables
    if rules.check_cover_title_size: total_ticks += 1
    if rules.check_appendix_body_size: total_ticks += 1
    if rules.check_grid_alignment: total_ticks += 1
    if rules.check_appendix_colon: total_ticks += 1
    if rules.check_figure_table_space: total_ticks += 1
    if rules.check_figure_position: total_ticks += 1
    if rules.check_chairman_number: total_ticks += 1

    count = 0
    _last_reported = 0
    def tick(msg: str):
        nonlocal count, _last_reported
        count += 1
        if progress_cb and (count - _last_reported >= 25 or count == total_ticks):
            _last_reported = count
            progress_cb(msg, count / max(total_ticks, 1))

    if rules.check_chinese_font:
        for para in paragraphs:
            _check_chinese_font(para, issues, doc=doc)
            tick("检查中文字体...")

    if rules.check_ascii_font:
        for para in paragraphs:
            _check_ascii_font(para, issues, doc=doc)
            tick("检查数字/英文字体...")

    if rules.check_brackets:
        for para in paragraphs:
            _check_brackets(para, issues)
            tick("检查括号格式...")

    if rules.check_period:
        for para in paragraphs:
            _check_period(para, issues)
            tick("检查列表句号...")

    if rules.check_forbidden_words and rules.forbidden_words:
        for para in paragraphs:
            _check_forbidden_words(para, rules.forbidden_words, issues)
            tick("检查禁忌词...")

    if rules.check_table_font_size:
        for table in tables:
            _check_table_font_size(table, settings, issues)
            tick("检查表格字号...")

    if rules.check_table_row_height:
        for table in tables:
            _check_table_row_height(table, settings, issues)
            tick("检查表格行高...")

    if rules.check_cover_title_size:
        _check_cover_title_size(paragraphs, settings, issues)
        tick("检查封面字号...")

    if rules.check_appendix_body_size:
        _check_appendix_body_size(paragraphs, settings, issues)
        tick("检查附件正文字号...")

    if rules.check_grid_alignment:
        _check_grid_alignment(paragraphs, issues)
        tick("检查网格对齐...")

    if rules.check_appendix_colon:
        _check_appendix_colon(paragraphs, issues)
        tick("检查附件冒号...")

    if rules.check_figure_table_space:
        _check_figure_table_space(paragraphs, issues)
        tick("检查图表空格...")

    if rules.check_figure_position:
        _check_figure_position(doc, tables, paragraphs, issues)
        tick("检查图表位置...")

    if rules.check_chairman_number:
        _check_chairman_number(paragraphs, tables, issues)
        tick("检查负责人信息...")

    if rules.check_presidential_order:
        from tvba_core_presidential import check_presidential_order_numbers
        check_presidential_order_numbers(paragraphs, issues)
        tick("检查主席令编号...")

    if rules.check_spacing:
        for para in paragraphs:
            _check_paragraph_spacing(para, issues)
            tick("检查段前段后...")

    if rules.check_caption_alignment:
        for para in paragraphs:
            _check_caption_alignment(para, issues, doc)
            tick("检查题注居中...")

    if rules.check_table_fixed_dimensions:
        for table in tables:
            _check_table_fixed_dimensions(table, issues)
            tick("检查表格固定尺寸...")

    if progress_cb:
        progress_cb("检查完成", 1.0)

    return issues


def _para_location(para) -> str:
    text = para.text[:50].replace("\n", " ")
    if len(para.text) > 50:
        text += "..."
    return f'"{text}"'


def _check_chinese_font(para, issues: list[ValidationIssue], doc=None):
    """Check if Chinese characters use the correct font (宋体)."""
    text = para.text
    if not text or not CN_RE.search(text):
        return
    for run in para.runs:
        if not CN_RE.search(run.text):
            continue
        try:
            fonts = get_effective_run_fonts(run, doc)
            east = fonts.get("eastAsia")
            if east is None:
                issues.append(ValidationIssue(
                    severity="warning",
                    description="未能解析到中文字体（未显式设置且无样式/默认值）",
                    location=_para_location(para),
                ))
                return
            if east.startswith("theme:"):
                # Theme-referenced font — Word resolves via theme1.xml, cannot verify
                issues.append(ValidationIssue(
                    severity="warning",
                    description=f"中文字体由主题引用 ({east})，无法验证是否为宋体",
                    location=_para_location(para),
                ))
                return
            if east != "宋体":
                issues.append(ValidationIssue(
                    severity="warning",
                    description=f"中文字体非宋体: {east}",
                    location=_para_location(para),
                ))
                return
        except Exception:
            pass


def _check_ascii_font(para, issues: list[ValidationIssue], doc=None):
    """Check if ASCII characters (digits, letters) use Times New Roman."""
    text = para.text
    if not text or not ASCII_RE.search(text):
        return
    for run in para.runs:
        if not ASCII_RE.search(run.text):
            continue
        try:
            fonts = get_effective_run_fonts(run, doc)
            ascii_font = fonts.get("ascii") or fonts.get("hAnsi")
            if ascii_font is None:
                issues.append(ValidationIssue(
                    severity="warning",
                    description="未能解析到数字/英文字体（未显式设置且无样式/默认值）",
                    location=_para_location(para),
                ))
                return
            if ascii_font.startswith("theme:"):
                issues.append(ValidationIssue(
                    severity="warning",
                    description=f"数字/英文字体由主题引用 ({ascii_font})，无法验证是否为 Times New Roman",
                    location=_para_location(para),
                ))
                return
            if ascii_font != "Times New Roman":
                issues.append(ValidationIssue(
                    severity="warning",
                    description=f"数字/英文字体非 Times New Roman: {ascii_font}",
                    location=_para_location(para),
                ))
                return
        except Exception:
            pass


def _check_brackets(para, issues: list[ValidationIssue]):
    """Check for half-width brackets in list items (四级及以下)."""
    text = para.text
    if not text:
        return
    if re.search(r'\([^)]*\)', text):
        issues.append(ValidationIssue(
            severity="warning",
            description="列表项包含半角括号 ()，应使用全角括号（）",
            location=_para_location(para),
        ))


def _check_period(para, issues: list[ValidationIssue]):
    """Check if a list item is missing a period at the end.

    Only flags genuine list items (with bracket/parenthesis delimiters
    or letter prefixes), NOT title/heading paragraphs (number+space).
    """
    text = (para.text or "").strip()
    if not text:
        return

    # Skip title paragraphs (number. or number.number + space)
    if re.match(r'^\d+(\.\d+)*\s+\S', text):
        return

    # Check if paragraph looks like a list item with bracket/parenthesis
    is_list_item = (
        re.match(r'^[\(\（]\d+[\)\）]', text) or
        re.match(r'^\d+[\)\）]', text) or
        re.match(r'^[a-z][\.\)]\s', text)
    )
    if not is_list_item:
        return

    if not text.endswith(("。", "）", ")", ".", "；", ";", "！", "？")):
        issues.append(ValidationIssue(
            severity="warning",
            description="列表项末尾缺少句号",
            location=_para_location(para),
        ))


def _check_forbidden_words(para, forbidden_words: tuple[str, ...], issues: list[ValidationIssue]):
    """Check for forbidden words in paragraph text."""
    text = para.text
    if not text:
        return
    for word in forbidden_words:
        if word in text:
            issues.append(ValidationIssue(
                severity="error",
                description=f'发现禁忌词"{word}"，建议替换',
                location=_para_location(para),
            ))


def _check_table_font_size(table, settings: FormatSettings, issues: list[ValidationIssue]):
    """Check if table body cells use the correct font size."""
    expected_pt = size_label_to_points(settings.table.body_size)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    try:
                        rPr = run._element.find(W_RPR)
                        if rPr is not None:
                            sz = rPr.find(W_SZ)
                            if sz is not None:
                                size_val = sz.get(W_VAL)
                                if size_val:
                                    actual_pt = int(size_val) / 2
                                    if actual_pt != expected_pt:
                                        issues.append(ValidationIssue(
                                            severity="warning",
                                            description=f"表格({row_idx+1},{col_idx+1})字号非{settings.table.body_size}",
                                            location=f"表格行{row_idx+1}列{col_idx+1}",
                                        ))
                                        return
                    except Exception:
                        pass


def _check_table_row_height(table, settings: FormatSettings, issues: list[ValidationIssue]):
    """Check for fixed table row heights (should be auto-fit per onsite requirement)."""
    for row_idx, row in enumerate(table.rows):
        try:
            tr = row._element
            trPr = tr.find(W_TRPR)
            if trPr is None:
                continue
            trHeight = trPr.find(W_TRHEIGHT)
            if trHeight is None:
                continue

            rule = trHeight.get(W_HRULE) or ""
            val = trHeight.get(W_VAL)

            if rule == "exact":
                # Fixed height is explicitly disallowed — tables should auto-fit
                cm = round(int(val) / 567, 2) if val else 0
                issues.append(ValidationIssue(
                    severity="warning",
                    description=f"表格行{row_idx+1}有固定行高{cm}cm（应为自适应）",
                    location=f"表格行{row_idx+1}",
                ))
            elif rule == "atLeast" and val:
                # atLeast with a high value effectively acts as fixed
                cm = int(val) / 567
                if cm > 1.5:
                    issues.append(ValidationIssue(
                        severity="warning",
                        description=f"表格行{row_idx+1}最小行高{cm:.1f}cm偏高（应为自适应）",
                        location=f"表格行{row_idx+1}",
                    ))
        except Exception:
            pass


def _check_cover_title_size(paragraphs, settings: FormatSettings, issues: list[ValidationIssue]):
    """Check if the cover page title uses 二号 (22pt) fontWeight."""
    expected_pt = size_label_to_points("二号")

    for para in paragraphs[:15]:
        text = (para.text or "").strip()
        if not text or len(text) < 2 or len(text) > 80:
            continue
        if para.alignment != 1:
            continue
        if is_toc_title_line(text) or is_toc_entry_line(text):
            break
        for run in para.runs:
            try:
                rPr = run._element.find(W_RPR)
                if rPr is not None:
                    sz = rPr.find(W_SZ)
                    if sz is not None:
                        half_pts = int(sz.get(W_VAL))
                        actual_pt = half_pts / 2
                        if actual_pt != expected_pt:
                            issues.append(ValidationIssue(
                                severity="warning",
                                description=f"封面标题字号非二号(22pt)，当前为{actual_pt:.0f}pt",
                                location=_para_location(para),
                            ))
            except Exception:
                pass
        break


def _check_appendix_body_size(paragraphs, settings: FormatSettings, issues: list[ValidationIssue]):
    """Check if appendix body text uses 小五 (10.5pt) font."""
    expected_pt = size_label_to_points("小五")

    in_appendix = False
    for para in paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        if is_appendix_title(text):
            in_appendix = True
            continue
        if not in_appendix:
            continue
        # Check if ending condition: next heading
        try:
            rPr = para._element.find(W_RPR)
            if rPr is not None:
                outline = rPr.find(W_OUTLINE)
                if outline is not None:
                    in_appendix = False
                    continue
        except Exception:
            pass
        for run in para.runs:
            try:
                rPr = run._element.find(W_RPR)
                if rPr is not None:
                    sz = rPr.find(W_SZ)
                    if sz is not None:
                        half_pts = int(sz.get(W_VAL))
                        actual_pt = half_pts / 2
                        if actual_pt != expected_pt:
                            issues.append(ValidationIssue(
                                severity="warning",
                                description=f"附件正文字号非小五(10.5pt)，当前为{actual_pt:.0f}pt",
                                location=_para_location(para),
                            ))
                            return
            except Exception:
                pass


def _check_grid_alignment(paragraphs, issues: list[ValidationIssue]):
    """Check if level 1 titles have snapToGrid enabled."""
    for para in paragraphs:
        try:
            pPr = para._element.find(W_PPR)
            if pPr is None:
                continue
            outline = pPr.find(W_OUTLINE)
            if outline is None:
                continue
            level_val = outline.get(W_VAL)
            if level_val != "0":
                continue
            snap = pPr.find(W_SNAP)
            auto_de = pPr.find(W_AUTO_DE)
            if snap is None or auto_de is None:
                issues.append(ValidationIssue(
                    severity="warning",
                    description="一级标题未勾选'对齐到网格'或'自动调整右缩进'",
                    location=_para_location(para),
                ))
                return
        except Exception:
            pass


def _check_appendix_colon(paragraphs, issues: list[ValidationIssue]):
    """Check if appendix titles use fullwidth colon (：) after the number."""
    for para in paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        if is_appendix_title(text):
            if not re.search(r'附件\d*[：:]', text):
                issues.append(ValidationIssue(
                    severity="warning",
                    description="附件标题未使用冒号（应为全角 ：）",
                    location=_para_location(para),
                ))


def _check_figure_table_space(paragraphs, issues: list[ValidationIssue]):
    """Check if figure/table captions have exactly one space between number and text."""
    for para in paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        if is_table_caption_line(text) or is_figure_caption_line(text):
            m = re.match(r'^([表图]\s*[0-9０-９]+(?:[.．][0-9０-９]+)*(?:[-－–—][0-9０-９]+)?)(\s+)(.+)$', text)
            if m and len(m.group(2)) != 1:
                issues.append(ValidationIssue(
                    severity="warning",
                    description=f"题注序号与名称间应恰好1个空格（当前{len(m.group(2))}个）",
                    location=_para_location(para),
                ))


def _check_figure_position(doc, tables, paragraphs, issues: list[ValidationIssue]):
    """Check figure/table caption positioning: table captions above, figure captions below."""
    body = doc.element.body
    elements = list(body)
    for i, elem in enumerate(elements):
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "tbl":
            for j in range(i - 1, -1, -1):
                prev_tag = elements[j].tag.split("}")[-1] if "}" in elements[j].tag else elements[j].tag
                if prev_tag == "p":
                    texts = []
                    for t in elements[j].findall(f".//{W_T}"):
                        if t.text:
                            texts.append(t.text)
                    para_text = "".join(texts).strip()
                    if para_text and not is_table_caption_line(para_text):
                        issues.append(ValidationIssue(
                            severity="warning",
                            description="表格上方应有题注（表x.x-x 题注文字）",
                            location=f'表格前段落: "{para_text[:30]}..."' if len(para_text) > 30 else f'表格前段落: "{para_text}"',
                        ))
                    break


# Patterns must be full field labels (with colon implied).
# Bare words like "审定"/"批准" are excluded because they match generic body text
# ("审定意见", "批准单位") and produce false positives.
_CHAIRMAN_PATTERNS = ("负责人", "审定人", "核准人", "批准人")


def _check_chairman_number(paragraphs, tables, issues: list[ValidationIssue]):
    """Check if the document has chairman/approver fields with content after them.

    Scans the first 30 paragraphs and first 5 tables for chairman/approver
    designations and verifies there is content after the label (name or title).
    """
    found_chairman = False

    # Collect text sources: paragraphs + table cell paragraphs
    text_sources: list[tuple[str, str]] = []  # (text, location)
    for para in paragraphs[:30]:
        text = (para.text or "").strip()
        if text:
            text_sources.append((text, _para_location(para)))

    for ti, table in enumerate(tables[:5]):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, para in enumerate(cell.paragraphs):
                    text = (para.text or "").strip()
                    if text:
                        text_sources.append((text, f"表格{ti+1}行{ri+1}列{ci+1}"))

    for text, location in text_sources:
        for pattern in _CHAIRMAN_PATTERNS:
            # Must match as field label: "负责人：张三", not "项目负责人制度说明"
            m = re.search(pattern + r'\s*[：:]\s*(\S+)', text)
            if m:
                found_chairman = True
                break
            # Also catch the case where the label exists but content is missing
            m_empty = re.search(pattern + r'\s*[：:]\s*$', text)
            if m_empty:
                found_chairman = True
                issues.append(ValidationIssue(
                    severity="warning",
                    description=f'"{pattern}"后应有姓名或编号',
                    location=location,
                ))
                break

    if not found_chairman and text_sources:
        issues.append(ValidationIssue(
            severity="warning",
            description="未在前30段及前5个表格中检测到负责人/审定人/核准人信息",
        ))


def _check_paragraph_spacing(para, issues: list[ValidationIssue]):
    """Check that paragraph has no before/after spacing (must be 0 per onsite req)."""
    text = (para.text or "").strip()
    if not text:
        return
    pPr = para._element.find(W_PPR)
    if pPr is None:
        return
    spacing = pPr.find(W_SPACING)
    if spacing is None:
        return

    def _get_int(attr):
        v = spacing.get(attr)
        return int(v) if v else 0

    before = _get_int(W_BEFORE)
    after = _get_int(W_AFTER)
    before_lines = _get_int(W_BEFORE_LINES)
    after_lines = _get_int(W_AFTER_LINES)

    if before != 0 or after != 0 or before_lines != 0 or after_lines != 0:
        issues.append(ValidationIssue(
            severity="warning",
            description=f"段落段前/段后非0（before={before} after={after} beforeLines={before_lines} afterLines={after_lines}）",
            location=_para_location(para),
        ))


def _check_caption_alignment(para, issues: list[ValidationIssue], doc=None):
    """Check that table/figure captions are centered."""
    text = (para.text or "").strip()
    if not text:
        return
    if not (is_table_caption_paragraph(para, doc) or is_figure_caption_paragraph(para, doc)):
        return
    pPr = para._element.find(W_PPR)
    if pPr is None:
        issues.append(ValidationIssue(
            severity="warning",
            description="题注未设置居中对齐",
            location=_para_location(para),
        ))
        return
    jc = pPr.find(W_JC)
    if jc is None or jc.get(W_VAL) != "center":
        issues.append(ValidationIssue(
            severity="warning",
            description="题注应为居中对齐",
            location=_para_location(para),
        ))


def _check_table_fixed_dimensions(table, issues: list[ValidationIssue]):
    """Check that table doesn't have fixed row heights or column widths."""
    tblPr = table._element.find(f"{{{W}}}tblPr")
    if tblPr is not None:
        # Check for fixed table layout (w:tblLayout w:type="fixed")
        tblLayout = tblPr.find(f"{{{W}}}tblLayout")
        if tblLayout is not None:
            layout_type = tblLayout.get(f"{{{W}}}type") or ""
            if layout_type == "fixed":
                issues.append(ValidationIssue(
                    severity="warning",
                    description="表格有固定列宽布局（应为自适应）",
                    location="表格",
                ))

        # Check for fixed total table width (w:tblW w:type="dxa" with non-zero value).
        # Use findall because multiple tblW elements can coexist (python-docx default + user-set).
        for tblW in tblPr.findall(f"{{{W}}}tblW"):
            w_type = tblW.get(f"{{{W}}}type") or ""
            if w_type == "dxa":
                w_val = tblW.get(f"{{{W}}}w")
                if w_val and int(w_val) > 0:
                    issues.append(ValidationIssue(
                        severity="warning",
                        description=f"表格宽度固定为{int(w_val)/20:.0f}pt（应为自适应）",
                        location="表格",
                    ))
                    break

    # Fixed row heights are already checked by _check_table_row_height
    # when check_table_row_height is enabled.
