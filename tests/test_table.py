import pytest
from docx import Document
from lxml import etree

from tvba_core_table import (
    is_table_caption_line,
    is_table_caption_paragraph,
    find_table_caption,
    apply_table_caption,
    apply_table_body,
    refresh_all,
)
from tvba_settings import TableSettings

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_NS = {"w": W}
W_VAL = f"{{{W}}}val"


def _ppr(para):
    pPr = para._element.find("w:pPr", W_NS)
    if pPr is None:
        pPr = etree.SubElement(para._element, f"{{{W}}}pPr")
    return pPr


def _set_dirty_caption_paragraph_format(para):
    pPr = _ppr(para)
    jc = pPr.find("w:jc", W_NS)
    if jc is None:
        jc = etree.SubElement(pPr, f"{{{W}}}jc")
    jc.set(W_VAL, "both")
    outline = pPr.find("w:outlineLvl", W_NS)
    if outline is None:
        outline = etree.SubElement(pPr, f"{{{W}}}outlineLvl")
    outline.set(W_VAL, "5")
    ind = pPr.find("w:ind", W_NS)
    if ind is None:
        ind = etree.SubElement(pPr, f"{{{W}}}ind")
    ind.set(f"{{{W}}}left", "2126")
    ind.set(f"{{{W}}}right", "567")
    ind.set(f"{{{W}}}firstLine", "482")


def _add_numbering_level(doc, *, num_id="902", abstract_id="902", ilvl="6", lvl_text="表%1.%2.%3-%7  "):
    numbering = doc.part.numbering_part._element
    abstract = etree.SubElement(numbering, f"{{{W}}}abstractNum")
    abstract.set(f"{{{W}}}abstractNumId", abstract_id)
    lvl = etree.SubElement(abstract, f"{{{W}}}lvl")
    lvl.set(f"{{{W}}}ilvl", ilvl)
    start = etree.SubElement(lvl, f"{{{W}}}start")
    start.set(W_VAL, "1")
    num_fmt = etree.SubElement(lvl, f"{{{W}}}numFmt")
    num_fmt.set(W_VAL, "decimal")
    lvl_text_elem = etree.SubElement(lvl, f"{{{W}}}lvlText")
    lvl_text_elem.set(W_VAL, lvl_text)

    num = etree.SubElement(numbering, f"{{{W}}}num")
    num.set(f"{{{W}}}numId", num_id)
    abstract_ref = etree.SubElement(num, f"{{{W}}}abstractNumId")
    abstract_ref.set(W_VAL, abstract_id)
    return num_id


def _set_paragraph_num_pr(para, *, num_id="902", ilvl="6"):
    pPr = _ppr(para)
    num_pr = pPr.find("w:numPr", W_NS)
    if num_pr is None:
        num_pr = etree.SubElement(pPr, f"{{{W}}}numPr")
    ilvl_elem = num_pr.find("w:ilvl", W_NS)
    if ilvl_elem is None:
        ilvl_elem = etree.SubElement(num_pr, f"{{{W}}}ilvl")
    ilvl_elem.set(W_VAL, ilvl)
    num_id_elem = num_pr.find("w:numId", W_NS)
    if num_id_elem is None:
        num_id_elem = etree.SubElement(num_pr, f"{{{W}}}numId")
    num_id_elem.set(W_VAL, num_id)

class TestIsTableCaptionLine:
    def test_biao_with_number_pattern(self):
        assert is_table_caption_line("表 1.1-1 示例表格") is True

    def test_biao_no_space(self):
        assert is_table_caption_line("表1.1-1 示例表格") is True

    def test_biao_simple_number(self):
        assert is_table_caption_line("表 1-1 示例表格") is True

    def test_starts_with_table(self):
        assert is_table_caption_line("Table 1-1 Example") is True

    def test_table_with_dot_number(self):
        assert is_table_caption_line("Table 1.1-1 Example") is True

    def test_no_prefix(self):
        assert is_table_caption_line("示例表格") is False

    def test_case_insensitive(self):
        assert is_table_caption_line("TABLE 1-1 Example") is True

    def test_simple_number_matches(self):
        assert is_table_caption_line("表 1 示例表格") is True

    def test_no_caption_text_fails(self):
        assert is_table_caption_line("表 1-1") is False

    def test_only_prefix_fails(self):
        assert is_table_caption_line("表 ") is False

    def test_table_without_number_pattern_fails(self):
        assert is_table_caption_line("Table Example") is False

    def test_fullwidth_space_separator(self):
        """Full-width space (U+3000) between number and text should match."""
        assert is_table_caption_line("表2.2.9-1　沿线经过环境敏感点统计表") is True

    def test_nonbreaking_space_separator(self):
        """Non-breaking space (U+00A0) between number and text should match."""
        assert is_table_caption_line("表2.2.9-1\xa0沿线经过环境敏感点统计表") is True

    def test_multi_dot_number_with_double_space(self):
        """Real captions like 表1.8.12-2  should be recognized."""
        assert is_table_caption_line("表1.8.12-2  本工程车辆配置一览表") is True

    @pytest.mark.parametrize(
        "caption",
        [
            "表1-1  测试表格",
            "表 1-1  测试表格",
            "表1.8-2  测试表格",
            "表 1.8-2  测试表格",
            "表1.8.4-2  本工程珠海站场和阀室设置",
            "表 1.8.4-2  本工程珠海站场和阀室设置",
            "表1.8.12-2  本工程车辆配置一览表",
            "表 1.8.12-2  本工程车辆配置一览表",
            "表１．８．１２－２  本工程车辆配置一览表",
            "表1.8.12—2  本工程车辆配置一览表",
        ],
    )
    def test_table_caption_number_category_matches(self, caption):
        """表 + optional space + dot-number + hyphen suffix captions are one category."""
        assert is_table_caption_line(caption) is True


class TestIsTableCaptionParagraph:
    def test_direct_numbering_with_table_label_is_caption(self):
        doc = Document()
        _add_numbering_level(doc)
        para = doc.add_paragraph("项目整体进度安排表")
        _set_paragraph_num_pr(para)

        assert is_table_caption_paragraph(para, doc) is True

    def test_direct_numbering_with_figure_label_is_not_table_caption(self):
        doc = Document()
        _add_numbering_level(doc, lvl_text="图%1.%2.%3-%7  ")
        para = doc.add_paragraph("工艺系统示意图")
        _set_paragraph_num_pr(para)

        assert is_table_caption_paragraph(para, doc) is False


class TestApplyTableCaption:
    def test_applies_font_and_bold(self):
        doc = Document()
        para = doc.add_paragraph("表 1.1-1 示例")
        settings = TableSettings(title_font="黑体", title_size="小四", title_bold=True)
        apply_table_caption(para, settings)
        run = para.runs[0]
        assert run.font.bold is True

    def test_multi_dot_caption_is_centered(self):
        doc = Document()
        para = doc.add_paragraph("表1.8.12-2  本工程车辆配置一览表")
        _set_dirty_caption_paragraph_format(para)
        settings = TableSettings(title_alignment="居中")

        apply_table_caption(para, settings)

        pPr = para._element.find("w:pPr", W_NS)
        jc = pPr.find("w:jc", W_NS)
        assert jc.get(W_VAL) == "center"

    @pytest.mark.parametrize(
        "caption",
        [
            "表1-1  测试表格",
            "表 1-1  测试表格",
            "表1.8-2  测试表格",
            "表 1.8-2  测试表格",
            "表1.8.4-2  本工程珠海站场和阀室设置",
            "表 1.8.4-2  本工程珠海站场和阀室设置",
            "表1.8.12-2  本工程车辆配置一览表",
            "表 1.8.12-2  本工程车辆配置一览表",
            "表１．８．１２－２  本工程车辆配置一览表",
            "表1.8.12—2  本工程车辆配置一览表",
        ],
    )
    def test_table_caption_category_clears_dirty_alignment_outline_and_indent(self, caption):
        doc = Document()
        para = doc.add_paragraph(caption)
        _set_dirty_caption_paragraph_format(para)
        settings = TableSettings()

        apply_table_caption(para, settings)

        pPr = para._element.find("w:pPr", W_NS)
        jc = pPr.find("w:jc", W_NS)
        outline = pPr.find("w:outlineLvl", W_NS)
        ind = pPr.find("w:ind", W_NS)
        assert jc.get(W_VAL) == "center"
        assert outline is None
        assert ind.get(f"{{{W}}}left") == "0"
        assert ind.get(f"{{{W}}}right") == "0"
        assert ind.get(f"{{{W}}}firstLine") is None
        assert ind.get(f"{{{W}}}hanging") is None

class TestApplyTableBody:
    def test_sets_borders(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        settings = TableSettings(line_width_pt=0.25, auto_fit_mode="window")
        apply_table_body(table, settings)
        tblPr = table._element.find(".//w:tblPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        borders = tblPr.find("w:tblBorders", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        assert borders is not None
        assert borders.find("w:top", W_NS).get(f"{{{W}}}sz") == "2"

    def test_applies_configured_row_height(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        settings = TableSettings(row_height_cm=0.6)

        apply_table_body(table, settings)

        tr_height = table.rows[0]._tr.find(".//w:trHeight", W_NS)
        assert tr_height.get(W_VAL) == "340"
        assert tr_height.get(f"{{{W}}}hRule") == "atLeast"

class TestRefreshAll:
    def test_finds_and_formats_table(self):
        doc = Document()
        doc.add_paragraph("表 1.1-1 测试表格")
        table = doc.add_table(rows=2, cols=2)
        settings = TableSettings()
        refresh_all(doc, settings)
        # Caption should be formatted
        para = doc.paragraphs[0]
        assert para.runs[0].font.bold is True

    def test_formats_standalone_multi_dot_caption(self):
        doc = Document()
        para = doc.add_paragraph("表1.8.12-2  本工程车辆配置一览表")
        para.alignment = 3
        settings = TableSettings(title_alignment="居中")

        refresh_all(doc, settings)

        pPr = para._element.find(".//w:pPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        jc = pPr.find("w:jc", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        assert jc.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "center"

    def test_formats_auto_numbered_caption_without_literal_table_prefix(self):
        doc = Document()
        _add_numbering_level(doc)
        para = doc.add_paragraph("项目整体进度安排表")
        _set_dirty_caption_paragraph_format(para)
        _set_paragraph_num_pr(para)
        doc.add_table(rows=2, cols=2)
        settings = TableSettings(title_alignment="居中")

        refresh_all(doc, settings)

        pPr = para._element.find(".//w:pPr", W_NS)
        jc = pPr.find("w:jc", W_NS)
        outline = pPr.find("w:outlineLvl", W_NS)
        assert jc.get(W_VAL) == "center"
        assert outline is None
        ind = pPr.find("w:ind", W_NS)
        assert ind.get(f"{{{W}}}left") == "0"
